#!/usr/bin/env python3
"""Generate Protocol of Disagreements (Протокол разногласий) as DOCX."""

import json
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Load JSON data
json_path = "/home/z/my-project/download/protocol_v2_analysis.json"
with open(json_path, "r", encoding="utf-8") as f:
    disagreements = json.load(f)

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15

# ===== TITLE =====
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_after = Pt(6)
title_run = title_para.add_run("ПРОТОКОЛ РАЗНОГЛАСИЙ")
title_run.bold = True
title_run.font.size = Pt(16)
title_run.font.name = 'Times New Roman'

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_para.paragraph_format.space_after = Pt(4)
subtitle_run = subtitle_para.add_run('к Договору генерального подряда №-___')
subtitle_run.bold = True
subtitle_run.font.size = Pt(13)
subtitle_run.font.name = 'Times New Roman'

sub2_para = doc.add_paragraph()
sub2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2_para.paragraph_format.space_after = Pt(12)
sub2_run = sub2_para.add_run('г. Москва')
sub2_run.font.size = Pt(11)
sub2_run.font.name = 'Times New Roman'

# Preamble
preamble = doc.add_paragraph()
preamble.paragraph_format.space_after = Pt(12)
preamble.paragraph_format.first_line_indent = Cm(1.25)
preamble_text = (
    "Настоящий Протокол разногласий составлен в связи с предложением ООО «[Генподрядчик]» "
    "(далее — «Генподрядчик») об изменении условий Договора генерального подряда №-___ "
    "(далее — «Договор»), заключённого между ООО «ИНВЕСТСТРОЙГРУПП» (далее — «Заказчик») "
    "и Генподрядчиком."
)
run = preamble.add_run(preamble_text)
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

# Severity legend
legend = doc.add_paragraph()
legend.paragraph_format.space_after = Pt(12)
legend_run = legend.add_run(
    "Уровни критичности: CRITICAL — условие создаёт экзистенциальный риск для Генподрядчика; "
    "HIGH — условие создаёт значительный финансовый риск; "
    "MEDIUM — условие создаёт операционный риск."
)
legend_run.font.size = Pt(9)
legend_run.font.name = 'Times New Roman'
legend_run.italic = True

# ===== MAIN TABLE =====
# Count items
total = len(disagreements)
critical_count = sum(1 for d in disagreements if d["severity"] == "CRITICAL")
high_count = sum(1 for d in disagreements if d["severity"] == "HIGH")
medium_count = sum(1 for d in disagreements if d["severity"] == "MEDIUM")

# Create table: columns = № | Пункт | Редакция Заказчика | Редакция Генподрядчика
# Plus justification row below each entry
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Set column widths
widths = [Cm(1.0), Cm(2.0), Cm(6.5), Cm(6.5)]
for i, width in enumerate(widths):
    table.columns[i].width = width

# Header row
header_cells = table.rows[0].cells
headers = ["№", "Пункт Договора", "Редакция Заказчика", "Редакция Генподрядчика"]
for i, header_text in enumerate(headers):
    cell = header_cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header_text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    # Gray header background
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# Fill data
for d in disagreements:
    # Main data row
    row = table.add_row()
    cells = row.cells

    # №
    cells[0].text = ""
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(d["id"]))
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    # Пункт
    cells[1].text = ""
    p = cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(d["clause"])
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.bold = True

    # Severity indicator
    severity_text = f"\n[{d['severity']}]"
    sev_run = p.add_run(severity_text)
    sev_run.font.size = Pt(8)
    sev_run.font.name = 'Times New Roman'
    if d["severity"] == "CRITICAL":
        sev_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        sev_run.bold = True
    elif d["severity"] == "HIGH":
        sev_run.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
        sev_run.bold = True
    else:
        sev_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    # Customer wording
    cells[2].text = ""
    p = cells[2].paragraphs[0]
    run = p.add_run(d["customer_wording"])
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    # Contractor proposal
    cells[3].text = ""
    p = cells[3].paragraphs[0]
    run = p.add_run(d["contractor_proposal"])
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.bold = True

    # Justification row (merged across all columns)
    just_row = table.add_row()
    just_cells = just_row.cells
    # Merge all cells
    merged_cell = just_cells[0].merge(just_cells[3])
    merged_cell.text = ""
    p = merged_cell.paragraphs[0]

    just_label = p.add_run("Обоснование: ")
    just_label.font.size = Pt(8)
    just_label.font.name = 'Times New Roman'
    just_label.bold = True
    just_label.italic = True

    just_text = p.add_run(d["justification"])
    just_text.font.size = Pt(8)
    just_text.font.name = 'Times New Roman'
    just_text.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    just_text.italic = True

    # Light background for justification row
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    merged_cell._tc.get_or_add_tcPr().append(shading)

# ===== SUMMARY SECTION =====
doc.add_paragraph()
summary_title = doc.add_paragraph()
summary_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = summary_title.add_run("СВОДКА АНАЛИЗА")
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

summary_para = doc.add_paragraph()
summary_para.paragraph_format.first_line_indent = Cm(1.25)
summary_text = (
    f"Всего выявлено пунктов разногласий: {total}\n"
    f"Из них:\n"
    f"  • CRITICAL (экзистенциальный риск): {critical_count}\n"
    f"  • HIGH (значительный финансовый риск): {high_count}\n"
    f"  • MEDIUM (операционный риск): {medium_count}\n\n"
    f"Предыдущий протокол содержал 18 пунктов. Настоящий протокол содержит {total} пунктов, "
    f"что на {total - 18} пунктов больше.\n\n"
    f"Новые ключевые находки, не вошедшие в предыдущий протокол:\n"
    f"  1. п. 1.7 — Обязательные условия Wildberries с приоритетом над Договором (CRITICAL)\n"
    f"  2. п. 3.4 — Бросовые работы от изменений Заказчика включены в твёрдую цену (CRITICAL)\n"
    f"  3. п. 4.5.1 — Убытки при приостановке по вине Заказчика ограничены 0,01% (CRITICAL)\n"
    f"  4. п. 10.2 — Заранее оценённые убытки 720 млн руб./мес., неприменимость ст. 333 ГК РФ (CRITICAL)\n"
    f"  5. п. 12.11 — Запрет на отказ от Договора по ст. 719 ГК РФ (CRITICAL)"
)
run = summary_para.add_run(summary_text)
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

# ===== SIGNATURE BLOCK =====
doc.add_paragraph()
doc.add_paragraph()

sig_title = doc.add_paragraph()
sig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sig_title.add_run("ПОДПИСИ СТОРОН")
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# Create signature table (borderless)
sig_table = doc.add_table(rows=6, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Remove borders
for row in sig_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

# Fill signature data
sig_data = [
    ("ЗАКАЗЧИК:", "ГЕНПОДРЯДЧИК:"),
    ("ООО «ИНВЕСТСТРОЙГРУПП»", "ООО «[Наименование]»"),
    ("", ""),
    ("_________________ / Грачёв И.В. /", "_________________ / _______________ /"),
    ("М.П.", "М.П."),
    ("Дата: «___» __________ 202__ г.", "Дата: «___» __________ 202__ г."),
]

for i, (left_text, right_text) in enumerate(sig_data):
    left_cell = sig_table.rows[i].cells[0]
    right_cell = sig_table.rows[i].cells[1]

    left_cell.text = ""
    right_cell.text = ""

    p_left = left_cell.paragraphs[0]
    p_right = right_cell.paragraphs[0]

    run_left = p_left.add_run(left_text)
    run_right = p_right.add_run(right_text)

    for run in [run_left, run_right]:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        if i == 0:
            run.bold = True

# Save
output_path = "/home/z/my-project/download/Протокол_разногласий_v2.docx"
doc.save(output_path)
print(f"DOCX saved to: {output_path}")
print(f"Total disagreements: {total}")
print(f"CRITICAL: {critical_count}, HIGH: {high_count}, MEDIUM: {medium_count}")
