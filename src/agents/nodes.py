"""
ASD v12.0 — Agent Nodes for LangGraph Workflow.

Each node represents an agent in the ASD pipeline.
All LLM calls go through llm_engine (unified interface).

Model Lineup (mac_studio):
    PM:     Llama 3.3 70B 4-bit
    ПТО:    Gemma 4 31B 4-bit (VLM, shared)
    Юрист:  Gemma 4 31B 4-bit (shared)
    Сметчик:Gemma 4 31B 4-bit (shared)
    Закупщик:Gemma 4 31B 4-bit (shared)
    Логист: Gemma 4 31B 4-bit (shared)
    Дело:   Gemma 4 E4B 4-bit
"""

import json
import logging
from datetime import datetime
from typing import Dict, Any, List, Optional
from src.core.llm_engine import llm_engine
from src.core.exceptions import (
    LLMUnavailableError,
    LLMResponseError,
    NetworkError,
    DocumentParseError,
)
from src.core.pm_agent import (
    compute_weighted_score,
    check_veto_rules,
    extract_legal_signal,
    extract_smeta_signal,
    extract_pto_signal,
    extract_procurement_signal,
    extract_logistics_signal,
    calculate_risk_level,
    DEFAULT_VETO_RULES,
)
import docx
import fitz  # PyMuPDF
import os

logger = logging.getLogger(__name__)
from src.db.models import AuditLog
from src.db.init_db import create_engine
from sqlalchemy.orm import sessionmaker
from src.config import settings
from src.utils.wiki_loader import load_wiki_page
from src.core.lessons_service import lessons_service
from src.core.rag_pipeline import rag_pipeline

# Setup DB session for Audit Logging
engine = create_engine(settings.database_url)
Session = sessionmaker(bind=engine)

# Orchestrator logic merged into pm_agent.py — ProjectManager handles
# weighted scoring + veto rules + LLM reasoning via PM Agent.


# =============================================================================
# LLM Fallback Tracking
# =============================================================================

async def _safe_agent_chat(
    agent: str,
    messages: List,
    state: Dict[str, Any],
    fallback_response: str = "{}",
    **kwargs,
) -> str:
    """
    Wrapper around llm_engine.safe_chat that records fallback events.
    
    If LLM is unavailable and fallback_response is used, this sets
    _llm_fallback_triggered=True in the state and logs the affected agent.
    """
    try:
        response = await llm_engine.chat(agent, messages, **kwargs)
        return response
    except (LLMUnavailableError, NetworkError) as e:
        logger.warning(
            "LLM unavailable for agent %s: %s. Using fallback response.",
            agent, str(e)
        )
        state["_llm_fallback_triggered"] = True
        fallback_agents = state.setdefault("_llm_fallback_agents", [])
        if agent not in fallback_agents:
            fallback_agents.append(agent)
        return fallback_response
    except Exception as e:
        logger.error(
            "Unexpected error in agent %s: %s. Using fallback response.",
            agent, str(e)
        )
        state["_llm_fallback_triggered"] = True
        fallback_agents = state.setdefault("_llm_fallback_agents", [])
        if agent not in fallback_agents:
            fallback_agents.append(agent)
        return fallback_response


def _compute_agent_confidence(
    state: Dict[str, Any],
    agent: str,
    result_text: str,
    parsed_data: Optional[Dict] = None,
) -> float:
    """
    Compute agent confidence based on response quality.
    
    Factors:
    - 0.00: LLM fallback was triggered (result is dummy data)
    - 0.30: Response couldn't be parsed as JSON
    - 0.50: Parsed but very short / minimal
    - 0.70: Parsed with expected structure
    - 0.90: Parsed with rich structure and expected keys present
    
    Args:
        state: Agent pipeline state
        agent: Agent name
        result_text: Raw LLM response text
        parsed_data: Parsed JSON (or None if parsing failed)
    
    Returns:
        Confidence score 0.0-1.0
    """
    # Fallback → zero confidence
    if state.get("_llm_fallback_triggered") and agent in state.get("_llm_fallback_agents", []):
        return 0.0
    
    # No response at all
    if not result_text or len(result_text.strip()) < 10:
        return 0.1
    
    # Couldn't parse JSON
    if parsed_data is None:
        return 0.3
    
    # Parsed but minimal structure
    if isinstance(parsed_data, dict) and len(parsed_data) <= 1:
        return 0.5
    
    # Rich structure with multiple keys
    if isinstance(parsed_data, dict) and len(parsed_data) >= 3:
        return 0.85
    
    # Default: structured but sparse
    return 0.7


def _compute_legal_confidence(
    state: Dict[str, Any],
    result: Any,
) -> float:
    """
    Compute legal agent confidence from analysis quality.
    
    Factors:
    - 0.0: LLM fallback triggered
    - 0.2: No findings at all (empty analysis)
    - 0.5: Few findings, mostly low/medium severity
    - 0.8: Multiple findings with high/critical severity → high confidence
    - 0.95: Definitively dangerous/rejected verdict with critical findings
    
    The logic: finding traps takes expertise, so having detailed findings
    with high severity = high confidence in the analysis.
    """
    if state.get("_llm_fallback_triggered") and "legal" in state.get("_llm_fallback_agents", []):
        return 0.0
    
    findings_count = getattr(result, "findings_count", 0) or len(getattr(result, "findings", []))
    critical_count = getattr(result, "critical_count", 0)
    high_count = getattr(result, "high_count", 0)
    verdict_value = getattr(getattr(result, "verdict", None), "value", "")
    
    if findings_count == 0:
        return 0.2
    
    # Strong signals → high confidence
    if verdict_value in ("dangerous", "rejected"):
        if critical_count >= 2:
            return 0.95
        return 0.85
    
    if critical_count >= 1 or high_count >= 3:
        return 0.80
    
    if findings_count >= 5:
        return 0.70
    
    if findings_count >= 1:
        return 0.55
    
    return 0.3


# =============================================================================
# Уровень 2: RAG-инъекция уроков (Lessons Learned) в контекст агентов
# =============================================================================

async def _get_rag_context(
    agent_name: str,
    task_description: str,
    project_id: Optional[int] = None,
    include_traps: bool = False,
) -> str:
    """
    Получить RAG-контекст для агента: документы проекта + граф знаний + БЛС.

    Использует rag_pipeline для сборки полного контекста.
    Возвращает пустую строку если RAG недоступен или нет результатов.
    """
    try:
        context = await rag_pipeline.get_agent_context(
            query=task_description,
            agent=agent_name,
            project_id=project_id,
            top_k=5,
            include_graph=True,
            include_traps=include_traps,
        )
        if context:
            logger.info("RAG context injected for %s: %d chars", agent_name, len(context))
        return context
    except (LLMUnavailableError, NetworkError) as e:
        logger.warning("RAG context unavailable for %s: %s", agent_name, e)
        return ""
    except Exception as e:
        logger.warning("RAG context failed for %s: %s", agent_name, e)
        return ""


async def _get_lessons_context(
    agent_name: str, 
    task_description: str, 
    work_type: str = "*"
) -> str:
    """
    Получить контекст из Lessons Learned для инъекции в промпт агента.
    
    Автоматически ищет релевантные уроки через RAG и формирует
    текстовый блок для добавления в промпт.
    
    Args:
        agent_name: Имя агента (ПТО, Юрист, Сметчик...)
        task_description: Описание задачи (для семантического поиска)
        work_type: Вид работ из WorkTypeRegistry
        
    Returns:
        Строка с уроками или пустая строка
    """
    try:
        context = await lessons_service.get_lessons_context(
            work_type=work_type,
            agent_name=agent_name,
            task_description=task_description,
            top_k=5,
        )
        if context:
            logger.info(f"Lessons injected for {agent_name}: {len(context)} chars")
        return context
    except (LLMUnavailableError, NetworkError) as e:
        logger.warning(f"Lessons unavailable for {agent_name}: {e}")
        return ""
    except Exception as e:
        logger.warning(f"Failed to inject lessons context for {agent_name}: {e}")
        return ""


def _extract_work_type(state: Dict[str, Any]) -> str:
    """Извлечь вид работ из стейта для RAG-поиска уроков."""
    intermediate = state.get("intermediate_data", {})
    # Извлекаем из intermediate_data или из описания задачи
    return intermediate.get("work_type", "*")


# =============================================================================
# PM Orchestrator Node — routing + verdict (v12.0: merged from HermesRouter)
# =============================================================================

async def hermes_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    [DEPRECATED] Legacy static router — superseded by PM-driven nodes_v2.

    This node uses a static routing map (archive→procurement→pto→smeta→legal→logistics)
    instead of the dynamic WorkPlan/TaskNode dispatch in pm_dispatch_router.

    Use pm_planning_node + pm_fan_out_router from nodes_v2.py instead.

    v12.0: Uses PM Agent logic (merged from HermesRouter):
      1. Weighted scoring (agent weights: Legal 0.35, Smeta 0.25, PTO 0.20, Procurement 0.12, Logistics 0.08)
      2. LLM reasoning for grey zone (0.3–0.7)
      3. Veto rules (legal DANGEROUS, margin<10%, critical_traps>=3, НМЦК<70%)

    Pipeline flow:
      start -> archive -> procurement -> pto -> smeta -> legal -> logistics -> verdict
    """
    import warnings
    warnings.warn(
        "hermes_node is deprecated. Use pm_planning_node + pm_fan_out_router from nodes_v2.py",
        DeprecationWarning, stacklevel=2,
    )
    logger.info("PM Orchestrator Router (v12.0) — DEPRECATED")
    next_step = state.get("next_step", "start")

    # Simple routing map for pipeline flow
    routing_map = {
        "start": "archive",
        "archive": "procurement",
        "procurement": "pto",
        "pto": "smeta",
        "smeta": "legal",
        "legal": "logistics",
    }

    if next_step in routing_map:
        return {
            "next_step": routing_map[next_step],
            "current_step": next_step,
        }
    elif next_step == "logistics":
        # All agents done -> compute verdict via weighted scoring + veto rules

        if state.get("_llm_fallback_triggered"):
            fallback_agents = state.get("_llm_fallback_agents", [])
            logger.warning(
                "LLM fallback detected for agents: %s. Verdict may be unreliable.",
                ", ".join(fallback_agents),
            )

        try:
            # Collect signals from all agents
            signals = [
                extract_legal_signal(state),
                extract_smeta_signal(state),
                extract_pto_signal(state),
                extract_procurement_signal(state),
                extract_logistics_signal(state),
            ]
            scoring = compute_weighted_score(signals)
            veto_id, veto_rules = check_veto_rules(state, DEFAULT_VETO_RULES)
            risk_level = calculate_risk_level(state, scoring)

            verdict = {
                "weighted_score": scoring.normalized_score,
                "zone": scoring.zone,
                "risk_level": risk_level.value if risk_level else "unknown",
                "veto_triggered": veto_id,
                "agent_signals": {s.agent_name: s.signal for s in signals},
                "agent_contributions": scoring.agent_contributions,
            }

            if veto_id:
                logger.warning("VETO triggered: %s — forcing no_go", veto_id)
                verdict["verdict"] = "no_go"
            elif scoring.zone == "go_zone":
                verdict["verdict"] = "go"
            elif scoring.zone == "no_go_zone":
                verdict["verdict"] = "no_go"
            else:
                verdict["verdict"] = "grey_zone"

            return {
                "next_step": "complete",
                "current_step": "verdict",
                "hermes_decision": verdict,
            }
        except (LLMUnavailableError, ValueError, KeyError) as e:
            logger.error(f"PM verdict computation failed: {e}")
            return {
                "next_step": "complete",
                "current_step": "verdict_fallback",
            }
    else:
        return {"next_step": "complete", "is_complete": True}


# =============================================================================
# Worker Nodes — специализированные агенты
# =============================================================================

async def archive_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """Узел Делопроизводителя: Регистрация и структурирование. (Gemma 4 E4B)"""
    logger.info("Archive Processing Starting")
    rules = load_wiki_page("Archive_Rules")
    prompt = (
        f"Инструкции:\n{rules}\n\n"
        f"Задача:\n{state['task_description']}\n\n"
        f"Зарегистрируйте входящий пакет документов."
    )

    messages = [{"role": "user", "content": prompt}]
    result_text = await _safe_agent_chat(
        "archive",
        messages,
        state,
        fallback_response='{"status": "registered", "doc_id": "REG-001"}',
    )

    try:
        archive_data = json.loads(result_text)
        confidence = _compute_agent_confidence(state, "archive", result_text, archive_data)
    except (json.JSONDecodeError, TypeError):
        archive_data = {"status": "unknown", "doc_id": "PARSE_ERROR"}
        confidence = _compute_agent_confidence(state, "archive", result_text, None)

    return {
        "archive_result": {
            "doc_id": archive_data.get("doc_id", "UNKNOWN"),
            "status": archive_data.get("status", "unknown"),
            "pages_registered": archive_data.get("pages", 0),
            "confidence_score": confidence,
        },
        "confidence_scores": {**state.get("confidence_scores", {}), "archive": confidence},
        "next_step": "archive",
    }


async def archive_ingest_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Узел инвентаризации: парсинг + индексация документа.
    Вызывается когда archive находит новый документ для обработки.

    Использует RAG Pipeline для полного цикла: file → parse → index.
    """
    logger.info("Archive Ingest: parsing and indexing documents")
    intermediate = state.get("intermediate_data", {})
    file_path = intermediate.get("file_path") or intermediate.get("contract_path")

    if not file_path or not os.path.exists(file_path):
        logger.warning("Archive Ingest: no file_path found in state")
        return {"next_step": "archive"}

    project_id = state["project_id"]

    # Полный цикл: парсинг + индексация через RAG Pipeline
    try:
        doc = await rag_pipeline.ingest(
            file_path=file_path,
            project_id=project_id,
            doc_type="unknown",
            embed=True,
            auto_classify=True,
        )
        if doc:
            logger.info(
                "Archive Ingest: doc #%d created (%s), %d chunks indexed",
                doc.id, doc.doc_type, len(doc.chunks),
            )
            return {
                "archive_result": {
                    "doc_id": str(doc.id),
                    "status": "indexed",
                    "pages_registered": len(doc.chunks),
                    "confidence_score": 0.9,
                },
                "intermediate_data": {
                    **intermediate,
                    "ingested_doc_id": doc.id,
                    "ingested_doc_type": doc.doc_type,
                },
                "next_step": "archive",
            }
    except (DocumentParseError, OSError, ValueError) as e:
        logger.error("Archive Ingest failed: %s", e)

    return {"next_step": "archive"}


async def procurement_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """Узел Закупщика: Анализ тендера и НМЦК. (Gemma 4 31B)"""
    logger.info("Procurement Analysis Starting")
    rules = load_wiki_page("Procurement_Rules")
    work_type = _extract_work_type(state)
    lessons = await _get_lessons_context("Закупщик", state['task_description'], work_type)
    
    prompt = (
        f"Инструкции:\n{rules}\n\n"
        f"{lessons}\n\n"
        f"Задача:\n{state['task_description']}\n\n"
        f"Оцените лот и его рентабельность."
    )

    messages = [{"role": "user", "content": prompt}]
    result_text = await _safe_agent_chat(
        "procurement",
        messages,
        state,
        fallback_response='{"lot_id": "T-100", "nmck": 50000000, "decision": "bid"}',
    )

    try:
        proc_data = json.loads(result_text)
        confidence = _compute_agent_confidence(state, "procurement", result_text, proc_data)
    except (json.JSONDecodeError, TypeError):
        proc_data = {}
        confidence = _compute_agent_confidence(state, "procurement", result_text, None)

    return {
        "procurement_result": {
            "lot_id": proc_data.get("lot_id", ""),
            "nmck": proc_data.get("nmck", 0),
            "nmck_vs_market": proc_data.get("nmck_vs_market", 0),
            "competitor_count": proc_data.get("competitor_count", 0),
            "decision": proc_data.get("decision", "bid"),
            "confidence_score": confidence,
        },
        "confidence_scores": {**state.get("confidence_scores", {}), "procurement": confidence},
        "next_step": "procurement",
    }


async def pto_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """Узел ПТО: Анализ объемов работ через LLM. (Gemma 4 31B VLM)"""
    logger.info("PTO Analysis Starting")
    rules = load_wiki_page("PTO_Rules")
    work_type = _extract_work_type(state)

    # v12.0: Извлечение текста из файла, если он передан в стейте
    intermediate = state.get("intermediate_data", {})
    file_path = intermediate.get("file_path")
    document_text = intermediate.get("document_text", "")

    if file_path and os.path.exists(file_path):
        logger.info("PTO reading file: %s", os.path.basename(file_path))
        if file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            document_text = "\n".join([para.text for para in doc.paragraphs])
        elif file_path.endswith(".pdf"):
            with fitz.open(file_path) as doc:
                document_text = "\n".join([page.get_text() for page in doc])
    
    # v12.0: RAG-инъекция уроков (Lessons Learned)
    lessons = await _get_lessons_context("ПТО", state['task_description'], work_type)
    # v12.0: RAG-контекст из документов проекта
    project_id = state.get("project_id")
    rag_ctx = await _get_rag_context("pto", state['task_description'], project_id, include_traps=True)
    
    prompt = (
        f"Инструкции:\n{rules}\n\n"
        f"{lessons}\n\n"
        f"{rag_ctx}\n\n"
        f"Задача:\n{state['task_description']}\n\n"
        f"Контекст документа:\n{document_text[:10000]}\n\n"
        f"Извлеките ВОР в формате JSON (укажите только JSON)."
    )

    messages = [{"role": "user", "content": prompt}]
    start_time = datetime.now()

    result_text = await _safe_agent_chat(
        "pto",
        messages,
        state,
        fallback_response='{"earthworks": 1500, "concrete": 200}',
    )

    duration = int((datetime.now() - start_time).total_seconds() * 1000)

    try:
        vor_data = json.loads(result_text)
        confidence = _compute_agent_confidence(state, "pto", result_text, vor_data)
    except (json.JSONDecodeError, TypeError):
        vor_data = {"raw_text": result_text}
        confidence = _compute_agent_confidence(state, "pto", result_text, None)

    with Session() as session:
        log = AuditLog(
            agent_name="PTO",
            action="extract_vor",
            input_data={"task": state["task_description"]},
            output_data=vor_data,
            duration_ms=duration,
        )
        session.add(log)
        session.commit()

    return {
        "vor_result": {
            "positions": vor_data.get("volumes", []),
            "confidence_score": confidence,
            "total_positions": len(vor_data.get("volumes", [])),
            "drawing_refs": [vor_data.get("source_drawing", "")],
            "unit_mismatches": 0,
        },
        "confidence_scores": {**state.get("confidence_scores", {}), "pto": confidence},
        "next_step": "pto",
    }


async def logistics_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """Узел Логиста: Поиск поставщиков и цен. (Gemma 4 31B)"""
    logger.info("Logistics Sourcing Starting")
    rules = load_wiki_page("Logistics_Rules")
    # Read VOR from typed state field (v2.0) or fall back to intermediate_data
    vor = state.get("vor_result") or state.get("intermediate_data", {}).get("vor", {})
    work_type = _extract_work_type(state)
    # v12.0: RAG-инъекция уроков (Lessons Learned)
    lessons = await _get_lessons_context("Логист", state['task_description'], work_type)

    prompt = (
        f"Инструкции:\n{rules}\n\n"
        f"{lessons}\n\n"
        f"ВОР:\n{json.dumps(vor, ensure_ascii=False)}\n\n"
        f"Найдите поставщиков для этих работ."
    )

    messages = [{"role": "user", "content": prompt}]
    result_text = await _safe_agent_chat(
        "logistics",
        messages,
        state,
        fallback_response='{"vendors": ["MetalInvest", "Severstal"], "best_price": 65000}',
    )

    try:
        log_data = json.loads(result_text)
        confidence = _compute_agent_confidence(state, "logistics", result_text, log_data)
    except (json.JSONDecodeError, TypeError):
        log_data = {}
        confidence = _compute_agent_confidence(state, "logistics", result_text, None)

    return {
        "logistics_result": {
            "vendors_found": len(log_data.get("vendors", [])),
            "best_price": log_data.get("best_price", 0),
            "delivery_available": log_data.get("delivery_available", True),
            "lead_time_days": log_data.get("lead_time_days", 14),
            "confidence_score": confidence,
        },
        "confidence_scores": {**state.get("confidence_scores", {}), "logistics": confidence},
        "next_step": "logistics",
    }


async def smeta_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """Узел Сметчика: Расчет стоимостей через LLM. (Gemma 4 31B)"""
    logger.info("Smeta Calculation Starting")
    rules = load_wiki_page("Smeta_Rules")
    # Read VOR from typed state field (v2.0) or fall back to intermediate_data
    vor = state.get("vor_result") or state.get("intermediate_data", {}).get("vor", {})
    work_type = _extract_work_type(state)
    # v12.0: RAG-инъекция уроков (Lessons Learned)
    lessons = await _get_lessons_context("Сметчик", state['task_description'], work_type)
    # v12.0: RAG-контекст
    project_id = state.get("project_id")
    rag_ctx = await _get_rag_context("smeta", state['task_description'], project_id, include_traps=True)

    prompt = (
        f"Инструкции:\n{rules}\n\n"
        f"{lessons}\n\n"
        f"{rag_ctx}\n\n"
        f"Доступные объемы:\n{json.dumps(vor, ensure_ascii=False)}\n\n"
        f"Оцените стоимость, верните валидный JSON."
    )
    messages = [{"role": "user", "content": prompt}]

    result_text = await _safe_agent_chat(
        "smeta",
        messages,
        state,
        fallback_response='{"earthworks_cost": 150000, "concrete_cost": 20000}',
    )

    try:
        smeta_data = json.loads(result_text)
        confidence = _compute_agent_confidence(state, "smeta", result_text, smeta_data)
    except (json.JSONDecodeError, TypeError):
        smeta_data = {"raw_costs": result_text}
        confidence = _compute_agent_confidence(state, "smeta", result_text, None)

    return {
        "smeta_result": {
            "total_cost": smeta_data.get("grand_totals", {}).get("total_with_vat", 0),
            "nmck": 0,
            "profit_margin_pct": 0,
            "fer_positions_used": len(smeta_data.get("grand_totals", {})),
            "confidence_score": confidence,
        },
        "confidence_scores": {**state.get("confidence_scores", {}), "smeta": confidence},
        "next_step": "smeta",
    }


async def legal_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Узел Юриста: Полный юридический анализ через LegalService. (Gemma 4 31B, 128K контекст)

    Поддерживает два режима:
    - Если в intermediate_data есть document_text/file_path — полный анализ
    - Иначе — быстрый обзор по task_description (Quick Review)

    v12.0: RAG-инъекция уроков + БЛС.
    v12.0: Сохраняет результат для последующей генерации протокола/претензии/иска.
    """
    logger.info("Legal Review Starting")

    from src.core.services.legal_service import legal_service
    from src.schemas.legal import LegalAnalysisRequest, ReviewType

    # Determine document source
    intermediate = state.get("intermediate_data", {})
    document_text = intermediate.get("document_text")
    file_path = intermediate.get("contract_path") or intermediate.get("file_path")
    document_id = intermediate.get("document_id")

    # v12.0: Извлечение текста из файла
    if file_path and os.path.exists(file_path):
        logger.info("Legal reading file: %s", os.path.basename(file_path))
        if file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            document_text = "\n".join([para.text for para in doc.paragraphs])
        elif file_path.endswith(".pdf"):
            with fitz.open(file_path) as doc:
                document_text = "\n".join([page.get_text() for page in doc])

    # v12.0: RAG + Lessons инъекция
    work_type = _extract_work_type(state)
    lessons = await _get_lessons_context("Юрист", state['task_description'], work_type)
    project_id = state.get("project_id")
    rag_ctx = await _get_rag_context("legal", state['task_description'], project_id, include_traps=True)
    if lessons and document_text:
        document_text = f"{lessons}\n\n{document_text}"
    if rag_ctx and document_text:
        document_text = f"{rag_ctx}\n\n{document_text}"

    try:
        request = LegalAnalysisRequest(
            document_id=document_id,
            document_text=document_text,
            file_path=None,
            review_type=ReviewType.CONTRACT,
        )

        result = await legal_service.analyze(request)
        findings_dicts = [f.model_dump() for f in result.findings]
        legal_confidence = _compute_legal_confidence(state, result)

        # Сохраняем полный результат в intermediate_data для генераторов документов
        return {
            "findings": findings_dicts,
            "legal_result": {
                "verdict": result.verdict.value,
                "findings_count": len(findings_dicts),
                "critical_count": result.critical_count,
                "high_count": result.high_count,
                "summary": result.summary,
                "protocol_items_count": len(result.protocol_items),
                "confidence_score": legal_confidence,
                "blc_matches": [],
            },
            "confidence_scores": {**state.get("confidence_scores", {}), "legal": legal_confidence},
            "intermediate_data": {
                **intermediate,
                "legal_verdict": result.verdict.value,
                "legal_summary": result.summary,
                "legal_critical_count": result.critical_count,
                "legal_high_count": result.high_count,
                "_legal_analysis_result": result.model_dump(),  # Полный результат для генераторов
            },
            "is_complete": True,
            "next_step": "legal",
        }

    except Exception as e:
        logger.error(f"Legal analysis failed: {e}")
        return {
            "findings": [{"trap": "Analysis Error", "risk": "High", "details": str(e)}],
            "is_complete": True,
            "next_step": "legal",
        }


# =============================================================================
# Legal Document Generation Nodes (v12.0)
# =============================================================================

async def legal_protocol_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Узел генерации протокола разногласий.

    Берёт результат legal_node из _legal_analysis_result,
    генерирует DOCX протокола разногласий (3 колонки).
    """
    logger.info("Legal Protocol Generation Starting")
    intermediate = state.get("intermediate_data", {})

    # Достаём сохранённый результат анализа
    analysis_data = intermediate.get("_legal_analysis_result")
    if not analysis_data:
        logger.warning("No _legal_analysis_result in state — running legal_node first")
        legal_state = await legal_node(state)
        analysis_data = legal_state.get("intermediate_data", {}).get("_legal_analysis_result")

    if not analysis_data:
        return {
            "intermediate_data": {**intermediate, "protocol_error": "No analysis result available"},
            "next_step": "legal",
        }

    try:
        from src.schemas.legal import LegalAnalysisResult
        from src.core.services.legal_documents import legal_doc_gen

        result = LegalAnalysisResult(**analysis_data)

        docx_path = await legal_doc_gen.generate_protocol(
            analysis_result=result,
            contract_number=intermediate.get("contract_number", ""),
            contract_date=intermediate.get("contract_date", ""),
            output_dir="/tmp/asd_docs",
        )

        logger.info("Protocol generated: %s", docx_path)
        return {
            "intermediate_data": {
                **intermediate,
                "protocol_docx_path": docx_path,
                "protocol_items_count": len(result.protocol_items),
            },
            "next_step": "legal",
        }
    except Exception as e:
        logger.error("Protocol generation failed: %s", e)
        return {
            "intermediate_data": {**intermediate, "protocol_error": str(e)},
            "next_step": "legal",
        }


async def legal_claim_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Узел генерации досудебной претензии.
    """
    logger.info("Legal Claim Generation Starting")
    intermediate = state.get("intermediate_data", {})

    analysis_data = intermediate.get("_legal_analysis_result")
    if not analysis_data:
        logger.warning("No analysis result for claim generation")
        return {"intermediate_data": {**intermediate, "claim_error": "No analysis"}, "next_step": "legal"}

    try:
        from src.schemas.legal import LegalAnalysisResult
        from src.core.services.legal_documents import legal_doc_gen

        result = LegalAnalysisResult(**analysis_data)

        claim_path = await legal_doc_gen.generate_claim(
            analysis_result=result,
            contract_number=intermediate.get("contract_number", ""),
            contract_date=intermediate.get("contract_date", ""),
            contract_subject=intermediate.get("contract_subject", ""),
            customer_name=intermediate.get("customer_name", ""),
            customer_inn=intermediate.get("customer_inn", ""),
            contractor_name=intermediate.get("contractor_name", ""),
            contractor_inn=intermediate.get("contractor_inn", ""),
            claim_amount=intermediate.get("claim_amount", 0.0),
            output_dir="/tmp/asd_docs",
        )

        logger.info("Claim generated: %s", claim_path)
        return {
            "intermediate_data": {**intermediate, "claim_path": claim_path},
            "next_step": "legal",
        }
    except Exception as e:
        logger.error("Claim generation failed: %s", e)
        return {"intermediate_data": {**intermediate, "claim_error": str(e)}, "next_step": "legal"}


async def legal_lawsuit_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Узел генерации искового заявления в арбитражный суд.
    """
    logger.info("Legal Lawsuit Generation Starting")
    intermediate = state.get("intermediate_data", {})

    analysis_data = intermediate.get("_legal_analysis_result")
    if not analysis_data:
        return {"intermediate_data": {**intermediate, "lawsuit_error": "No analysis"}, "next_step": "legal"}

    try:
        from src.schemas.legal import LegalAnalysisResult
        from src.core.services.legal_documents import legal_doc_gen

        result = LegalAnalysisResult(**analysis_data)

        lawsuit_path = await legal_doc_gen.generate_lawsuit(
            analysis_result=result,
            case_facts=intermediate.get("case_facts", ""),
            contract_number=intermediate.get("contract_number", ""),
            contract_date=intermediate.get("contract_date", ""),
            contractor_name=intermediate.get("contractor_name", ""),
            contractor_inn=intermediate.get("contractor_inn", ""),
            contractor_ogrn=intermediate.get("contractor_ogrn", ""),
            customer_name=intermediate.get("customer_name", ""),
            customer_inn=intermediate.get("customer_inn", ""),
            claim_amount=intermediate.get("claim_amount", 0.0),
            output_dir="/tmp/asd_docs",
        )

        logger.info("Lawsuit generated: %s", lawsuit_path)
        return {
            "intermediate_data": {**intermediate, "lawsuit_path": lawsuit_path},
            "next_step": "legal",
        }
    except Exception as e:
        logger.error("Lawsuit generation failed: %s", e)
        return {"intermediate_data": {**intermediate, "lawsuit_error": str(e)}, "next_step": "legal"}


# =============================================================================
# PTO Agent Nodes (v12.0) — Инвентаризация и верификация ИД
# =============================================================================

async def pto_inventory_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Узел инвентаризации ПТО: классификация документов по матрице 344/пр.

    Принимает список документов проекта, классифицирует каждый по 13 позициям,
    строит матрицу полноты и возвращает дельту.
    """
    logger.info("PTO Inventory Starting")
    from src.core.services.pto_agent import pto_agent, CompReport, CompletenessGap

    intermediate = state.get("intermediate_data", {})
    project_id = state["project_id"]

    # Получаем список видов работ из контекста
    work_types = intermediate.get("work_types", []) or _extract_work_types(state)

    # Получаем список документов проекта (из БД или из state)
    available_docs = intermediate.get("available_docs", [])
    if not available_docs:
        try:
            from src.core.document_repository import document_repo
            db_docs = document_repo.list_by_project(project_id)
            available_docs = [
                {"doc_type": d.doc_type, "doc_id": str(d.id), "name": d.filename}
                for d in db_docs
            ]
        except Exception as e:
            logger.warning("PTO: cannot load docs from DB: %s", e)

    # Генерируем отчёт
    try:
        report = await pto_agent.generate_completeness_report(
            project_id=project_id,
            work_types=work_types,
            available_docs=available_docs,
        )

        report_text = pto_agent.format_report(report)

        logger.info(
            "PTO inventory: %.1f%% complete, %d critical gaps, %d AOSR trails",
            report.completeness_pct, len(report.critical_gaps), len(report.aosr_trails),
        )

        return {
            "intermediate_data": {
                **intermediate,
                "pto_report": report_text,
                "pto_completeness_pct": report.completeness_pct,
                "pto_critical_gaps": len(report.critical_gaps),
                "pto_total_gaps": len(report.gaps),
                "pto_aosr_count": len(report.aosr_trails),
                "work_types": work_types,
            },
            "compliance_delta": {
                "completeness_pct": str(report.completeness_pct),
                "critical_gaps": str(len(report.critical_gaps)),
                "total_gaps": str(len(report.gaps)),
            },
            "next_step": "pto",
        }
    except Exception as e:
        logger.error("PTO inventory failed: %s", e)
        return {
            "intermediate_data": {**intermediate, "pto_error": str(e)},
            "next_step": "pto",
        }


async def pto_verify_trail_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Узел верификации шлейфа АОСР: проверка полного комплекта документов к акту.

    Для каждого АОСР проверяет: ВОР, ИГС, ИС, документы качества, протоколы, уведомление.
    """
    logger.info("PTO Trail Verification Starting")
    from src.core.services.pto_agent import pto_agent

    intermediate = state.get("intermediate_data", {})
    work_types = intermediate.get("work_types", []) or _extract_work_types(state)

    # Строим все шлейфы
    all_trails = []
    for wt in work_types:
        aosr_list = pto_agent.get_required_aosr_list(wt)
        for aosr in aosr_list:
            if aosr.get("mandatory", True):
                trail = pto_agent.build_aosr_trail(aosr["name"], wt)
                all_trails.append(trail)

    # Проверяем каждый шлейф (без БД — эвристически)
    incomplete_trails = []
    for trail in all_trails:
        if not trail.is_complete:
            incomplete_trails.append({
                "aosr_name": trail.aosr_name,
                "work_type": trail.work_type,
                "missing": [
                    {"type": i.item_type, "name": i.name}
                    for i in trail.missing_mandatory
                ],
            })

    logger.info(
        "PTO trail verify: %d trails, %d incomplete",
        len(all_trails), len(incomplete_trails),
    )

    return {
        "intermediate_data": {
            **intermediate,
            "pto_trails_total": len(all_trails),
            "pto_trails_incomplete": len(incomplete_trails),
            "pto_incomplete_trails": incomplete_trails,
        },
        "next_step": "pto",
    }


def _extract_work_types(state: Dict[str, Any]) -> List[str]:
    """Извлечь список видов работ из описания задачи или контекста."""
    task = state.get("task_description", "").lower()
    intermediate = state.get("intermediate_data", {})

    work_types = []
    # Простые эвристики
    kw_to_wt = {
        "фундамент": "фундаменты_монолитные",
        "сваи": "фундаменты_свайные",
        "котлован": "земляные_выемки",
        "бетон": "бетонные",
        "металлоконструкц": "металлоконструкции",
        "кладка": "каменная_кладка",
        "отделк": "отделка_стены_потолки",
        "пол": "отделка_полы",
        "электромонтаж": "электромонтаж_внутренние",
        "отоплен": "отопление",
        "вентиляц": "вентиляция",
        "водоснабжен": "водоснабжение",
        "канализац": "канализация",
    }
    for kw, wt in kw_to_wt.items():
        if kw in task:
            work_types.append(wt)

    return work_types if work_types else ["бетонные", "фундаменты_монолитные"]


# =============================================================================
# Smeta & Delo Agent Nodes (v12.0)
# =============================================================================

async def smeta_estimate_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """Узел Сметчика: построение сметы и анализ рентабельности."""
    logger.info("Smeta Estimate Starting")
    from src.core.services.smeta_agent import smeta_agent, SmetaAgent

    intermediate = state.get("intermediate_data", {})
    project_id = state["project_id"]
    task = state["task_description"]

    # Получаем ВОР (из ПТО или из state)
    vor = state.get("vor_result") or intermediate.get("vor", {})
    vor_positions = vor.get("positions", []) if isinstance(vor, dict) else []

    if not vor_positions:
        # Создаём демо-позиции для теста
        vor_positions = [
            {"code": "ФЕР06-01-001", "name": "Бетонная подготовка", "unit": "м³", "quantity": 50},
            {"code": "ФЕР06-01-015", "name": "Армирование", "unit": "т", "quantity": 12},
            {"code": "ФЕР06-01-020", "name": "Бетонирование", "unit": "м³", "quantity": 200},
        ]

    nmck = intermediate.get("nmck", 0.0) or float(intermediate.get("nmck_value", 15_000_000))

    try:
        estimate = smeta_agent.build_estimate(
            project_id=project_id,
            title=f"Смета — {task[:60]}",
            vor_positions=vor_positions,
            region_coeff=intermediate.get("region_coeff", 1.0),
            index_coeff=intermediate.get("index_coeff", 1.0),
        )

        margin = smeta_agent.analyze_margin(estimate, nmck)

        logger.info(
            "Smeta: grand_total=%.0f, margin=%.1f%%, zone=%s",
            estimate.grand_total, margin.margin_pct, margin.margin_zone.value,
        )

        return {
            "smeta_result": {
                "total_cost": estimate.grand_total,
                "nmck": nmck,
                "profit_margin_pct": margin.margin_pct,
                "fer_positions_used": len(estimate.lines),
                "confidence_score": 0.85,
                "low_margin_positions": margin.low_margin_lines[:5],
            },
            "intermediate_data": {
                **intermediate,
                "smeta_estimate": estimate.to_dict(),
                "smeta_margin": margin.to_dict(),
                "smeta_grand_total": estimate.grand_total,
                "smeta_margin_pct": margin.margin_pct,
            },
            "next_step": "smeta",
        }
    except Exception as e:
        logger.error("Smeta estimate failed: %s", e)
        return {"intermediate_data": {**intermediate, "smeta_error": str(e)}, "next_step": "smeta"}


async def smeta_compare_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """Узел Сметчика: сравнение ВОР ↔ КС-2."""
    logger.info("Smeta VOR/KS-2 Comparison Starting")
    from src.core.services.smeta_agent import smeta_agent

    intermediate = state.get("intermediate_data", {})

    vor_positions = intermediate.get("vor_positions", [])
    ks2_positions = intermediate.get("ks2_positions", [])

    try:
        comparison = smeta_agent.compare_vor_ks2(vor_positions, ks2_positions)

        return {
            "intermediate_data": {
                **intermediate,
                "smeta_comparison": comparison,
                "smeta_has_discrepancies": comparison["has_discrepancies"],
            },
            "next_step": "smeta",
        }
    except Exception as e:
        logger.error("Smeta compare failed: %s", e)
        return {"intermediate_data": {**intermediate, "smeta_error": str(e)}, "next_step": "smeta"}


async def delo_registry_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """Узел Делопроизводителя: создание и обновление реестра ИД."""
    logger.info("Delo Registry Starting")
    from src.core.services.delo_agent import delo_agent, DocStatus

    intermediate = state.get("intermediate_data", {})
    project_id = state["project_id"]

    try:
        # Создаём/получаем реестр
        registry = delo_agent.get_registry(project_id)
        if not registry:
            registry = delo_agent.create_registry(
                project_id, f"Проект #{project_id}"
            )

        # Авто-регистрация документов из доступных
        available_docs = intermediate.get("available_docs", [])
        for doc in available_docs[:20]:  # Не больше 20 за раз
            if not any(e.doc_name == doc.get("name", "") for e in registry.entries):
                delo_agent.register_document(
                    project_id=project_id,
                    doc_type=doc.get("doc_type", "unknown"),
                    doc_name=doc.get("name", doc.get("filename", "")),
                    category_344=doc.get("category_344", "act_aosr"),
                )

        report = delo_agent.generate_registry_report(project_id)
        stats = delo_agent.get_completion_stats(project_id)

        logger.info(
            "Delo: %d docs, %.1f%% complete, %d overdue",
            stats.get("total", 0), stats.get("completion_pct", 0), stats.get("overdue", 0),
        )

        return {
            "intermediate_data": {
                **intermediate,
                "delo_report": report,
                "delo_stats": stats,
                "delo_total_docs": stats.get("total", 0),
            },
            "next_step": "archive",
        }
    except Exception as e:
        logger.error("Delo registry failed: %s", e)
        return {"intermediate_data": {**intermediate, "delo_error": str(e)}, "next_step": "archive"}


# =============================================================================
# Procurement & Logistics Nodes (v12.0)
# =============================================================================

async def procurement_analyze_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """Узел Закупщика: анализ тендера и поиск поставщиков."""
    logger.info("Procurement Analysis Starting")
    from src.core.services.procurement_logistics import procurement_agent, TenderInfo, TenderDecision

    intermediate = state.get("intermediate_data", {})
    project_id = state["project_id"]

    # Данные тендера
    tender = TenderInfo(
        lot_id=intermediate.get("lot_id", f"LOT-{project_id:04d}"),
        title=state["task_description"][:100],
        nmck=float(intermediate.get("nmck", 15_000_000)),
        customer=intermediate.get("customer_name", ""),
        region=intermediate.get("region", "Москва"),
    )

    # Себестоимость из Сметчика
    estimated_cost = intermediate.get("smeta_grand_total", tender.nmck * 0.75)

    try:
        analysis = procurement_agent.analyze_tender(
            tender, estimated_cost, market_work_type=intermediate.get("work_type", "")
        )

        # Поиск поставщиков
        materials = intermediate.get("materials", ["Бетон", "Арматура", "Металлоконструкции"])
        quotes = procurement_agent.search_suppliers(materials, tender.region)
        comparison = procurement_agent.compare_quotes(quotes)

        logger.info(
            "Procurement: decision=%s, margin=%.1f%%, suppliers=%d",
            analysis.decision.value, analysis.margin_pct, len(quotes),
        )

        return {
            "procurement_result": {
                "lot_id": tender.lot_id,
                "nmck": tender.nmck,
                "nmck_vs_market": analysis.nmck_vs_market_pct,
                "competitor_count": analysis.competitor_count,
                "decision": analysis.decision.value,
                "confidence_score": 0.8,
            },
            "intermediate_data": {
                **intermediate,
                "procurement_analysis": analysis.to_dict(),
                "procurement_quotes": comparison,
                "procurement_decision": analysis.decision.value,
            },
            "next_step": "procurement",
        }
    except Exception as e:
        logger.error("Procurement analysis failed: %s", e)
        return {"intermediate_data": {**intermediate, "procurement_error": str(e)}, "next_step": "procurement"}


async def logistics_plan_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """Узел Логиста: планирование доставки."""
    logger.info("Logistics Planning Starting")
    from src.core.services.procurement_logistics import logistics_agent, TransportType, SupplierQuote

    intermediate = state.get("intermediate_data", {})
    project_id = state["project_id"]

    # Поставщики из Закупщика
    quotes_data = intermediate.get("procurement_quotes", {}).get("suppliers", [])
    quotes = [
        SupplierQuote(
            supplier_name=q.get("name", ""),
            material_name="",
            unit="",
            quantity=1,
            unit_price=q.get("price", 0),
            total=q.get("price", 0),
            delivery_days=q.get("delivery_days", 14),
            rating=q.get("rating", 5),
        )
        for q in quotes_data
    ]

    destination = intermediate.get("region", "Москва")

    try:
        plan = logistics_agent.build_logistics_plan(project_id, quotes, destination)

        logger.info(
            "Logistics: %d routes, %.0f ₽ transport cost, %d days max lead",
            len(plan.routes), plan.total_transport_cost, plan.max_lead_time_days,
        )

        return {
            "logistics_result": {
                "vendors_found": plan.vendor_count,
                "best_price": min((r.total_cost for r in plan.routes), default=0),
                "delivery_available": True,
                "lead_time_days": plan.max_lead_time_days,
                "confidence_score": 0.75,
            },
            "intermediate_data": {
                **intermediate,
                "logistics_plan": plan.to_dict(),
                "logistics_total_cost": plan.total_transport_cost,
            },
            "next_step": "logistics",
        }
    except Exception as e:
        logger.error("Logistics planning failed: %s", e)
        return {"intermediate_data": {**intermediate, "logistics_error": str(e)}, "next_step": "logistics"}
