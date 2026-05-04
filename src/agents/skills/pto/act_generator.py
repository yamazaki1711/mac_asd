"""
MAC_ASD v13.0 — PTO Act Generator Skill.

Генерация актов исполнительной документации в формате DOCX.
Поддерживает: АОСР, входной контроль, скрытые работы, освидетельствование.

Использует docxtpl (Jinja2 inside DOCX) с fallback на python-docx.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional

from src.agents.skills.common.base import SkillBase, SkillResult, SkillStatus

logger = logging.getLogger(__name__)


# ── Template paths (relative to project root) ──────────────────────────────

class ActType(str, Enum):
    AOSR = "aosr"                       # Акт освидетельствования скрытых работ
    INCOMING_CONTROL = "incoming_control"  # Акт входного контроля
    HIDDEN_WORKS = "hidden_works"       # Акт скрытых работ (синоним АОСР)
    INSPECTION = "inspection"           # Акт освидетельствования (АООК)


TEMPLATE_MAP = {
    ActType.AOSR: "library/templates/acts/aosr/3_AOSR.docx",
    ActType.INCOMING_CONTROL: "library/templates/acts/aosr/1_AOGRO.docx",
    ActType.HIDDEN_WORKS: "library/templates/acts/aosr/3_AOSR.docx",
    ActType.INSPECTION: "library/templates/acts/aosr/4_AOOK.docx",
}

DEFAULT_OUTPUT_DIR = "data/exports/acts"


class PTO_ActGenerator(SkillBase):
    """
    Навык ПТО: генерация актов ИД в формате DOCX.

    Принимает тип акта и контекстные данные, заполняет шаблон DOCX
    через docxtpl (Jinja2) или генерирует документ через python-docx
    при отсутствии шаблона.
    """

    skill_id = "PTO_ActGenerator"
    description = "Генерация актов исполнительной документации (АОСР, входной контроль, скрытые работы)"
    agent = "pto"

    def validate_input(self, params: Dict[str, Any]) -> Dict[str, Any]:
        act_type = params.get("act_type")
        if not act_type:
            return {"valid": False, "errors": ["Параметр 'act_type' обязателен"]}

        valid_types = {t.value for t in ActType}
        if act_type not in valid_types:
            return {"valid": False, "errors": [
                f"Неизвестный тип акта: {act_type}. Допустимые: {sorted(valid_types)}"
            ]}

        return {"valid": True}

    async def _execute(self, params: Dict[str, Any]) -> SkillResult:
        act_type = ActType(params["act_type"])
        context = params.get("context", {})
        output_dir = params.get("output_dir", DEFAULT_OUTPUT_DIR)
        custom_template = params.get("template_path")

        # Resolve template
        template_path = custom_template or TEMPLATE_MAP.get(act_type)
        if template_path:
            template_full = Path(template_path)
            if not template_full.is_absolute():
                from src.config import settings
                template_full = Path(settings.BASE_DIR) / template_path

        # Gather context
        context_data = self._gather_context(act_type, context)

        # Generate timestamped filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{act_type.value}_{timestamp}.docx"

        # Ensure output directory
        output_full = Path(output_dir)
        if not output_full.is_absolute():
            from src.config import settings
            output_full = Path(settings.BASE_DIR) / output_dir
        output_full.mkdir(parents=True, exist_ok=True)
        output_path = output_full / filename

        # Render
        template_used = None
        try:
            if template_path and template_full.exists():
                self._render_docxtpl(template_full, context_data, output_path)
                template_used = str(template_path)
            else:
                logger.info("Template not found — generating plain DOCX")
                self._render_plain_docx(act_type, context_data, output_path)
        except Exception as e:
            logger.warning("DOCX render failed: %s. Generating plain DOCX fallback.", e)
            try:
                self._render_plain_docx(act_type, context_data, output_path)
                template_used = None
            except Exception as fallback_e:
                return SkillResult(
                    status=SkillStatus.ERROR,
                    skill_id=self.skill_id,
                    errors=[f"DOCX generation failed: {fallback_e}"],
                )

        file_size = output_path.stat().st_size if output_path.exists() else 0

        return SkillResult(
            status=SkillStatus.SUCCESS,
            skill_id=self.skill_id,
            data={
                "act_type": act_type.value,
                "file_path": str(output_path),
                "filename": filename,
                "template_used": template_used,
                "size_bytes": file_size,
                "note": (
                    "Файл сохранён локально. "
                    "Для загрузки в Google Диск используйте asd_drive_upload."
                ),
            },
        )

    # ── Context gathering ──────────────────────────────────────────────────

    @staticmethod
    def _gather_context(act_type: ActType, provided: Dict[str, Any]) -> Dict[str, Any]:
        """Build template context from provided data with sensible defaults."""
        now = datetime.now()
        defaults = {
            "act_number": provided.get("act_number", ""),
            "act_date": provided.get("act_date", now.strftime("%d.%m.%Y")),
            "project_name": provided.get("project_name", ""),
            "object_name": provided.get("object_name", ""),
            "customer_name": provided.get("customer_name", ""),
            "contractor_name": provided.get("contractor_name", "ООО «КСК №1»"),
            "work_description": provided.get("work_description", ""),
            "work_type": provided.get("work_type", ""),
            "volume": provided.get("volume", ""),
            "unit": provided.get("unit", ""),
            "project_documentation": provided.get("project_documentation", ""),
            "materials": provided.get("materials", []),
            "commission_members": provided.get("commission_members", []),
            "normative_docs": provided.get("normative_docs", []),
            "executive_schemes": provided.get("executive_schemes", []),
            "test_reports": provided.get("test_reports", []),
            "defects": provided.get("defects", []),
            "decision": provided.get("decision", "Работы выполнены в соответствии с проектной документацией"),
            "signatures": provided.get("signatures", []),
        }

        # Add act-type-specific defaults
        if act_type == ActType.AOSR:
            defaults.setdefault("act_title", "АКТ ОСВИДЕТЕЛЬСТВОВАНИЯ СКРЫТЫХ РАБОТ")
        elif act_type == ActType.INCOMING_CONTROL:
            defaults.setdefault("act_title", "АКТ ВХОДНОГО КОНТРОЛЯ")
        elif act_type == ActType.HIDDEN_WORKS:
            defaults.setdefault("act_title", "АКТ СКРЫТЫХ РАБОТ")
        elif act_type == ActType.INSPECTION:
            defaults.setdefault("act_title", "АКТ ОСВИДЕТЕЛЬСТВОВАНИЯ ОТВЕТСТВЕННЫХ КОНСТРУКЦИЙ")

        return defaults

    # ── DOCX rendering ─────────────────────────────────────────────────────

    @staticmethod
    def _render_docxtpl(template_path: Path, context: Dict[str, Any], output_path: Path):
        """Render DOCX from Jinja2 template using docxtpl."""
        from docxtpl import DocxTemplate

        doc = DocxTemplate(str(template_path))
        doc.render(context)
        doc.save(str(output_path))
        logger.info("DOCX rendered from template: %s", output_path)

    @staticmethod
    def _render_plain_docx(act_type: ActType, context: Dict[str, Any], output_path: Path):
        """Generate a plain DOCX from scratch using python-docx (fallback)."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(context.get("act_title", f"АКТ — {act_type.value}"), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Meta
        doc.add_paragraph(f"№ {context.get('act_number', '___')}")
        doc.add_paragraph(f"Дата: {context.get('act_date', '___')}")

        # Project info
        doc.add_heading("Объект", level=2)
        doc.add_paragraph(f"Проект: {context.get('project_name', '___')}")
        doc.add_paragraph(f"Объект: {context.get('object_name', '___')}")

        # Parties
        doc.add_heading("Стороны", level=2)
        doc.add_paragraph(f"Заказчик: {context.get('customer_name', '___')}")
        doc.add_paragraph(f"Подрядчик: {context.get('contractor_name', '___')}")

        # Work description
        doc.add_heading("Выполненные работы", level=2)
        doc.add_paragraph(context.get("work_description", "Описание работ не указано"))

        vol = context.get("volume", "")
        unit = context.get("unit", "")
        if vol:
            doc.add_paragraph(f"Объём: {vol} {unit}".strip())

        # Materials
        materials = context.get("materials", [])
        if materials:
            doc.add_heading("Материалы", level=2)
            for mat in materials:
                if isinstance(mat, str):
                    doc.add_paragraph(f"• {mat}", style="List Bullet")
                elif isinstance(mat, dict):
                    doc.add_paragraph(
                        f"• {mat.get('name', '')} — {mat.get('quantity', '')} {mat.get('unit', '')}".strip(),
                        style="List Bullet",
                    )

        # Commission
        commission = context.get("commission_members", [])
        if commission:
            doc.add_heading("Комиссия", level=2)
            for member in commission:
                if isinstance(member, str):
                    doc.add_paragraph(member)
                elif isinstance(member, dict):
                    doc.add_paragraph(
                        f"{member.get('role', '')}: {member.get('name', '')}"
                    )

        # Decision
        doc.add_heading("Решение", level=2)
        doc.add_paragraph(context.get("decision", "Работы выполнены в соответствии с требованиями."))

        # Signatures
        sigs = context.get("signatures", [])
        if sigs:
            doc.add_heading("Подписи", level=2)
            for sig in sigs:
                if isinstance(sig, str):
                    doc.add_paragraph(f"__________________ {sig}")
                elif isinstance(sig, dict):
                    doc.add_paragraph(
                        f"__________________ {sig.get('name', '')} ({sig.get('role', '')})"
                    )
        else:
            # Default signature block
            doc.add_paragraph("")
            doc.add_paragraph("__________________ Подрядчик")
            doc.add_paragraph("__________________ Заказчик")
            doc.add_paragraph("__________________ Стройконтроль")

        doc.save(str(output_path))
        logger.info("Plain DOCX generated: %s", output_path)
