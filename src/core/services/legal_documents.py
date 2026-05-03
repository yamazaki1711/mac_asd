"""
ASD v12.0 — Legal Document Generators.

Генерация юридических документов на основе результатов LegalService:
  1. Протокол разногласий (DOCX) — 3-колоночный формат
  2. Досудебная претензия — требование об устранении нарушений
  3. Исковое заявление — в арбитражный суд

Каждый документ опирается на конкретные статьи закона (ГК РФ, ГрК РФ, ФЗ-44/223),
анализ БЛС и судебную практику.

Usage:
    from src.core.services.legal_documents import LegalDocumentGenerator

    gen = LegalDocumentGenerator(llm_engine=llm_engine)
    
    # Протокол разногласий
    docx_path = await gen.generate_protocol(result, output_dir="/tmp")
    
    # Претензия
    claim = await gen.generate_claim(analysis_result, contract_info)
    
    # Иск
    lawsuit = await gen.generate_lawsuit(analysis_result, case_info)
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from src.schemas.legal import (
    LegalAnalysisResult,
    LegalFinding,
    LegalSeverity,
    LegalVerdict,
    ProtocolDisagreements,
    ProtocolItem,
    ProtocolPartyInfo,
)

logger = logging.getLogger(__name__)


# =============================================================================
# LLM Prompts
# =============================================================================

PROTOCOL_INTRO_PROMPT = """Ты — юрист строительной субподрядной организации. 
Составь пояснительную записку к протоколу разногласий.

Договор: {contract_number} от {contract_date}
Стороны: Подрядчик ({contractor_name}) и Заказчик ({customer_name})

Найдено проблем: {total_issues} (критических: {critical}, высоких: {high})

Ключевые риски:
{key_risks}

Напиши краткую пояснительную записку (3-5 предложений) в деловом стиле.
Укажи, что протокол составлен в соответствии с ГК РФ и позицией ВС РФ.
Ответ: только текст записки, без заголовков.
"""

CLAIM_GENERATION_PROMPT = """Ты — юрист строительной субподрядной организации.
Составь досудебную претензию на основе результатов юридического анализа.

ДАННЫЕ ДЛЯ ПРЕТЕНЗИИ:
- Адресат (Заказчик): {customer_name}, ИНН {customer_inn}
- Адрес: {customer_address}
- Отправитель (Подрядчик): {contractor_name}, ИНН {contractor_inn}
- Адрес: {contractor_address}
- Договор: № {contract_number} от {contract_date}
- Предмет договора: {contract_subject}
- Сумма требований: {claim_amount} руб.

НАРУШЕНИЯ:
{violations}

СУММА ТРЕБОВАНИЙ:
{amount_breakdown}

ПРАВОВОЕ ОБОСНОВАНИЕ:
{legal_basis}

Составь текст досудебной претензии. Строгий деловой стиль.
Структура:
1. Шапка (кому, от кого)
2. Описание обязательств (договор)
3. Описание нарушений (факты)
4. Правовое обоснование (статьи)
5. Требования (конкретно)
6. Расчёт суммы требований
7. Предупреждение об обращении в суд
8. Приложения

Формат — чистый текст для вставки в документ.
"""

LAWSUIT_GENERATION_PROMPT = """Ты — юрист, специализирующийся на арбитражных спорах 
в строительстве. Составь исковое заявление в арбитражный суд.

ДАННЫЕ ДЕЛА:
- Истец: {contractor_name}, ИНН {contractor_inn}, ОГРН {contractor_ogrn}
- Адрес истца: {contractor_address}
- Ответчик: {customer_name}, ИНН {customer_inn}, ОГРН {customer_ogrn}
- Адрес ответчика: {customer_address}
- Договор: № {contract_number} от {contract_date}
- Цена иска: {claim_amount} руб.
- Госпошлина: {state_duty} руб.
- Арбитражный суд: {court_name}
- Подсудность: {jurisdiction_basis}

ОБСТОЯТЕЛЬСТВА ДЕЛА:
{case_facts}

ПРАВОВОЕ ОБОСНОВАНИЕ:
{legal_basis}

ДОСУДЕБНОЕ УРЕГУЛИРОВАНИЕ:
{pre_trial_info}

Составь текст искового заявления. Структура (АПК РФ ст. 125):
1. Шапка (суд, истец, ответчик, цена иска, госпошлина)
2. Описательная часть (обстоятельства дела)
3. Мотивировочная часть (правовое обоснование)
4. Просительная часть (конкретные требования)
5. Приложения
6. Подпись, дата

Ссылайся на конкретные статьи: ГК РФ (гл. 37), АПК РФ, Постановления Пленума ВС РФ.
Формат — чистый текст.
"""

MOTIVATED_REFUSAL_PROMPT = """Ты — юрист строительной субподрядной организации.
Составь мотивированный отказ от подписания актов выполненных работ 
(КС-2, КС-3, АОСР) в соответствии со ст. 753 ГК РФ.

ДАННЫЕ:
- Отправитель (Подрядчик): {contractor_name}, ИНН {contractor_inn}
- Адрес: {contractor_address}
- Адресат (Заказчик): {customer_name}, ИНН {customer_inn}
- Адрес: {customer_address}
- Договор: № {contract_number} от {contract_date}
- Предмет договора: {contract_subject}
- Акты, от подписания которых отказ: {act_list}

ОСНОВАНИЯ ДЛЯ ОТКАЗА:
{refusal_grounds}

ПРАВОВОЕ ОБОСНОВАНИЕ:
{legal_basis}

Составь текст мотивированного отказа. Строгий деловой стиль.
Структура:
1. Шапка (кому, от кого, исх. номер, дата)
2. Ссылка на договор и направленные акты
3. Перечень актов с обоснованием отказа по каждому
4. Мотивировочная часть: конкретные причины
5. Правовое обоснование (ст. 753 ГК РФ)
6. Требование: подписать акты в течение 5 рабочих дней
7. Предупреждение об одностороннем подписании (п. 4 ст. 753 ГК РФ)
8. Приложения

Формат — чистый текст.
"""


# =============================================================================
# Legal DOCX Formatter
# =============================================================================

class LegalDocxFormatter:
    """Единый форматер DOCX для юридических документов ASD.

    Стандарт: Times New Roman 12pt, A4, левое поле 3 см (ГОСТ Р 7.0.97),
    межстрочный 1.15.
    """
    FONT_NAME = "Times New Roman"
    FONT_SIZE = 12
    FONT_SIZE_SMALL = 10
    MARGIN_LEFT_CM = 3.0
    MARGIN_RIGHT_CM = 1.5
    MARGIN_TOP_CM = 2.0
    MARGIN_BOTTOM_CM = 2.0

    def __init__(self):
        self._doc = None

    def create_document(self):
        import docx as docx_lib
        from docx.shared import Pt, Cm
        doc = docx_lib.Document()
        for section in doc.sections:
            section.left_margin = Cm(self.MARGIN_LEFT_CM)
            section.right_margin = Cm(self.MARGIN_RIGHT_CM)
            section.top_margin = Cm(self.MARGIN_TOP_CM)
            section.bottom_margin = Cm(self.MARGIN_BOTTOM_CM)
        style = doc.styles["Normal"]
        style.font.name = self.FONT_NAME
        style.font.size = Pt(self.FONT_SIZE)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(6)
        self._doc = doc
        return doc

    def add_title(self, text, level=0):
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if level == 0:
            p = self._doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = self.FONT_NAME
        else:
            h = self._doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = self.FONT_NAME

    def add_paragraph(self, text, bold=False, italic=False):
        p = self._doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = self.FONT_NAME
        return p

    def add_section_header(self, text):
        from docx.shared import Pt
        p = self._doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(text)
        run.bold = True
        run.font.name = self.FONT_NAME

    def add_bullet(self, text):
        p = self._doc.add_paragraph(text, style="List Bullet")
        for run in p.runs:
            run.font.name = self.FONT_NAME

    def add_signature_block(self, party_name, label="", date_str=None):
        if date_str is None:
            date_str = datetime.now().strftime("%d.%m.%Y")
        self.add_paragraph("")
        prefix = f"{label}: " if label else ""
        self.add_paragraph(f"{prefix}________________ /_____________/")
        self.add_paragraph(f"Дата: {date_str}")
        if party_name:
            self.add_paragraph(f"({party_name})", italic=True)

    def add_attachment_list(self, attachments):
        self.add_section_header("ПРИЛОЖЕНИЯ:")
        for i, att in enumerate(attachments, 1):
            self.add_paragraph(f"{i}. {att}")

    def save(self, filepath):
        Path(filepath).parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(filepath)
        return filepath


# =============================================================================
# Legal Document Generator
# =============================================================================

class LegalDocumentGenerator:
    """
    Генератор юридических документов ASD v12.0.

    Принимает результат LegalService.analyze() и генерирует:
      - Протокол разногласий (DOCX, 3 колонки)
      - Досудебную претензию
      - Исковое заявление
    """

    def __init__(self, llm_engine=None):
        from src.core.llm_engine import llm_engine as default_llm
        self._llm = llm_engine or default_llm
        self._formatter = LegalDocxFormatter()

    # =========================================================================
    # Protocol of Disagreements (DOCX)
    # =========================================================================

    async def generate_protocol(
        self,
        analysis_result: LegalAnalysisResult,
        contract_number: str = "",
        contract_date: str = "",
        customer_info: Optional[ProtocolPartyInfo] = None,
        contractor_info: Optional[ProtocolPartyInfo] = None,
        output_dir: str = "/tmp",
    ) -> str:
        """
        Сгенерировать протокол разногласий в формате DOCX.

        Args:
            analysis_result: результат юридической экспертизы
            contract_number: номер договора
            contract_date: дата договора
            customer_info: реквизиты Заказчика
            contractor_info: реквизиты Подрядчика
            output_dir: директория для сохранения

        Returns:
            Путь к сгенерированному DOCX-файлу
        """
        # Формируем ProtocolDisagreements из findings
        protocol = self._build_protocol_from_findings(
            analysis_result, contract_number, contract_date,
            customer_info, contractor_info,
        )

        # Генерируем пояснительную записку через LLM
        summary = await self._generate_protocol_summary(
            protocol, analysis_result,
        )
        protocol.summary = summary

        # Экспорт в DOCX
        docx_path = self._export_protocol_docx(protocol, output_dir)

        logger.info(
            "Protocol generated: %d items → %s",
            protocol.total_items, docx_path,
        )
        return docx_path

    def _build_protocol_from_findings(
        self,
        result: LegalAnalysisResult,
        contract_number: str,
        contract_date: str,
        customer_info: Optional[ProtocolPartyInfo],
        contractor_info: Optional[ProtocolPartyInfo],
    ) -> ProtocolDisagreements:
        """Собрать ProtocolDisagreements из findings анализа."""

        # Фильтруем: только HIGH и CRITICAL
        actionable = [
            f for f in result.findings
            if f.severity in (LegalSeverity.CRITICAL, LegalSeverity.HIGH)
        ]

        # Если уже есть protocol_items в результате — используем их
        if result.protocol_items:
            items = result.protocol_items
        else:
            # Строим из findings
            items = []
            for i, f in enumerate(actionable, 1):
                contractor_edit = f.contractor_edit or self._default_edit(f)
                items.append(ProtocolItem(
                    row_number=i,
                    clause_ref=f.clause_ref,
                    customer_text=f.issue,  # Проблемная формулировка
                    contractor_text=contractor_edit,
                    legal_basis=f.legal_basis,
                    severity=f.severity,
                    blc_match=f.blc_match,
                ))

        return ProtocolDisagreements(
            protocol_title=f"Протокол разногласий к Договору № {contract_number}",
            contract_number=contract_number,
            contract_date=contract_date,
            customer_info=customer_info or ProtocolPartyInfo(name="Заказчик"),
            contractor_info=contractor_info or ProtocolPartyInfo(name="Подрядчик"),
            items=items,
            total_items=len(items),
        )

    async def _generate_protocol_summary(
        self,
        protocol: ProtocolDisagreements,
        result: LegalAnalysisResult,
    ) -> str:
        """Генерация пояснительной записки через LLM."""
        key_risks = "\n".join(
            f"- [{f.severity.value}] п. {f.clause_ref}: {f.issue[:120]}"
            for f in result.findings
            if f.severity in (LegalSeverity.CRITICAL, LegalSeverity.HIGH)
        )[:1500]

        prompt = PROTOCOL_INTRO_PROMPT.format(
            contract_number=protocol.contract_number or "___",
            contract_date=protocol.contract_date or "___",
            contractor_name=protocol.contractor_info.name,
            customer_name=protocol.customer_info.name,
            total_issues=protocol.total_items,
            critical=result.critical_count,
            high=result.high_count,
            key_risks=key_risks or "Не указаны",
        )

        try:
            response = await self._llm.chat(
                "legal",
                [{"role": "user", "content": prompt}],
            )
            return response.strip()
        except Exception as e:
            logger.warning("Protocol summary LLM failed: %s", e)
            return (
                "Настоящий протокол разногласий составлен в соответствии с "
                "действующим законодательством РФ. Просим рассмотреть "
                "предложенные редакции пунктов договора."
            )

    def _export_protocol_docx(
        self, protocol: ProtocolDisagreements, output_dir: str
    ) -> str:
        """Экспорт протокола разногласий в DOCX (3 колонки, единый стиль)."""
        try:
            import docx as docx_lib
            from docx.shared import Inches, Pt
        except ImportError:
            logger.warning("python-docx not available — generating text file instead")
            return self._export_protocol_text(protocol, output_dir)

        fmt = self._formatter
        doc = fmt.create_document()

        # Title
        fmt.add_title(protocol.protocol_title)
        fmt.add_paragraph(
            f"к Договору № {protocol.contract_number or '___'} "
            f"от {protocol.contract_date or '___'}",
            italic=True,
        )
        fmt.add_paragraph("")

        # Parties
        cust = protocol.customer_info
        contr = protocol.contractor_info
        fmt.add_section_header("СТОРОНЫ:")
        fmt.add_paragraph(
            f"Заказчик: {cust.name}"
            + (f", ИНН {cust.inn}" if cust.inn else "")
            + (f", {cust.legal_address}" if cust.legal_address else "")
        )
        fmt.add_paragraph(
            f"Подрядчик: {contr.name}"
            + (f", ИНН {contr.inn}" if contr.inn else "")
            + (f", {contr.legal_address}" if contr.legal_address else "")
        )
        fmt.add_paragraph("")

        # Summary
        if protocol.summary:
            fmt.add_paragraph(protocol.summary, italic=True)
            fmt.add_paragraph("")

        # Table: 3 columns
        if protocol.items:
            fmt.add_section_header("ТАБЛИЦА РАЗНОГЛАСИЙ:")
            fmt.add_paragraph("")
            table = doc.add_table(rows=1, cols=3, style="Table Grid")

            hdr = table.rows[0].cells
            hdr[0].text = "Пункт, статья договора"
            hdr[1].text = "Редакция Заказчика"
            hdr[2].text = "Редакция Подрядчика"
            for cell in hdr:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.name = fmt.FONT_NAME
                        run.font.size = Pt(fmt.FONT_SIZE_SMALL)

            for item in protocol.items:
                row = table.add_row().cells
                sev = item.severity.value if item.severity else "high"
                row[0].text = (
                    f"{item.row_number}. {item.clause_ref}\n"
                    f"[{sev}]\n"
                    f"{item.legal_basis}"
                )
                row[1].text = item.customer_text
                row[2].text = item.contractor_text
                for cell in row:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = fmt.FONT_NAME
                            run.font.size = Pt(fmt.FONT_SIZE_SMALL)

            for row_obj in table.rows:
                row_obj.cells[0].width = Inches(2.0)
                row_obj.cells[1].width = Inches(2.5)
                row_obj.cells[2].width = Inches(2.5)

        fmt.add_paragraph("")
        fmt.add_paragraph(f"Всего пунктов разногласий: {protocol.total_items}", bold=True)

        # Signatures
        fmt.add_signature_block(contr.name, "Подрядчик")
        fmt.add_signature_block(cust.name, "Заказчик")

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = str(Path(output_dir) / f"protocol_disagreements_{ts}.docx")
        return fmt.save(filepath)

    def _export_protocol_text(
        self, protocol: ProtocolDisagreements, output_dir: str
    ) -> str:
        """Fallback: текстовый протокол если python-docx недоступен."""
        lines = [protocol.protocol_title, "=" * 60, ""]
        lines.append(f"Заказчик: {protocol.customer_info.name}")
        lines.append(f"Подрядчик: {protocol.contractor_info.name}")
        if protocol.summary:
            lines.extend(["", protocol.summary])
        lines.extend(["", "-" * 60])

        for item in protocol.items:
            lines.append(
                f"\n{item.row_number}. {item.clause_ref} "
                f"[{item.severity.value if item.severity else 'high'}]"
            )
            lines.append(f"   Основание: {item.legal_basis}")
            lines.append(f"   Заказчик:  {item.customer_text[:200]}")
            lines.append(f"   Подрядчик: {item.contractor_text[:200]}")

        lines.extend(["", f"Всего: {protocol.total_items} пунктов"])

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = str(Path(output_dir) / f"protocol_disagreements_{ts}.txt")
        Path(filepath).write_text("\n".join(lines), encoding="utf-8")
        return filepath

    # =========================================================================
    # Pre-trial Claim (Досудебная претензия)
    # =========================================================================

    async def generate_claim(
        self,
        analysis_result: LegalAnalysisResult,
        contract_number: str = "",
        contract_date: str = "",
        contract_subject: str = "",
        customer_name: str = "",
        customer_inn: str = "",
        customer_address: str = "",
        contractor_name: str = "",
        contractor_inn: str = "",
        contractor_address: str = "",
        claim_amount: float = 0.0,
        output_dir: str = "/tmp",
    ) -> str:
        """
        Сгенерировать досудебную претензию.

        Args:
            analysis_result: результат юридического анализа
            contract_number: номер договора
            contract_date: дата договора
            contract_subject: предмет договора
            customer_name: Заказчик
            customer_inn: ИНН Заказчика
            customer_address: адрес Заказчика
            contractor_name: Подрядчик
            contractor_inn: ИНН Подрядчика
            contractor_address: адрес Подрядчика
            claim_amount: сумма требований
            output_dir: директория для сохранения

        Returns:
            Путь к текстовому файлу претензии
        """
        # Собираем нарушения
        violations = self._format_violations(analysis_result)
        legal_basis = self._format_legal_basis(analysis_result)

        # Расчёт суммы
        amount_breakdown = (
            f"Основной долг: {claim_amount:,.2f} руб.\n"
            f"Проценты по ст. 395 ГК РФ: расчёт прилагается\n"
            f"Судебные расходы: будут предъявлены дополнительно"
        )

        prompt = CLAIM_GENERATION_PROMPT.format(
            customer_name=customer_name or "Заказчик",
            customer_inn=customer_inn or "___",
            customer_address=customer_address or "___",
            contractor_name=contractor_name or "Подрядчик",
            contractor_inn=contractor_inn or "___",
            contractor_address=contractor_address or "___",
            contract_number=contract_number or "___",
            contract_date=contract_date or "___",
            contract_subject=contract_subject or "выполнение строительных работ",
            claim_amount=f"{claim_amount:,.2f}",
            violations=violations,
            amount_breakdown=amount_breakdown,
            legal_basis=legal_basis,
        )

        try:
            claim_text = await self._llm.chat(
                "legal",
                [{"role": "user", "content": prompt}],
            )
        except Exception as e:
            logger.error("Claim generation LLM failed: %s", e)
            claim_text = self._fallback_claim(
                contractor_name, customer_name, contract_number, claim_amount, violations, legal_basis
            )

        return self._export_claim_docx(
            claim_text=claim_text,
            contractor_name=contractor_name,
            contractor_inn=contractor_inn,
            contractor_address=contractor_address,
            customer_name=customer_name,
            customer_inn=customer_inn,
            customer_address=customer_address,
            contract_number=contract_number,
            contract_date=contract_date,
            claim_amount=claim_amount,
            output_dir=output_dir,
        )

    def _fallback_claim(
        self, contractor: str, customer: str, contract_num: str,
        amount: float, violations: str, legal_basis: str,
    ) -> str:
        """Fallback-текст претензии если LLM недоступна."""
        today = datetime.now().strftime('%d.%m.%Y')
        return f"""ДОСУДЕБНАЯ ПРЕТЕНЗИЯ

От: {contractor}
Кому: {customer}

По договору № {contract_num}

Уважаемый Заказчик!

На основании договора № {contract_num} Подрядчиком были выполнены работы.
Однако Заказчиком допущены следующие нарушения:

{violations}

ПРАВОВОЕ ОБОСНОВАНИЕ:
{legal_basis}

ТРЕБОВАНИЕ:
Оплатить задолженность в размере {amount:,.2f} руб. в течение 10 календарных дней с момента получения претензии.

В случае неисполнения требований Подрядчик будет вынужден обратиться в Арбитражный суд.

Подрядчик: ______________ /_____________/
Дата: {today}
"""

    # =========================================================================
    # Lawsuit (Исковое заявление)
    # =========================================================================

    async def generate_lawsuit(
        self,
        analysis_result: LegalAnalysisResult,
        case_facts: str = "",
        contract_number: str = "",
        contract_date: str = "",
        contractor_name: str = "",
        contractor_inn: str = "",
        contractor_ogrn: str = "",
        contractor_address: str = "",
        customer_name: str = "",
        customer_inn: str = "",
        customer_ogrn: str = "",
        customer_address: str = "",
        claim_amount: float = 0.0,
        court_name: str = "Арбитражный суд города Москвы",
        jurisdiction_basis: str = "ст. 35 АПК РФ — по месту нахождения ответчика",
        pre_trial_info: str = "",
        output_dir: str = "/tmp",
    ) -> str:
        """
        Сгенерировать исковое заявление в арбитражный суд.

        Returns:
            Путь к текстовому файлу искового заявления
        """
        # Расчёт госпошлины (ст. 333.21 НК РФ)
        state_duty = self._calc_state_duty(claim_amount)

        legal_basis = self._format_legal_basis(analysis_result)

        prompt = LAWSUIT_GENERATION_PROMPT.format(
            contractor_name=contractor_name or "Истец",
            contractor_inn=contractor_inn or "___",
            contractor_ogrn=contractor_ogrn or "___",
            contractor_address=contractor_address or "___",
            customer_name=customer_name or "Ответчик",
            customer_inn=customer_inn or "___",
            customer_ogrn=customer_ogrn or "___",
            customer_address=customer_address or "___",
            contract_number=contract_number or "___",
            contract_date=contract_date or "___",
            claim_amount=f"{claim_amount:,.2f}",
            state_duty=f"{state_duty:,.2f}",
            court_name=court_name,
            jurisdiction_basis=jurisdiction_basis,
            case_facts=case_facts or self._format_case_facts(analysis_result),
            legal_basis=legal_basis,
            pre_trial_info=pre_trial_info or "Досудебная претензия направлена, ответ не получен.",
        )

        try:
            lawsuit_text = await self._llm.chat(
                "legal",
                [{"role": "user", "content": prompt}],
            )
        except Exception as e:
            logger.error("Lawsuit generation LLM failed: %s", e)
            lawsuit_text = self._fallback_lawsuit(
                contractor_name, customer_name, contract_number,
                claim_amount, state_duty, court_name, legal_basis, case_facts
            )

        return self._export_lawsuit_docx(
            lawsuit_text=lawsuit_text,
            contractor_name=contractor_name,
            contractor_inn=contractor_inn,
            contractor_ogrn=contractor_ogrn,
            contractor_address=contractor_address,
            customer_name=customer_name,
            customer_inn=customer_inn,
            customer_ogrn=customer_ogrn,
            customer_address=customer_address,
            contract_number=contract_number,
            contract_date=contract_date,
            claim_amount=claim_amount,
            state_duty=state_duty,
            court_name=court_name,
            output_dir=output_dir,
        )

    def _fallback_lawsuit(
        self, contractor: str, customer: str, contract_num: str,
        amount: float, duty: float, court: str,
        legal_basis: str, facts: str,
    ) -> str:
        """Fallback-текст иска если LLM недоступна."""
        return f"""В {court}

ИСТЕЦ: {contractor}
ОТВЕТЧИК: {customer}

ЦЕНА ИСКА: {amount:,.2f} руб.
ГОСПОШЛИНА: {duty:,.2f} руб.

ИСКОВОЕ ЗАЯВЛЕНИЕ
о взыскании задолженности по договору подряда

ОБСТОЯТЕЛЬСТВА ДЕЛА:
{facts}

ПРАВОВОЕ ОБОСНОВАНИЕ:
{legal_basis}

ПРОШУ СУД:
1. Взыскать с Ответчика задолженность в размере {amount:,.2f} руб.
2. Взыскать с Ответчика проценты по ст. 395 ГК РФ.
3. Взыскать с Ответчика расходы по уплате госпошлины.

ПРИЛОЖЕНИЯ:
1. Копия договора № {contract_num}
2. Акты КС-2, КС-3
3. Досудебная претензия
4. Расчёт задолженности
5. Квитанция об уплате госпошлины
6. Выписки из ЕГРЮЛ

Истец: ______________ /_____________/
Дата: {datetime.now().strftime('%d.%m.%Y')}
"""

    # =========================================================================
    # DOCX Export Methods
    # =========================================================================

    def _export_claim_docx(
        self, claim_text, contractor_name, contractor_inn, contractor_address,
        customer_name, customer_inn, customer_address,
        contract_number, contract_date, claim_amount, output_dir,
    ):
        """Экспорт претензии в DOCX с единым стилем."""
        try:
            import docx
        except ImportError:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            filepath = str(Path(output_dir) / f"pre_trial_claim_{ts}.txt")
            Path(filepath).write_text(claim_text, encoding="utf-8")
            return filepath

        fmt = self._formatter
        doc = fmt.create_document()
        fmt.add_title("ДОСУДЕБНАЯ ПРЕТЕНЗИЯ")
        fmt.add_paragraph("")
        fmt.add_section_header("АДРЕСАТ:")
        fmt.add_paragraph(f"{customer_name or 'Заказчик'}")
        if customer_inn:
            fmt.add_paragraph(f"ИНН: {customer_inn}")
        if customer_address:
            fmt.add_paragraph(f"Адрес: {customer_address}")
        fmt.add_paragraph("")
        fmt.add_section_header("ОТПРАВИТЕЛЬ:")
        fmt.add_paragraph(f"{contractor_name or 'Подрядчик'}")
        if contractor_inn:
            fmt.add_paragraph(f"ИНН: {contractor_inn}")
        if contractor_address:
            fmt.add_paragraph(f"Адрес: {contractor_address}")
        fmt.add_paragraph("")
        fmt.add_paragraph(
            f"По договору № {contract_number or '___'} от {contract_date or '___'}",
            bold=True,
        )
        fmt.add_paragraph("")

        for para in claim_text.split("\n"):
            para = para.strip()
            if not para:
                fmt.add_paragraph("")
                continue
            if any(para.startswith(f"{i}.") for i in range(1, 10)):
                fmt.add_section_header(para)
            elif para.startswith("- ") or para.startswith("• "):
                fmt.add_bullet(para[2:])
            else:
                fmt.add_paragraph(para)

        fmt.add_attachment_list([
            f"Копия договора № {contract_number or '___'}",
            "Акты выполненных работ (КС-2, КС-3)",
            "Расчёт задолженности",
            "Расчёт процентов по ст. 395 ГК РФ",
            "Документы, подтверждающие полномочия подписанта",
        ])
        fmt.add_signature_block(contractor_name or "Подрядчик", "Подрядчик")

        logger.info("Claim DOCX generated: %s", claim_amount)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = str(Path(output_dir) / f"pre_trial_claim_{ts}.docx")
        return fmt.save(filepath)

    def _export_lawsuit_docx(
        self, lawsuit_text, contractor_name, contractor_inn, contractor_ogrn,
        contractor_address, customer_name, customer_inn, customer_ogrn,
        customer_address, contract_number, contract_date,
        claim_amount, state_duty, court_name, output_dir,
    ):
        """Экспорт искового заявления в DOCX."""
        try:
            import docx
        except ImportError:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            filepath = str(Path(output_dir) / f"lawsuit_{ts}.txt")
            Path(filepath).write_text(lawsuit_text, encoding="utf-8")
            return filepath

        fmt = self._formatter
        doc = fmt.create_document()
        fmt.add_paragraph(f"В {court_name}", bold=True)
        fmt.add_paragraph("")
        fmt.add_section_header("ИСТЕЦ:")
        fmt.add_paragraph(f"{contractor_name or 'Истец'}")
        if contractor_inn:
            fmt.add_paragraph(f"ИНН: {contractor_inn}")
        if contractor_ogrn:
            fmt.add_paragraph(f"ОГРН: {contractor_ogrn}")
        if contractor_address:
            fmt.add_paragraph(f"Адрес: {contractor_address}")
        fmt.add_paragraph("")
        fmt.add_section_header("ОТВЕТЧИК:")
        fmt.add_paragraph(f"{customer_name or 'Ответчик'}")
        if customer_inn:
            fmt.add_paragraph(f"ИНН: {customer_inn}")
        if customer_ogrn:
            fmt.add_paragraph(f"ОГРН: {customer_ogrn}")
        if customer_address:
            fmt.add_paragraph(f"Адрес: {customer_address}")
        fmt.add_paragraph("")
        fmt.add_paragraph(f"ЦЕНА ИСКА: {claim_amount:,.2f} руб.", bold=True)
        fmt.add_paragraph(f"ГОСПОШЛИНА: {state_duty:,.2f} руб.", bold=True)
        fmt.add_paragraph("")
        fmt.add_title("ИСКОВОЕ ЗАЯВЛЕНИЕ")
        fmt.add_paragraph(
            f"о взыскании задолженности по договору подряда "
            f"№ {contract_number or '___'} от {contract_date or '___'}",
            italic=True,
        )
        fmt.add_paragraph("")

        for para in lawsuit_text.split("\n"):
            para = para.strip()
            if not para:
                fmt.add_paragraph("")
                continue
            if any(para.startswith(f"{i}.") for i in range(1, 15)):
                fmt.add_section_header(para)
            elif para.startswith("- ") or para.startswith("• "):
                fmt.add_bullet(para[2:])
            else:
                fmt.add_paragraph(para)

        fmt.add_attachment_list([
            "Копия договора подряда",
            "Акты КС-2, справки КС-3",
            "Досудебная претензия с доказательством направления",
            "Расчёт задолженности и процентов по ст. 395 ГК РФ",
            "Квитанция об уплате госпошлины",
            "Выписка из ЕГРЮЛ на Истца",
            "Выписка из ЕГРЮЛ на Ответчика",
            "Документы, подтверждающие полномочия подписанта",
        ])
        fmt.add_signature_block(contractor_name or "Истец", "Истец")

        logger.info("Lawsuit DOCX generated: %s", claim_amount)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = str(Path(output_dir) / f"lawsuit_{ts}.docx")
        return fmt.save(filepath)

    # =========================================================================
    # Motivated Refusal
    # =========================================================================

    async def generate_motivated_refusal(
        self,
        analysis_result: LegalAnalysisResult,
        act_numbers: Optional[List[str]] = None,
        refusal_grounds: str = "",
        contract_number: str = "",
        contract_date: str = "",
        contract_subject: str = "",
        customer_name: str = "",
        customer_inn: str = "",
        customer_address: str = "",
        contractor_name: str = "",
        contractor_inn: str = "",
        contractor_address: str = "",
        output_dir: str = "/tmp",
    ) -> str:
        """Сгенерировать мотивированный отказ от подписания актов (ст. 753 ГК РФ)."""
        act_list = ", ".join(act_numbers) if act_numbers else "АОСР, КС-2, КС-3"
        grounds = refusal_grounds or self._format_refusal_grounds(analysis_result)
        legal_basis = self._format_legal_basis(analysis_result)

        prompt = MOTIVATED_REFUSAL_PROMPT.format(
            contractor_name=contractor_name or "Подрядчик",
            contractor_inn=contractor_inn or "___",
            contractor_address=contractor_address or "___",
            customer_name=customer_name or "Заказчик",
            customer_inn=customer_inn or "___",
            customer_address=customer_address or "___",
            contract_number=contract_number or "___",
            contract_date=contract_date or "___",
            contract_subject=contract_subject or "выполнение строительных работ",
            act_list=act_list,
            refusal_grounds=grounds,
            legal_basis=legal_basis,
        )

        try:
            refusal_text = await self._llm.chat(
                "legal", [{"role": "user", "content": prompt}]
            )
        except Exception as e:
            logger.error("Motivated refusal LLM failed: %s", e)
            refusal_text = self._fallback_refusal(
                contractor_name, customer_name, contract_number,
                act_list, grounds, legal_basis,
            )

        return self._export_refusal_docx(
            refusal_text=refusal_text,
            contractor_name=contractor_name,
            contractor_inn=contractor_inn,
            contractor_address=contractor_address,
            customer_name=customer_name,
            customer_inn=customer_inn,
            customer_address=customer_address,
            contract_number=contract_number,
            contract_date=contract_date,
            act_numbers=act_numbers or [],
            output_dir=output_dir,
        )

    def _export_refusal_docx(
        self, refusal_text, contractor_name, contractor_inn, contractor_address,
        customer_name, customer_inn, customer_address,
        contract_number, contract_date, act_numbers, output_dir,
    ):
        """Экспорт мотивированного отказа в DOCX."""
        try:
            import docx
        except ImportError:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            filepath = str(Path(output_dir) / f"motivated_refusal_{ts}.txt")
            Path(filepath).write_text(refusal_text, encoding="utf-8")
            return filepath

        fmt = self._formatter
        doc = fmt.create_document()
        today = datetime.now().strftime("%d.%m.%Y")
        fmt.add_paragraph(f"Исх. № ___ от {today}", bold=True)
        fmt.add_paragraph("")
        fmt.add_section_header("АДРЕСАТ:")
        fmt.add_paragraph(f"{customer_name or 'Заказчик'}")
        if customer_inn:
            fmt.add_paragraph(f"ИНН: {customer_inn}")
        if customer_address:
            fmt.add_paragraph(f"Адрес: {customer_address}")
        fmt.add_paragraph("")
        fmt.add_section_header("ОТПРАВИТЕЛЬ:")
        fmt.add_paragraph(f"{contractor_name or 'Подрядчик'}")
        if contractor_inn:
            fmt.add_paragraph(f"ИНН: {contractor_inn}")
        if contractor_address:
            fmt.add_paragraph(f"Адрес: {contractor_address}")
        fmt.add_paragraph("")
        fmt.add_title("МОТИВИРОВАННЫЙ ОТКАЗ")
        fmt.add_paragraph(
            f"от подписания актов выполненных работ "
            f"по договору № {contract_number or '___'} от {contract_date or '___'}",
            italic=True,
        )
        fmt.add_paragraph("")

        for para in refusal_text.split("\n"):
            para = para.strip()
            if not para:
                fmt.add_paragraph("")
                continue
            if any(para.startswith(f"{i}.") for i in range(1, 15)):
                fmt.add_section_header(para)
            elif para.startswith("- ") or para.startswith("• "):
                fmt.add_bullet(para[2:])
            else:
                fmt.add_paragraph(para)

        if act_numbers:
            fmt.add_section_header("АКТЫ, ОТ ПОДПИСАНИЯ КОТОРЫХ ПОДРЯДЧИК ОТКАЗЫВАЕТСЯ:")
            for act in act_numbers:
                fmt.add_bullet(act)

        attachments = [f"Копия договора № {contract_number or '___'}"]
        for act in act_numbers:
            attachments.append(f"Копия акта {act}")
        attachments.extend([
            "Документы, подтверждающие основания для отказа",
            "Документы, подтверждающие полномочия подписанта",
        ])
        fmt.add_attachment_list(attachments)
        fmt.add_signature_block(contractor_name or "Подрядчик", "Подрядчик")

        logger.info("Motivated refusal DOCX generated")
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = str(Path(output_dir) / f"motivated_refusal_{ts}.docx")
        return fmt.save(filepath)

    def _fallback_refusal(
        self, contractor, customer, contract_num, act_list, grounds, legal_basis,
    ):
        """Fallback-текст мотивированного отказа если LLM недоступна."""
        today = datetime.now().strftime("%d.%m.%Y")
        return (
            f"МОТИВИРОВАННЫЙ ОТКАЗ\nот подписания актов выполненных работ\n\n"
            f"Исх. № ___ от {today}\n\nОт: {contractor}\nКому: {customer}\n\n"
            f"По договору № {contract_num}\n\n"
            f"Подрядчик ОТКАЗЫВАЕТСЯ от подписания актов ({act_list})\n"
            f"по следующим основаниям:\n{grounds}\n\n"
            f"ПРАВОВОЕ ОБОСНОВАНИЕ:\n{legal_basis}\n\n"
            f"Ст. 753 ГК РФ: заказчик вправе отказаться от приёмки только\n"
            f"при обнаружении недостатков, исключающих использование результата.\n\n"
            f"ТРЕБОВАНИЕ: подписать акты в течение 5 рабочих дней.\n"
            f"При необоснованном уклонении акты считаются подписанными\n"
            f"в одностороннем порядке (п. 4 ст. 753 ГК РФ).\n\n"
            f"Подрядчик: ______________ /_____________/\nДата: {today}"
        )

    # =========================================================================
    # Helpers
    # =========================================================================

    @staticmethod
    def _format_violations(result: LegalAnalysisResult) -> str:
        """Форматировать список нарушений для претензии."""
        violations = []
        for i, f in enumerate(result.findings, 1):
            if f.severity in (LegalSeverity.CRITICAL, LegalSeverity.HIGH):
                violations.append(
                    f"{i}. {f.issue}\n"
                    f"   Пункт договора: {f.clause_ref}\n"
                    f"   Основание: {f.legal_basis}"
                )
        return "\n\n".join(violations) if violations else "Нарушения не детализированы."

    @staticmethod
    def _format_legal_basis(result: LegalAnalysisResult) -> str:
        """Форматировать правовое обоснование."""
        refs = result.normative_refs or []
        # Добавляем ссылки из findings
        for f in result.findings:
            if f.legal_basis and f.legal_basis not in refs:
                refs.append(f.legal_basis)

        if not refs:
            refs = [
                "ст. 309 ГК РФ (обязательства должны исполняться надлежащим образом)",
                "ст. 310 ГК РФ (недопустимость одностороннего отказа от исполнения)",
                "ст. 711 ГК РФ (порядок оплаты работы)",
                "ст. 395 ГК РФ (проценты за пользование чужими денежными средствами)",
            ]

        return "\n".join(f"- {ref}" for ref in refs[:10])

    @staticmethod
    def _format_case_facts(result: LegalAnalysisResult) -> str:
        """Форматировать обстоятельства дела."""
        critical_findings = [
            f for f in result.findings
            if f.severity == LegalSeverity.CRITICAL
        ]
        if critical_findings:
            return "\n".join(
                f"- {f.issue} (п. {f.clause_ref})" for f in critical_findings[:5]
            )
        return result.summary or "Обстоятельства изложены в прилагаемых документах."

    @staticmethod
    def _format_refusal_grounds(result: LegalAnalysisResult) -> str:
        """Форматировать основания для мотивированного отказа."""
        findings = [
            f for f in result.findings
            if f.severity in (LegalSeverity.CRITICAL, LegalSeverity.HIGH)
        ]
        if not findings:
            return (
                "1. Работы выполнены с нарушением требований ПД и НД.\n"
                "2. Отсутствует полный комплект ИД (Приказ 344/пр).\n"
                "3. Результат работ не может быть использован по назначению."
            )
        return "\n".join(
            f"{i}. {f.issue} (п. {f.clause_ref})"
            for i, f in enumerate(findings, 1)
        )

    @staticmethod
    def _calc_state_duty(claim_amount: float) -> float:
        """Расчёт госпошлины по ст. 333.21 НК РФ (арбитраж)."""
        if claim_amount <= 100_000:
            return max(2000, claim_amount * 0.04)
        elif claim_amount <= 200_000:
            return 4000 + (claim_amount - 100_000) * 0.03
        elif claim_amount <= 1_000_000:
            return 7000 + (claim_amount - 200_000) * 0.02
        elif claim_amount <= 2_000_000:
            return 23_000 + (claim_amount - 1_000_000) * 0.01
        else:
            return min(200_000, 33_000 + (claim_amount - 2_000_000) * 0.005)

    @staticmethod
    def _default_edit(finding: LegalFinding) -> str:
        """Дефолтная редакция Подрядчика если не указана явно."""
        if finding.auto_fixable:
            return (
                f"Предлагается исключить/изменить данный пункт в соответствии с "
                f"{finding.legal_basis}. {finding.recommendation}"
            )
        return f"Требуется переформулировать с учётом {finding.legal_basis}. {finding.recommendation}"


# Синглтон
legal_doc_gen = LegalDocumentGenerator()
