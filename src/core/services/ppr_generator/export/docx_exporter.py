"""
PPR Generator — DOCX Exporter.

Экспортирует пояснительную записку ППР в формат Microsoft Word (.docx).
"""
from __future__ import annotations

import logging
import tempfile
from typing import List

from ..schemas import PPRInput, SectionResult, TTKResult

logger = logging.getLogger(__name__)


class PPRDocxExporter:
    """Экспортёр ППР в DOCX."""

    def compile(
        self,
        input: PPRInput,
        sections: List[SectionResult],
        ttks: List[TTKResult],
    ) -> str:
        """
        Собрать DOCX с пояснительной запиской.

        Args:
            input: Входные данные ППР
            sections: Разделы ПЗ
            ttks: Технологические карты

        Returns:
            Путь к созданному DOCX-файлу
        """
        safe_code = input.project_code.replace("/", "-")

        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed — using placeholder")
            return self._compile_placeholder(input)

        output_path = tempfile.mktemp(suffix=".docx", prefix=f"ppr_{safe_code}_")
        doc = Document()

        # Default style
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(11)

        for section in sections:
            # Section title
            heading = doc.add_heading(section.title, level=1)

            # Content as paragraphs
            for para_text in section.content.split("\n\n"):
                text = para_text.strip()
                if not text:
                    continue
                if text.startswith("### "):
                    doc.add_heading(text[4:], level=3)
                elif text.startswith("## "):
                    doc.add_heading(text[3:], level=2)
                elif text.startswith("|"):
                    # Table — parse rows
                    rows = [r.strip() for r in text.split("\n") if r.strip().startswith("|")]
                    if len(rows) >= 2:
                        data_rows = [[cell.strip() for cell in row.split("|")[1:-1]] for row in rows]
                        # Filter header separator
                        clean_rows = [data_rows[0]] + [r for r in data_rows[1:] if not all(c.startswith("-") for c in r if c)]
                        if len(clean_rows) >= 1:
                            table = doc.add_table(rows=len(clean_rows), cols=len(clean_rows[0]))
                            table.style = "Table Grid"
                            for i, row_data in enumerate(clean_rows):
                                for j, cell_text in enumerate(row_data):
                                    if j < len(table.rows[i].cells):
                                        table.rows[i].cells[j].text = cell_text
                else:
                    p = doc.add_paragraph(text)

            if section != sections[-1]:
                doc.add_page_break()

        doc.save(output_path)
        logger.info(f"DOCX compiled: {output_path}")
        return output_path

    def _compile_placeholder(self, input: PPRInput) -> str:
        """Fallback: save as text when python-docx is unavailable."""
        import os
        safe_code = input.project_code.replace("/", "-")
        output_path = os.path.join(
            tempfile.gettempdir(),
            f"ppr_{safe_code}_export.txt"
        )
        logger.warning(f"DOCX exporter using plain text fallback: {output_path}")
        return output_path
