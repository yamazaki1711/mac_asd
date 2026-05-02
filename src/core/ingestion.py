"""
ASD v12.0 — Document Ingestion Pipeline.

Превращает поток сканов/файлов в структурированные данные для агентов:

  PDF/Изображение → OCR → Текст
  Текст → DocumentClassifier → Тип документа (АОСР/ТТН/Сертификат/КС-2/...)
  Тип + Текст → EntityExtractor → Сущности (даты, количества, номера, партии)
  Сущности → GraphService → Узлы и рёбра в NetworkX

Pipeline:
  IngestionPipeline
    ├─ scan_folder(folder) → список документов
    ├─ classify(text) → DocumentType
    ├─ extract(text, doc_type) → dict сущностей
    └─ ingest_to_graph(entities) → заполнение NetworkX
"""

from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


# =============================================================================
# Document Types
# =============================================================================

class DocumentType(str, Enum):
    AOSR = "aosr"                       # Акт освидетельствования скрытых работ
    AOOK = "aook"                       # Акт освидетельствования ответственных конструкций
    KS2 = "ks2"                         # Акт о приёмке выполненных работ (КС-2)
    KS3 = "ks3"                         # Справка о стоимости (КС-3)
    VOR = "vor"                         # Ведомость объёмов работ
    CERTIFICATE = "certificate"         # Сертификат/паспорт качества
    TTN = "ttn"                         # Товарно-транспортная накладная
    UPD = "upd"                         # Универсальный передаточный документ
    CONTRACT = "contract"               # Договор / допсоглашение
    CLAIM = "claim"                     # Претензия
    LETTER = "letter"                   # Деловое письмо / уведомление
    EMAIL = "email"                     # Email-переписка
    PHOTO = "photo"                     # Фотоотчёт
    JOURNAL = "journal"                 # Журнал работ (ОЖР, ЖВК, ЖБР)
    EXECUTIVE_SCHEME = "executive_scheme"  # Исполнительная схема (ИГС/ИС)
    DRAWING = "drawing"                 # Чертёж (КМ, КЖ, АР)
    UNKNOWN = "unknown"                 # Нераспознанный документ


# =============================================================================
# Classification Keywords — что искать в тексте для определения типа
# =============================================================================

DOCUMENT_KEYWORDS: Dict[DocumentType, List[Dict[str, Any]]] = {
    DocumentType.AOSR: [
        {"keywords": ["акт освидетельствования скрытых работ", "аоср", "скрытых работ"],
         "weight": 10, "must_have": ["акт", "освидетельствования"]},
        {"keywords": ["акт освидетельствования выполненных работ", "аоср"],
         "weight": 5},
    ],
    DocumentType.AOOK: [
        {"keywords": ["акт освидетельствования ответственных конструкций", "аоок",
                      "ответственных конструкций"], "weight": 10,
         "must_have": ["ответственных", "конструкций"]},
    ],
    DocumentType.KS2: [
        {"keywords": ["кс-2", "kc-2", "акт о приёмке выполненных работ", "форма № кс-2",
                      "выполненных работ", "акт о приемке"],
         "weight": 10,
         "must_have": ["кс-2", "kc-2", "акт о приёмке выполненных работ", "акт о приемке выполненных работ"]},
        # Fallback без must_have для плохого OCR (Tesseract искажает заголовки КС-форм)
        {"keywords": ["акт о приёмке", "выполненных работ", "сметная стоимость",
                      "единичная расценка", "позиция по смете", "всего по акту"],
         "weight": 8},
    ],
    DocumentType.KS3: [
        {"keywords": ["кс-3", "kc-3", "справка о стоимости", "форма № кс-3",
                      "справка о стоимости выполненных работ"],
         "weight": 10,
         "must_have": ["кс-3", "kc-3", "справка о стоимости"]},
        # Fallback без must_have для плохого OCR
        {"keywords": ["справка о стоимости", "выполненных работ и затрат",
                      "справка", "с начала проведения работ", "с начала года"],
         "weight": 8},
    ],
    DocumentType.CERTIFICATE: [
        {"keywords": ["сертификат качества", "паспорт качества", "сертификат соответствия",
                      "декларация о соответствии", "свидетельство о качестве",
                      "паспорт", "сертификат"],
         "weight": 10},
        {"keywords": ["качества", "партия №", "плавка", "евраз", "северсталь"],
         "weight": 5},
    ],
    DocumentType.TTN: [
        {"keywords": ["товарно-транспортная накладная", "ттн", "транспортная накладная",
                      "форма № 1-т"], "weight": 10,
         "must_have": ["накладная"]},
    ],
    DocumentType.UPD: [
        {"keywords": ["универсальный передаточный документ", "упд", "счёт-фактура",
                      "форма № у-1"], "weight": 10},
    ],
    DocumentType.CONTRACT: [
        {"keywords": ["договор подряда", "договор", "контракт", "дополнительное соглашение",
                      "допсоглашение", "субподряд"],
         "weight": 10, "must_have": ["договор"]},
        {"keywords": ["договор", "контракт", "субподряд"],
         "weight": 5},
    ],
    DocumentType.CLAIM: [
        {"keywords": ["претензия", "досудебная претензия", "требование об уплате",
                      "исковое заявление", "арбитражный суд"], "weight": 10},
    ],
    DocumentType.JOURNAL: [
        {"keywords": ["журнал", "ожр", "жвк", "жбр", "общий журнал работ",
                      "входного контроля", "бетонных работ", "сварочных работ"],
         "weight": 8, "must_have": ["журнал"]},
    ],
    DocumentType.EXECUTIVE_SCHEME: [
        {"keywords": ["исполнительная схема", "игс", "геодезическая схема",
                      "схема результатов", "исполнительная геодезическая"],
         "weight": 10},
    ],
    DocumentType.EMAIL: [
        {"keywords": ["from:", "to:", "subject:", "re:", "fwd:", "@"],
         "weight": 12, "must_have": ["from:"]},
    ],
    DocumentType.VOR: [
        {"keywords": ["ведомость объёмов работ", "вор", "ведомость объёмов"],
         "weight": 10, "must_have": ["ведомость"]},
    ],
    DocumentType.LETTER: [
        {"keywords": ["уведомление", "исх. №", "письмо", "направляем вам",
                      "просим вас", "сообщаем"], "weight": 6},
    ],
    DocumentType.DRAWING: [
        {"keywords": ["чертёж", "км", "кж", "ар", "генплан", "разрез", "спецификация",
                      "формат а", "лист 1"], "weight": 6},
    ],
}

# =============================================================================
# Extraction Patterns — регулярки для извлечения сущностей
# =============================================================================

@dataclass
class ExtractionPattern:
    """Шаблон для извлечения сущности из текста."""
    field: str
    patterns: List[str]  # Regex patterns
    doc_types: List[DocumentType]  # Для каких типов документов применять
    required: bool = False
    transform: Optional[str] = None  # "date", "float", "int"


EXTRACTION_PATTERNS: List[ExtractionPattern] = [
    # ── Даты ──
    ExtractionPattern("date", [
        r"от\s+«?(\d{2}[./]\d{2}[./]\d{2,4})»?",
        r"(\d{2}[./]\d{2}[./]\d{2,4})\s*г\.",
        r"дата:\s*(\d{2}[./]\d{2}[./]\d{2,4})",
        r"дата составления\D*(\d{2}[./]\d{2}[./]\d{2,4})",
    ], list(DocumentType), transform="date"),

    ExtractionPattern("date_issue", [
        r"дата выдачи\D*(\d{2}[./]\d{2}[./]\d{2,4})",
        r"выдан\D+(\d{2}[./]\d{2}[./]\d{2,4})",
    ], [DocumentType.CERTIFICATE], transform="date"),

    # ── Номера документов ──
    ExtractionPattern("document_number", [
        r"№\s*([А-Яа-яA-Za-z0-9/-]+)",
        r"номер\D{0,5}№\s*([А-Яа-яA-Za-z0-9/-]+)",
        r"№\s*([\d]+[-/][\d]+[-/][\d]+)",
    ], list(DocumentType)),

    ExtractionPattern("aosr_number", [
        r"АОСР\s*№\s*([\d/-]+)",
        r"акт\s*№\s*([\d/-]+)",
    ], [DocumentType.AOSR, DocumentType.AOOK]),

    ExtractionPattern("contract_number", [
        r"договор\D+№\s*([\d/-]+[А-Яа-яA-Za-z]*)",
        r"контракт\D+№\s*([\d/-]+[А-Яа-яA-Za-z]*)",
    ], [DocumentType.CONTRACT]),

    # ── Материалы ──
    ExtractionPattern("material_name", [
        r"(шпунт\s+[ЛЛ]арсена\s+\S+)",
        r"(шпунт\s+Л\d[\s-]*У?М?)",
        r"(арматура\s+А\d+\S*)",
        r"(бетон\s+(?:кл|В)\d+\S*)",
        r"(сталь\s+\S+\d*)",
        r"материал:\s*(.+?)(?:[,;.]|$)",
    ], [DocumentType.CERTIFICATE, DocumentType.TTN, DocumentType.AOSR, DocumentType.VOR]),

    # ── Количества и единицы ──
    ExtractionPattern("quantity", [
        r"(?:количество|объём|кол-во|в количестве)\D*([\d\s,.]+)\s*(?:шт|тн?|м3|м²|м2|п\.?\s*м|кг|компл)",
        r"(\d+[\s,.]*\d*)\s*(?:штук|тонн|шпунтин)",
    ], [DocumentType.TTN, DocumentType.CERTIFICATE, DocumentType.AOSR, DocumentType.KS2],
        transform="float"),

    ExtractionPattern("unit", [
        r"(?:шт|тн?|м3|м³|м2|м²|п\.?\s*м|кг|компл|секц)",
    ], [DocumentType.CERTIFICATE, DocumentType.TTN, DocumentType.VOR]),

    # ── Партия ──
    ExtractionPattern("batch_number", [
        r"(?:парти[яи]|плавк[аи]|batch)\D+№?\s*([\d/А-ЯA-Z-]+)",
        r"(?:партия|плавка)\s+№?\s*([\d]+)",
        r"№\s*партии\s*[:\s]*([\d/А-ЯA-Z-]+)",
    ], [DocumentType.CERTIFICATE, DocumentType.TTN]),

    ExtractionPattern("batch_size", [
        r"Количество:\s*(\d+[\s,.]*\d*)\s*(?:шт|штук)",
        r"Партия\s+№\s*[\w/-]+\s*\n\s*Количество:\s*(\d+[\s,.]*\d*)\s*(?:шт|штук)",
        r"(?:(?:парти[яи]|в количестве)\D*?)(\d+[\s,.]*\d*)\s*(?:шт|штук|тн|тонн)",
        r"(\d+)\s*шпунтин",
    ], [DocumentType.CERTIFICATE, DocumentType.TTN], transform="float"),

    # ──Поставщик / Заказчик ──
    ExtractionPattern("supplier_name", [
        r"(?:поставщик|изготовитель|производитель):\s*(.+?)(?:[,;\n]|$)",
        r"(?:ООО|ЗАО|АО|ПАО)\s*«(.+?)»",
    ], [DocumentType.CERTIFICATE, DocumentType.TTN]),

    ExtractionPattern("customer_name", [
        r"(?:заказчик|генподрядчик):\s*(.+?)(?:[,;\n]|$)",
    ], [DocumentType.CONTRACT, DocumentType.KS2, DocumentType.AOSR]),

    # ── ГОСТ ──
    ExtractionPattern("gost", [
        r"(ГОСТ(?:\s*Р)?\s*[\d.]+[-–]\d+)",
        r"(ГОСТ\s+(?:Р\s+)?[\d.]+)",
        r"(ТУ\s+[\d-]+)",
    ], [DocumentType.CERTIFICATE, DocumentType.CONTRACT]),

    # ── Вид работ ──
    ExtractionPattern("work_type", [
        r"(?:вид работ|наименование работ):\s*(.+?)(?:[,;\n]|$)",
        r"(?:погружение|забивка|вибропогружение|устройство|монтаж|бетонирование|армирование)\s+(.+?)(?:[,;\n]|$)",
    ], [DocumentType.AOSR, DocumentType.VOR, DocumentType.KS2]),
]


# =============================================================================
# Extracted Document
# =============================================================================

@dataclass
class ExtractedDocument:
    """Документ, извлечённый из скана/файла."""
    file_path: Path
    doc_type: DocumentType
    classification_confidence: float
    raw_text: str = ""
    entities: Dict[str, Any] = field(default_factory=dict)
    errors: List[str] = field(default_factory=list)
    page_count: int = 1
    processed_at: str = field(default_factory=lambda: datetime.now().isoformat())

    # VLM-поля (для сканированных документов)
    vlm_classified: bool = False          # классифицирован через VLM, а не keyword
    scan_info: Optional[Dict[str, Any]] = None  # ScanInfo as dict
    embedded_refs: List[Dict[str, str]] = field(default_factory=list)
    # embedded_refs: [{"type": "certificate", "identifier": "№21514", ...}]


# =============================================================================
# Document Classifier
# =============================================================================

class DocumentClassifier:
    """
    Классификатор строительных документов.

    Использует keyword-based подход с весами для определения типа документа
    по текстовому содержимому. Не требует LLM.
    """

    def classify(self, text: str) -> Tuple[DocumentType, float]:
        """
        Определить тип документа по тексту.

        Для каждого doc_type пробует все rule_sets; confidence = лучший результат
        среди rule_sets (а не среднее по всем rule_sets).

        Args:
            text: Извлечённый текст документа

        Returns:
            (DocumentType, confidence: 0.0–1.0)
        """
        text_lower = text.lower()
        best_per_type: Dict[DocumentType, float] = {}

        for doc_type, rule_sets in DOCUMENT_KEYWORDS.items():
            best_confidence = 0.0

            for rule in rule_sets:
                # Проверка must_have
                if "must_have" in rule:
                    if not all(mh.lower() in text_lower for mh in rule["must_have"]):
                        continue

                # Подсчёт keyword matches для этого rule_set
                matched = 0
                total = 0
                for kw in rule["keywords"]:
                    total += rule["weight"]
                    if kw.lower() in text_lower:
                        matched += rule["weight"]

                if total > 0:
                    conf = matched / total
                    if conf > best_confidence:
                        best_confidence = conf

            if best_confidence > 0:
                best_per_type[doc_type] = best_confidence

        if not best_per_type:
            return (DocumentType.UNKNOWN, 0.0)

        # Наилучшее совпадение
        best_type = max(best_per_type, key=best_per_type.get)
        confidence = best_per_type[best_type]

        # Порог уверенности (снижен для keyword-based классификатора)
        if confidence < 0.15:
            return (DocumentType.UNKNOWN, confidence)

        return (best_type, min(confidence, 0.95))


# =============================================================================
# Entity Extractor
# =============================================================================

class EntityExtractor:
    """
    Извлечение сущностей из текста документа.

    Использует regex-шаблоны с привязкой к типу документа.
    Не требует LLM — работает быстро на больших объёмах.
    """

    def extract(self, text: str, doc_type: DocumentType) -> Dict[str, Any]:
        """
        Извлечь сущности из текста.

        Args:
            text: Извлечённый текст
            doc_type: Тип документа

        Returns:
            Словарь сущностей {field: value|list}
        """
        entities: Dict[str, Any] = {
            "doc_type": doc_type.value,
            "raw_matches": {},
        }

        patterns_to_apply = [
            p for p in EXTRACTION_PATTERNS
            if not p.doc_types or doc_type in p.doc_types or list(DocumentType) == p.doc_types
        ]

        for ext_pat in patterns_to_apply:
            matches = []
            for pattern in ext_pat.patterns:
                found = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
                for m in found:
                    val = m.strip() if isinstance(m, str) else m[0].strip()
                    if val and len(val) > 1:
                        # Transform
                        try:
                            if ext_pat.transform == "date":
                                val = self._normalize_date(val)
                            elif ext_pat.transform == "float":
                                val = float(re.sub(r'[\s,]', '', val.replace(',', '.')))
                            elif ext_pat.transform == "int":
                                val = int(re.sub(r'[\s,]', '', val))
                        except (ValueError, AttributeError):
                            pass
                        matches.append(val)

            if matches:
                # Дедикация: если одно значение — строка, если много — список
                entities[ext_pat.field] = matches[0] if len(matches) == 1 else matches[:5]

        return entities

    def _normalize_date(self, raw: str) -> str:
        """Нормализовать дату в ISO-формат (YYYY-MM-DD)."""
        raw = raw.strip()
        for fmt in ["%d.%m.%Y", "%d.%m.%y", "%d/%m/%Y", "%d/%m/%y", "%Y-%m-%d"]:
            try:
                dt = datetime.strptime(raw, fmt)
                return dt.strftime("%Y-%m-%d")
            except ValueError:
                continue
        # Если не смогли — возвращаем как есть (нестандартный формат)
        return raw


# =============================================================================
# OCR Engine
# =============================================================================

class OCREngine:
    """
    Извлечение текста из PDF и изображений.

    Использует:
      - PyMuPDF (fitz) для текстовых PDF — быстро, без OCR
      - rapidocr для изображений и сканов без текстового слоя
      - Tesseract как резервный вариант
    """

    def extract_text(self, file_path: Path) -> Tuple[str, int]:
        """
        Извлечь текст из файла.

        Returns:
            (text, page_count)
        """
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            return self._extract_pdf(file_path)
        elif suffix in ('.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.webp'):
            return self._extract_image(file_path)
        elif suffix in ('.txt', '.md'):
            text = file_path.read_text(encoding='utf-8', errors='replace')
            return (text, 1)
        elif suffix in ('.docx',):
            return self._extract_docx(file_path)
        elif suffix in ('.xlsx', '.xls'):
            return self._extract_excel(file_path)
        else:
            logger.warning("Unsupported file type: %s", suffix)
            return ("", 0)

    def _extract_pdf(self, file_path: Path) -> Tuple[str, int]:
        """Извлечь текст из PDF через PyMuPDF + rapidocr для сканов."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF not installed")
            return ("", 0)

        doc = fitz.open(str(file_path))
        full_text = []
        page_count = len(doc)

        for page_num in range(page_count):
            page = doc[page_num]
            text = page.get_text()

            # Если текстовый слой пустой — это скан, нужен OCR
            if not text.strip() or len(text.strip()) < 50:
                try:
                    text = self._ocr_page(page)
                except Exception as e:
                    logger.warning("OCR failed for page %d of %s: %s",
                                   page_num + 1, file_path.name, e)
                    text = ""

            full_text.append(text)

        doc.close()
        return ("\n".join(full_text), page_count)

    def _ocr_page(self, page) -> str:
        """OCR страницы. Tesseract (rus+eng) primary для кириллицы. RapidOCR GPU — fallback."""
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")

        # Primary: Tesseract (native Cyrillic support, slower but accurate)
        import tempfile, subprocess
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            f.write(img_bytes)
            tmp_path = f.name
        try:
            result = subprocess.run(
                ['tesseract', tmp_path, 'stdout', '-l', 'rus+eng', '--psm', '6'],
                capture_output=True, text=True, timeout=30
            )
            text = result.stdout.strip()
            if len(text) > 20:
                return text
        except Exception as e:
            logger.debug("Tesseract failed: %s", e)
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

        # Fallback: RapidOCR GPU (onnxruntime-gpu, CUDA/TensorRT)
        # NOTE: Chinese-optimized model misreads Cyrillic as Latin.
        # PaddleOCR v5 with native Cyrillic will replace this.
        try:
            from rapidocr import RapidOCR
            engine = RapidOCR()
            output = engine(img_bytes)
            if hasattr(output, 'txts') and output.txts:
                return "\n".join(str(t) for t in output.txts if t)
        except ImportError:
            pass
        except Exception as e:
            logger.debug("RapidOCR failed: %s", e)

        return ""

    def _extract_image(self, file_path: Path) -> Tuple[str, int]:
        """OCR изображения через rapidocr."""
        try:
            from rapidocr import RapidOCR
            engine = RapidOCR()
            result, _ = engine(str(file_path))
            if result:
                lines = [item[1] for item in result if item[1]]
                return ("\n".join(lines), 1)
        except ImportError:
            pass

        # Tesseract fallback
        import subprocess
        result = subprocess.run(
            ['tesseract', str(file_path), 'stdout', '-l', 'rus+eng', '--psm', '6'],
            capture_output=True, text=True, timeout=30
        )
        return (result.stdout, 1)

    def _extract_docx(self, file_path: Path) -> Tuple[str, int]:
        """Извлечь текст из DOCX."""
        try:
            from docx import Document
            doc = Document(str(file_path))
            text = "\n".join(p.text for p in doc.paragraphs)
            return (text, 1)
        except ImportError:
            logger.warning("python-docx not installed")
            return ("", 0)

    def _extract_excel(self, file_path: Path) -> Tuple[str, int]:
        """Извлечь текст из Excel."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            texts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                texts.append(f"--- Лист: {sheet_name} ---")
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip():
                        texts.append(row_text)
            return ("\n".join(texts), len(wb.sheetnames))
        except ImportError:
            logger.warning("openpyxl not installed")
            return ("", 0)


# =============================================================================
# Ingestion Pipeline — Orchestrator
# =============================================================================

class IngestionPipeline:
    """
    Сквозной конвейер приёма документов:

    1. Сканирование папки — сбор файлов
    2. OCR — извлечение текста
    3. Классификация — определение типа документа
    4. Извлечение сущностей — даты, номера, количества, партии
    5. Заполнение NetworkX-графа
    """

    def __init__(self, enable_vlm: bool = True):
        self.classifier = DocumentClassifier()
        self.extractor = EntityExtractor()
        self.ocr = OCREngine()
        self.documents: List[ExtractedDocument] = []
        self.stats: Dict[str, Any] = {}

        # VLM-интеграция
        self.enable_vlm = enable_vlm
        self._scan_detector = None
        self._vlm_classifier = None

    @property
    def scan_detector(self):
        if self._scan_detector is None:
            from src.core.scan_detector import scan_detector
            self._scan_detector = scan_detector
        return self._scan_detector

    @property
    def vlm_classifier(self):
        if self._vlm_classifier is None:
            from src.core.vlm_classifier import vlm_classifier
            self._vlm_classifier = vlm_classifier
        return self._vlm_classifier

    def scan_folder(
        self,
        folder: Path,
        recursive: bool = True,
        file_types: Optional[List[str]] = None,
    ) -> List[ExtractedDocument]:
        """
        Пакетная обработка папки с документами.

        Args:
            folder: путь к папке
            recursive: рекурсивно обходить подпапки
            file_types: список расширений (['pdf', 'jpg', 'png', 'docx', ...]).
                        По умолчанию: pdf, jpg, jpeg, png, tiff, tif, docx, xlsx, txt

        Returns:
            Список ExtractedDocument
        """
        if file_types is None:
            file_types = ['pdf', 'jpg', 'jpeg', 'png', 'tiff', 'tif', 'bmp', 'docx', 'xlsx', 'txt']

        folder = Path(folder)
        if not folder.exists():
            logger.error("Folder not found: %s", folder)
            return []

        # Сбор файлов
        files = []
        glob_pattern = "**/*" if recursive else "*"
        for file_path in sorted(folder.glob(glob_pattern)):
            if file_path.is_file() and file_path.suffix.lower().lstrip('.') in file_types:
                files.append(file_path)

        logger.info("Ingestion: found %d files in %s", len(files), folder)
        return self.process_files(files)

    def process_files(self, files: List[Path]) -> List[ExtractedDocument]:
        """
        Обработать список файлов: OCR → classify → extract.
        """
        self.documents = []
        type_counts: Dict[str, int] = {}
        errors_total = 0

        for file_path in files:
            try:
                doc = self.process_single(file_path)
                self.documents.append(doc)
                type_counts[doc.doc_type.value] = type_counts.get(doc.doc_type.value, 0) + 1
                errors_total += len(doc.errors)
            except Exception as e:
                logger.error("Failed to process %s: %s", file_path.name, e)
                self.documents.append(ExtractedDocument(
                    file_path=file_path,
                    doc_type=DocumentType.UNKNOWN,
                    classification_confidence=0.0,
                    errors=[str(e)],
                ))

        self.stats = {
            "total_files": len(files),
            "processed": len(self.documents),
            "type_counts": type_counts,
            "total_errors": errors_total,
            "completed_at": datetime.now().isoformat(),
        }

        logger.info("Ingestion complete: %s", self.stats)
        return self.documents

    def process_single(self, file_path: Path) -> ExtractedDocument:
        """
        Обработать один файл: OCR → classify → extract.
        Сканированные PDF направляются на VLM-классификацию.
        """
        # Шаг 1: OCR
        text, page_count = self.ocr.extract_text(file_path)

        # Шаг 1.5: Детекция скана (для статистики)
        scan_info = None
        vlm_classified = False
        embedded_refs = []

        if file_path.suffix.lower() == '.pdf':
            scan_info = self.scan_detector.detect(file_path, text)

        if not text.strip():
            return ExtractedDocument(
                file_path=file_path,
                doc_type=DocumentType.UNKNOWN,
                classification_confidence=0.0,
                raw_text=text,
                page_count=page_count,
                errors=["No text extracted"],
                scan_info={
                    "is_scanned": scan_info.is_scanned if scan_info else False,
                    "file_size_kb": scan_info.file_size_bytes // 1024 if scan_info else 0,
                } if scan_info else None,
            )

        # Шаг 2: Классификация (keyword, затем опциональный LLM fallback)
        doc_type, confidence = self.classifier.classify(text)

        try:
            import asyncio
            from src.core.hybrid_classifier import hybrid_classifier
            if hasattr(hybrid_classifier, 'classify'):
                result = asyncio.run(hybrid_classifier.classify(text, enable_llm=False))
                if result and result.confidence > confidence:
                    doc_type = DocumentType(result.doc_type) if result.doc_type in DocumentType.__members__ else doc_type
                    confidence = result.confidence
        except Exception:
            pass  # Keyword classification wins

        # VLM fallback: если keyword-классификатор не уверен — пробуем VLM
        # VLM fallback #1: keyword confidence < 0.5
        # VLM fallback #2: filename содержит КС-формы (кс2/кс3/кс6/кс-2/кс-3/кс-6, латиница тоже)
        #   КС-формы почти всегда сканированы с плохим OCR, keyword путает с contract
        vlm_trigger = (
            self.enable_vlm and file_path.suffix.lower() == '.pdf' and (
                confidence < 0.5 or
                any(kw in file_path.name.lower() for kw in
                    ['кс-2', 'кс-3', 'кс-6а', 'кс-6', 'кс2', 'кс3', 'кс6а', 'кс6',
                     'ks-2', 'ks-3', 'ks-6', 'ks2', 'ks3', 'ks6'])
            )
        )
        if vlm_trigger:
            logger.info("Low keyword confidence (%.2f) for %s — trying VLM",
                       confidence, file_path.name)
            try:
                import asyncio
                vlm_result = asyncio.run(
                    self.vlm_classifier.classify_document(file_path)
                )
                if vlm_result.doc_type and vlm_result.doc_type != "Неизвестно":
                    doc_type = self._map_vlm_type(vlm_result.doc_type)
                    confidence = vlm_result.confidence
                    vlm_classified = True
                    embedded_refs = vlm_result.embedded_refs
                    logger.info("VLM fallback: %s → %s (confidence: %.2f)",
                               file_path.name, doc_type.value, confidence)
            except Exception as e:
                logger.warning("VLM fallback failed for %s: %s", file_path.name, e)

        # Шаг 3: Извлечение сущностей
        entities = self.extractor.extract(text, doc_type)

        return ExtractedDocument(
            file_path=file_path,
            doc_type=doc_type,
            classification_confidence=confidence,
            raw_text=text[:2000],  # Первые 2000 символов — для отладки
            entities=entities,
            page_count=page_count,
            vlm_classified=vlm_classified,
            embedded_refs=embedded_refs,
            scan_info={
                "is_scanned": scan_info.is_scanned if scan_info else False,
                "file_size_kb": scan_info.file_size_bytes // 1024 if scan_info else 0,
                "text_chars": scan_info.text_chars if scan_info else len(text),
            } if scan_info else None,
        )

    def _map_vlm_type(self, vlm_type: str) -> DocumentType:
        """Преобразовать строку типа от VLM в DocumentType enum."""
        type_map = {
            "АОСР": DocumentType.AOSR,
            "АООК": DocumentType.AOOK,
            "КС-2": DocumentType.KS2,
            "КС-3": DocumentType.KS3,
            "КС-6а": DocumentType.KS2,  # КС-6а — журнал учёта, ближе к КС-2
            "Сертификат": DocumentType.CERTIFICATE,
            "Паспорт": DocumentType.CERTIFICATE,
            "Счёт": DocumentType.UPD,
            "УПД": DocumentType.UPD,
            "Договор": DocumentType.CONTRACT,
            "Протокол": DocumentType.CLAIM,
            "Журнал": DocumentType.JOURNAL,
            "Исполнительная схема": DocumentType.EXECUTIVE_SCHEME,
            "Чертёж": DocumentType.DRAWING,
            "Письмо": DocumentType.LETTER,
            "Неизвестно": DocumentType.UNKNOWN,
        }
        return type_map.get(vlm_type, DocumentType.UNKNOWN)

    def _first_str(self, value, default: str = "") -> str:
        """Извлечь первую строку из значения (может быть list или str)."""
        if isinstance(value, list):
            return str(value[0]) if value else default
        return str(value) if value else default

    def _first_float(self, value, default: float = 0.0) -> float:
        """Извлечь первое число из значения."""
        if isinstance(value, list):
            for v in value:
                if isinstance(v, (int, float)):
                    return float(v)
            return default
        return float(value) if value else default

    def ingest_to_graph(self, project_id: str = "") -> int:
        """
        Заполнить NetworkX-граф извлечёнными документами.

        Для каждого ExtractedDocument создаёт соответствующий узел
        (AOSR/Batch/Certificate/...) и рёбра связей.
        UNKNOWN-документы также попадают в граф как generic document nodes.

        Returns:
            Количество добавленных узлов
        """
        from src.core.graph_service import graph_service

        nodes_added = 0
        for doc in self.documents:
            doc_id = f"doc_{doc.file_path.stem}_{nodes_added}"
            entities = doc.entities

            try:
                if doc.doc_type == DocumentType.AOSR:
                    graph_service.add_aosr(
                        aosr_id=doc_id,
                        work_type=self._first_str(entities.get("work_type", "")),
                        description=f"Извлечено из: {doc.file_path.name}",
                        date=self._first_str(entities.get("date", "")),
                        project_id=project_id,
                    )
                    nodes_added += 1
                    mat_name = self._first_str(entities.get("material_name", ""))
                    if mat_name:
                        mat_id = f"mat_{mat_name.replace(' ', '_')}"
                        if not graph_service.graph.has_node(mat_id):
                            graph_service.add_material(mat_id, mat_name)
                            nodes_added += 1

                elif doc.doc_type == DocumentType.CERTIFICATE:
                    graph_service.add_certificate(
                        cert_id=doc_id,
                        material_name=self._first_str(entities.get("material_name", "")),
                        batch_number=self._first_str(entities.get("batch_number", "")),
                        batch_size=self._first_float(entities.get("batch_size", 0.0)),
                        unit=self._first_str(entities.get("unit", "")),
                        supplier=self._first_str(entities.get("supplier_name", "")),
                        issue_date=self._first_str(entities.get("date_issue", "")),
                        gost=self._first_str(entities.get("gost", "")),
                    )
                    nodes_added += 1
                    batch_num = self._first_str(entities.get("batch_number", ""))
                    if batch_num:
                        batch_id = f"batch_{batch_num}"
                        if graph_service.graph.has_node(batch_id):
                            graph_service.link_certificate_to_batch(doc_id, batch_id)

                elif doc.doc_type == DocumentType.TTN:
                    graph_service.add_ttn(
                        ttn_id=doc_id,
                        supplier=self._first_str(entities.get("supplier_name", "")),
                        date=self._first_str(entities.get("date", "")),
                        material_list=self._first_str(entities.get("material_name", "")),
                    )
                    nodes_added += 1

                elif doc.doc_type in (DocumentType.KS2, DocumentType.CONTRACT,
                                      DocumentType.CLAIM, DocumentType.VOR,
                                      DocumentType.EXECUTIVE_SCHEME):
                    graph_service.add_document(doc_id, {
                        "doc_type": doc.doc_type.value,
                        "file_name": doc.file_path.name,
                        **entities,
                    })
                    nodes_added += 1

                else:
                    # UNKNOWN and other types — still create a document node
                    graph_service.add_document(doc_id, {
                        "doc_type": doc.doc_type.value,
                        "file_name": doc.file_path.name,
                        "raw_text_sample": (doc.raw_text or "")[:500],
                        "confidence": doc.classification_confidence,
                    })
                    nodes_added += 1

                # Provenance: документ → файл (scan node must be created first)
                scan_id = f"scan_{doc.file_path.name}"
                if not graph_service.graph.has_node(scan_id):
                    graph_service.add_scan(scan_id, str(doc.file_path))
                    nodes_added += 1
                if graph_service.graph.has_node(doc_id):
                    graph_service.link_document_to_scan(doc_id, scan_id)

            except Exception as e:
                logger.error("Failed to ingest %s to graph: %s", doc_id, e)

        logger.info("Ingested %d nodes from %d documents", nodes_added, len(self.documents))
        return nodes_added

    def get_inventory_report(self) -> Dict[str, Any]:
        """
        Сформировать отчёт об инвентаризации:
        что найдено, в каком состоянии, чего не хватает.
        Включает данные VLM-классификации и встроенные ссылки.
        """
        doc_types_found = {}
        for doc in self.documents:
            dt = doc.doc_type.value
            doc_types_found[dt] = doc_types_found.get(dt, 0) + 1

        # Сбор VLM-данных
        vlm_classified_count = sum(1 for d in self.documents if d.vlm_classified)
        scanned_count = sum(
            1 for d in self.documents
            if d.scan_info and d.scan_info.get("is_scanned")
        )
        all_embedded_refs = []
        for d in self.documents:
            for ref in d.embedded_refs:
                all_embedded_refs.append({
                    **ref,
                    "found_in": str(d.file_path),
                })

        return {
            "total_processed": len(self.documents),
            "doc_types_found": doc_types_found,
            "confidence_distribution": {
                "high": sum(1 for d in self.documents if d.classification_confidence >= 0.7),
                "medium": sum(1 for d in self.documents if 0.3 <= d.classification_confidence < 0.7),
                "low": sum(1 for d in self.documents if d.classification_confidence < 0.3),
            },
            "unknown_docs": [
                str(d.file_path) for d in self.documents
                if d.doc_type == DocumentType.UNKNOWN
            ],
            "stats": self.stats,
            # VLM-данные
            "vlm_stats": {
                "vlm_classified": vlm_classified_count,
                "scanned_detected": scanned_count,
                "embedded_references": all_embedded_refs,
                "embedded_reference_count": len(all_embedded_refs),
            },
        }


# =============================================================================
# Singleton
# =============================================================================

ingestion_pipeline = IngestionPipeline()
