"""
ASD v12.0 — Output Pipeline (Генерация готовых документов).

Генерирует готовые к печати и сдаче стройконтролю документы:
  - АОСР (Акт освидетельствования скрытых работ)
  - КС-2 (Акт о приёмке выполненных работ)
  - КС-3 (Справка о стоимости)
  - Реестр ИД (опись + титульный лист)
  - Протокол разногласий
  - Досудебная претензия

Все документы — DOCX, готовые к печати или отправке.
Нумерация — сквозная, по проекту, с учётом системы нумерации заказчика.

Pipeline:
  OutputPipeline
    ├─ generate_aosr(data) → DOCX
    ├─ generate_ks2(data)  → DOCX
    ├─ generate_ks3(data)  → DOCX
    ├─ generate_id_register(project) → DOCX (реестр + опись)
    ├─ bundle_package(docs) → ZIP готового комплекта
    └─ NumberingService — сквозная нумерация
"""

from __future__ import annotations

import logging
import os
import shutil
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    Document = None  # type: ignore
    Pt = Cm = Inches = RGBColor = Emu = None  # type: ignore
    WD_ALIGN_PARAGRAPH = WD_TABLE_ALIGNMENT = WD_ORIENT = None  # type: ignore
    qn = None  # type: ignore

logger = logging.getLogger(__name__)


# =============================================================================
# Document Numbering Service
# =============================================================================

@dataclass
class DocumentNumber:
    """Номер документа в системе нумерации заказчика."""
    prefix: str       # "АОСР", "КС2", "КС3"
    project_code: str  # Код проекта у заказчика (напр. "СК-2025")
    sequence: int      # Порядковый номер
    suffix: str = ""   # Доп. суффикс (напр. "/1" для доработки)

    def __str__(self) -> str:
        base = f"{self.prefix}-{self.project_code}-{self.sequence:04d}"
        return base + self.suffix if self.suffix else base


class NumberingService:
    """
    Сквозная нумерация документов по проекту.

    Хранит состояние в ~/.hermes/asd_numbering.json.
    """

    _instance = None
    _state: Dict[str, Dict[str, int]] = {}  # {project_code: {prefix: last_seq}}

    def __init__(self, state_file: str = ""):
        self.state_file = Path(state_file or os.path.expanduser("~/.hermes/asd_numbering.json"))
        self._load()

    def _load(self):
        import json
        try:
            if self.state_file.exists():
                with open(self.state_file) as f:
                    self._state = json.load(f)
        except (OSError, json.JSONDecodeError, ValueError) as e:
            logger.debug("Failed to load output pipeline state: %s", e)
            self._state = {}

    def _save(self):
        import json
        self.state_file.parent.mkdir(parents=True, exist_ok=True)
        with open(self.state_file, "w") as f:
            json.dump(self._state, f, indent=2)

    def next_number(self, project_code: str, prefix: str) -> DocumentNumber:
        """Выдать следующий номер для префикса."""
        if project_code not in self._state:
            self._state[project_code] = {}
        if prefix not in self._state[project_code]:
            self._state[project_code][prefix] = 0
        self._state[project_code][prefix] += 1
        self._save()
        return DocumentNumber(
            prefix=prefix,
            project_code=project_code,
            sequence=self._state[project_code][prefix],
        )


# =============================================================================
# A4 Template Base
# =============================================================================

class A4Template:
    """Базовый класс для создания DOCX-документов формата А4."""

    FONT_MAIN = "Times New Roman"

    def __init__(self):
        if not _HAS_DOCX:
            raise ImportError(
                "python-docx is required for document generation. "
                "Install it with: pip install python-docx"
            )
        self.doc = Document()
        self._setup_page()

    @property
    def FONT_SIZE_BODY(self):
        return Pt(12)

    @property
    def FONT_SIZE_SMALL(self):
        return Pt(10)

    @property
    def FONT_SIZE_TITLE(self):
        return Pt(14)

    @property
    def MARGIN_TOP(self):
        return Cm(2)

    @property
    def MARGIN_BOTTOM(self):
        return Cm(2)

    @property
    def MARGIN_LEFT(self):
        return Cm(3)    # 3 см слева — под подшивку

    @property
    def MARGIN_RIGHT(self):
        return Cm(1.5)

    def _setup_page(self):
        """Настройка страницы А4."""
        for section in self.doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = self.MARGIN_TOP
            section.bottom_margin = self.MARGIN_BOTTOM
            section.left_margin = self.MARGIN_LEFT
            section.right_margin = self.MARGIN_RIGHT

    def _add_paragraph(self, text: str, bold: bool = False, alignment: int = WD_ALIGN_PARAGRAPH.LEFT,
                       font_size: Pt = None, space_after: Pt = Pt(6)) -> "Paragraph":
        """Добавить параграф с форматированием."""
        p = self.doc.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_after = space_after
        run = p.add_run(text)
        run.font.name = self.FONT_MAIN
        run.font.size = font_size or self.FONT_SIZE_BODY
        run.bold = bold
        # Для кириллицы нужно явно задать шрифт для run
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_MAIN)
        return p

    def _add_table(self, headers: List[str], rows: List[List[str]],
                   col_widths: List[Cm] = None) -> "Table":
        """Добавить таблицу."""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заголовки
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = self.FONT_SIZE_SMALL
                    run.font.name = self.FONT_MAIN

        # Данные
        for row_idx, row_data in enumerate(rows):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = self.FONT_SIZE_SMALL
                        run.font.name = self.FONT_MAIN

        # Ширины колонок
        if col_widths:
            for row in table.rows:
                for i, width in enumerate(col_widths):
                    row.cells[i].width = width

        return table

    def save(self, path: Path) -> Path:
        """Сохранить документ."""
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        logger.info("Document saved: %s", path)
        return path


# =============================================================================
# AOSR Generator
# =============================================================================

class AOSRGenerator(A4Template):
    """
    Генератор Акта освидетельствования скрытых работ (АОСР).

    Форма по Приказу 344/пр (Приложение 3).
    """

    def generate(self, data: Dict[str, Any]) -> Path:
        """
        Сгенерировать АОСР.

        data:
          - aosr_number: str (напр. "АОСР-СК-2025-0001")
          - project_name: str
          - object_address: str
          - work_type: str
          - work_start: str
          - work_end: str
          - materials: list[str]
          - certificates: list[str]
          - design_docs: list[str]
          - executor_company: str ("КСК №1")
          - customer_company: str ("ГенПодряд")
          - developer_company: str ("Заказчик")
          - commission_members: list[{name, role, company}]
          - decision: str ("разрешается" / "запрещается")
          - date: str (дата подписания)
        """
        number = data.get("aosr_number", "АОСР-???-0000")

        # ── Заголовок ──
        self._add_paragraph("Приложение 3", font_size=Pt(10),
                            alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        self._add_paragraph(
            f"к приказу Минстроя России\nот 16.05.2023 № 344/пр",
            font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )

        self.doc.add_paragraph()  # отступ

        self._add_paragraph(
            f"АКТ ОСВИДЕТЕЛЬСТВОВАНИЯ СКРЫТЫХ РАБОТ\n{number}",
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=self.FONT_SIZE_TITLE
        )

        self.doc.add_paragraph()

        # ── Реквизиты ──
        self._add_paragraph(f"Объект капитального строительства: {data.get('project_name', '')}")
        self._add_paragraph(f"Адрес объекта: {data.get('object_address', '')}")

        # ── Комиссия ──
        self._add_paragraph("\nКомиссия в составе:", bold=True)
        members = data.get("commission_members", [])
        for i, m in enumerate(members, 1):
            self._add_paragraph(
                f"  {i}. {m.get('name', '')} — {m.get('role', '')} "
                f"({m.get('company', '')})"
            )

        # ── Освидетельствование ──
        self._add_paragraph("\nПроизвела освидетельствование скрытых работ:", bold=True)
        self._add_paragraph(f"Наименование работ: {data.get('work_type', '')}")
        self._add_paragraph(f"Дата начала работ: {data.get('work_start', '')}")
        self._add_paragraph(f"Дата окончания работ: {data.get('work_end', '')}")

        # ── Материалы ──
        materials = data.get("materials", [])
        if materials:
            self._add_paragraph("\nПри выполнении работ применены:", bold=True)
            for mat in materials:
                self._add_paragraph(f"  • {mat}")

        # ── Документы ──
        certs = data.get("certificates", [])
        if certs:
            self._add_paragraph("\nДокументы, подтверждающие качество:", bold=True)
            for c in certs:
                self._add_paragraph(f"  • {c}")

        design_docs = data.get("design_docs", [])
        if design_docs:
            self._add_paragraph("\nПроектная документация:", bold=True)
            for d in design_docs:
                self._add_paragraph(f"  • {d}")

        # ── Решение ──
        self.doc.add_paragraph()
        self._add_paragraph("Решение комиссии:", bold=True)
        decision = data.get("decision", "разрешается")
        self._add_paragraph(
            f"Выполнение последующих работ {decision}.",
            bold=(decision == "запрещается")
        )

        # ── Подписи ──
        self.doc.add_paragraph()
        self._add_paragraph(f"Дата составления акта: {data.get('date', '__________')}")

        self.doc.add_paragraph()
        for m in members:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(f"\n{m.get('name', '')}  __________________  (подпись)")
            run.font.name = self.FONT_MAIN
            run.font.size = self.FONT_SIZE_BODY

        # ── Сохранение ──
        output_dir = data.get("output_dir", "/tmp/asd_output")
        output_path = Path(output_dir) / f"{number}.docx"
        return self.save(output_path)


# =============================================================================
# KS-2 Generator
# =============================================================================

class KS2Generator(A4Template):
    """
    Генератор Акта о приёмке выполненных работ (КС-2).

    Унифицированная форма по Постановлению Госкомстата №100.
    Ориентация — альбомная для широкой таблицы.
    """

    def generate(self, data: Dict[str, Any]) -> Path:
        """
        Сгенерировать КС-2.

        data:
          - ks2_number: str
          - ks3_number: str (связанная справка КС-3)
          - project_name: str
          - object_address: str
          - investor: str
          - customer: str ("Генподрядчик")
          - contractor: str ("КСК №1")
          - contract_number: str
          - contract_date: str
          - period: str ("01.04.2025 – 30.04.2025")
          - date: str (дата составления)
          - lines: list[{row, code, name, unit, quantity, unit_price, total}]
          - overhead_pct: float
          - profit_pct: float
          - vat_pct: float (20.0)
          - output_dir: str
        """
        # Альбомная ориентация для КС-2
        for section in self.doc.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(29.7)
            section.page_height = Cm(21)

        number = data.get("ks2_number", "КС2-???-0000")

        # ── Шапка ──
        self._add_paragraph(
            f"Унифицированная форма № КС-2\n"
            f"Утверждена постановлением Госкомстата России от 11.11.99 № 100",
            font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )

        # Таблица реквизитов (2 колонки)
        self._add_paragraph(
            f"АКТ О ПРИЁМКЕ ВЫПОЛНЕННЫХ РАБОТ {number}",
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=self.FONT_SIZE_TITLE
        )

        # ── Основная таблица КС-2 ──
        lines = data.get("lines", [])
        headers = ["№ п/п", "Шифр\nрасценки", "Наименование работ", "Ед.\nизм.",
                    "Кол-во", "Цена за ед.\n(руб.)", "Стоимость\n(руб.)"]
        rows = []
        for line in lines:
            rows.append([
                str(line.get("row", "")),
                str(line.get("code", "")),
                str(line.get("name", "")),
                str(line.get("unit", "")),
                str(line.get("quantity", "")),
                f'{line.get("unit_price", 0):.2f}',
                f'{line.get("total", 0):.2f}',
            ])

        # Итого
        subtotal = data.get("subtotal", sum(l.get("total", 0) for l in lines))
        rows.append(["", "", "ИТОГО прямые затраты", "", "", "", f"{subtotal:.2f}"])
        rows.append(["", "", f"Накладные расходы ({data.get('overhead_pct', 0)}%)",
                      "", "", "", f"{data.get('overhead', 0):.2f}"])
        rows.append(["", "", f"Сметная прибыль ({data.get('profit_pct', 0)}%)",
                      "", "", "", f"{data.get('profit', 0):.2f}"])
        rows.append(["", "", f"НДС {data.get('vat_pct', 20)}%",
                      "", "", "", f"{data.get('vat', 0):.2f}"])
        rows.append(["", "", "ВСЕГО по акту", "", "", "",
                      f"{data.get('grand_total', 0):.2f}"])

        col_widths = [Cm(1.2), Cm(2.5), Cm(8), Cm(1.5), Cm(1.5), Cm(2.5), Cm(2.5)]
        self._add_table(headers, rows, col_widths)

        # ── Подписи ──
        self.doc.add_paragraph()
        self._add_paragraph(f"Сдал: {data.get('contractor', 'ООО «КСК №1»')}  _______________  __________")
        self._add_paragraph(f"Принял: {data.get('customer', 'Заказчик')}  _______________  __________")
        self._add_paragraph(f"Дата: {data.get('date', '__________')}")

        output_path = Path(data.get("output_dir", "/tmp/asd_output")) / f"{number}.docx"
        return self.save(output_path)


# =============================================================================
# KS-3 Generator
# =============================================================================

class KS3Generator(A4Template):
    """Генератор Справки о стоимости выполненных работ (КС-3)."""

    def generate(self, data: Dict[str, Any]) -> Path:
        """Сгенерировать КС-3."""
        number = data.get("ks3_number", "КС3-???-0000")

        self._add_paragraph(
            "Унифицированная форма № КС-3",
            font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )

        self._add_paragraph(
            f"СПРАВКА О СТОИМОСТИ ВЫПОЛНЕННЫХ РАБОТ И ЗАТРАТ\n{number}",
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=self.FONT_SIZE_TITLE
        )
        self.doc.add_paragraph()

        self._add_paragraph(f"Заказчик: {data.get('customer', '')}")
        self._add_paragraph(f"Подрядчик: {data.get('contractor', '')}")
        self._add_paragraph(f"Стройка: {data.get('project_name', '')}")
        self._add_paragraph(f"Договор подряда: {data.get('contract_number', '')} "
                            f"от {data.get('contract_date', '')}")
        self._add_paragraph(f"Отчётный период: {data.get('period', '')}")

        self.doc.add_paragraph()

        # Таблица
        grand_total = data.get("grand_total", 0)
        headers = ["№", "Показатель", "С начала работ", "За отчётный период"]
        rows = [
            ["1", "Стоимость выполненных работ и затрат",
             f"{data.get('total_since_start', grand_total):.2f}",
             f"{grand_total:.2f}"],
            ["2", f"в т.ч. НДС {data.get('vat_pct', 20)}%",
             f"{data.get('vat_total', grand_total * 0.1667):.2f}",
             f"{data.get('vat_period', grand_total * 0.1667):.2f}"],
            ["3", "Всего с НДС",
             f"{data.get('total_since_start', grand_total):.2f}",
             f"{grand_total:.2f}"],
        ]
        self._add_table(headers, rows, [Cm(1), Cm(8), Cm(4), Cm(4)])

        # Подписи
        self.doc.add_paragraph()
        self._add_paragraph(f"Сдал: {data.get('contractor', '')}  _______________  __________")
        self._add_paragraph(f"Принял: {data.get('customer', '')}  _______________  __________")
        self._add_paragraph(f"Дата: {data.get('date', '')}")

        output_path = Path(data.get("output_dir", "/tmp/asd_output")) / f"{number}.docx"
        return self.save(output_path)


# =============================================================================
# ID Register Generator (Реестр ИД)
# =============================================================================

class IDRegisterGenerator(A4Template):
    """
    Генератор Реестра исполнительной документации.

    Формирует:
      - Титульный лист
      - Реестр (таблица всех документов)
      - Опись передаваемых документов
    """

    def generate(self, project_data: Dict[str, Any]) -> Path:
        """
        Сгенерировать реестр ИД.

        project_data:
          - project_name: str
          - project_code: str
          - customer: str
          - contractor: str
          - date: str
          - documents: list[{number, name, pages, date, status, note}]
          - output_dir: str
        """
        code = project_data.get("project_code", "???")

        # ── Титульный лист ──
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self._add_paragraph(
            "РЕЕСТР ИСПОЛНИТЕЛЬНОЙ ДОКУМЕНТАЦИИ",
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(16)
        )
        self.doc.add_paragraph()

        self._add_paragraph(
            f"Объект: {project_data.get('project_name', '')}",
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
        self._add_paragraph(
            f"Заказчик: {project_data.get('customer', '')}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
        self._add_paragraph(
            f"Подрядчик: {project_data.get('contractor', 'ООО «КСК №1»')}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
        self._add_paragraph(
            f"Дата составления: {project_data.get('date', '__________')}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )

        self.doc.add_page_break()

        # ── Таблица реестра ──
        self._add_paragraph("РЕЕСТР ДОКУМЕНТОВ", bold=True, font_size=Pt(14))
        self.doc.add_paragraph()

        docs = project_data.get("documents", [])
        headers = ["№ п/п", "Номер документа", "Наименование", "Стр.", "Дата", "Статус", "Примечание"]
        rows = []
        for i, doc in enumerate(docs, 1):
            rows.append([
                str(i),
                doc.get("number", ""),
                doc.get("name", doc.get("number", "")),
                str(doc.get("pages", "")),
                doc.get("date", ""),
                doc.get("status", ""),
                doc.get("note", ""),
            ])

        self._add_table(headers, rows,
                        [Cm(1), Cm(3.5), Cm(6), Cm(1), Cm(2.5), Cm(2), Cm(3)])

        self.doc.add_paragraph()
        self._add_paragraph(f"Всего документов: {len(docs)}")
        self._add_paragraph(f"Всего листов: {sum(d.get('pages', 0) for d in docs)}")

        # ── Опись ──
        self.doc.add_page_break()
        self._add_paragraph(
            f"ОПИСЬ ДОКУМЕНТОВ, ПЕРЕДАВАЕМЫХ ЗАКАЗЧИКУ\n"
            f"по объекту: {project_data.get('project_name', '')}",
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
        self.doc.add_paragraph()

        for i, doc in enumerate(docs, 1):
            self._add_paragraph(f"{i}. {doc.get('number', '')} — {doc.get('name', '')} "
                                f"({doc.get('pages', '')} л.)")

        self.doc.add_paragraph()
        self._add_paragraph("Документацию сдал:")
        self._add_paragraph(f"{project_data.get('contractor', 'ООО «КСК №1»')}  "
                            f"_______________  /_________/")
        self._add_paragraph("\nДокументацию принял:")
        self._add_paragraph(f"{project_data.get('customer', '')}  "
                            f"_______________  /_________/")
        self._add_paragraph(f"\nДата: {project_data.get('date', '__________')}")

        output_path = Path(project_data.get("output_dir", "/tmp/asd_output")) / \
                      f"Реестр_ИД_{code}.docx"
        return self.save(output_path)


# =============================================================================
# Output Pipeline — Orchestrator
# =============================================================================

class OutputPipeline:
    """
    Сквозной конвейер генерации выходных документов.

    Принимает данные от агентов → генерирует готовые DOCX → пакует в ZIP.
    """

    def __init__(self):
        self.numbering = NumberingService()
        self.aosr_gen = AOSRGenerator()
        self.ks2_gen = KS2Generator()
        self.ks3_gen = KS3Generator()
        self.register_gen = IDRegisterGenerator()

    def generate_aosr_package(
        self, project_code: str, aosr_data: Dict[str, Any]
    ) -> Path:
        """Сгенерировать один АОСР."""
        num = self.numbering.next_number(project_code, "АОСР")
        aosr_data["aosr_number"] = str(num)
        aosr_data["output_dir"] = aosr_data.get("output_dir", f"/tmp/asd_output/{project_code}")
        return self.aosr_gen.generate(aosr_data)

    def generate_ks_package(
        self, project_code: str, estimate_data: Dict[str, Any]
    ) -> Tuple[Path, Path]:
        """Сгенерировать пару КС-2 + КС-3."""
        num_ks2 = self.numbering.next_number(project_code, "КС2")
        num_ks3 = self.numbering.next_number(project_code, "КС3")

        estimate_data["ks2_number"] = str(num_ks2)
        estimate_data["ks3_number"] = str(num_ks3)
        estimate_data["output_dir"] = estimate_data.get("output_dir", f"/tmp/asd_output/{project_code}")

        ks2_path = self.ks2_gen.generate(estimate_data)
        ks3_path = self.ks3_gen.generate(estimate_data)
        return (ks2_path, ks3_path)

    def generate_id_register(self, project_data: Dict[str, Any]) -> Path:
        """Сгенерировать реестр ИД."""
        return self.register_gen.generate(project_data)

    def bundle_package(
        self, project_code: str, doc_paths: List[Path], output_dir: str = ""
    ) -> Path:
        """
        Упаковать все документы в ZIP-архив для передачи заказчику.
        """
        output_dir = Path(output_dir or f"/tmp/asd_output/{project_code}")
        output_dir.mkdir(parents=True, exist_ok=True)

        zip_name = f"ИД_{project_code}_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
        zip_path = output_dir / zip_name

        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for doc_path in doc_paths:
                if doc_path.exists():
                    zf.write(doc_path, doc_path.name)

        logger.info("Package bundled: %s (%d files)", zip_path, len(doc_paths))
        return zip_path

    # -------------------------------------------------------------------------
    # Auto-export from DeloAgent
    # -------------------------------------------------------------------------

    def generate_from_registry(self, registry_data: Dict[str, Any]) -> List[Path]:
        """
        Генерирует полный комплект ИД из данных реестра (DeloAgent).

        Автоматически создаёт:
          - Реестр ИД (титульный лист + опись)
          - Все АОСР из реестра
          - КС-2/КС-3 (если есть данные)

        Args:
            registry_data: результат DeloAgent.export_registry_for_output()

        Returns:
            Список путей к сгенерированным файлам
        """
        paths: List[Path] = []
        project_code = registry_data.get("project_code", "UNKNOWN")
        output_dir = registry_data.get("output_dir", f"/tmp/asd_output/{project_code}")

        # 1. Реестр ИД
        register_path = self.generate_id_register(registry_data)
        paths.append(register_path)

        # 2. АОСР из реестра
        for doc in registry_data.get("documents", []):
            if "аоср" in doc.get("name", "").lower():
                aosr_data = {
                    "aosr_number": doc["number"],
                    "project_name": registry_data.get("project_name", ""),
                    "object_address": registry_data.get("object_address", ""),
                    "work_type": doc.get("name", ""),
                    "decision": "разрешается",
                    "date": doc.get("date", ""),
                    "executor_company": registry_data.get("contractor", "ООО «КСК №1»"),
                    "customer_company": registry_data.get("customer", "Заказчик"),
                    "commission_members": [
                        {"name": "", "role": "Представитель заказчика",
                         "company": registry_data.get("customer", "")},
                        {"name": "", "role": "Представитель подрядчика",
                         "company": registry_data.get("contractor", "")},
                    ],
                    "output_dir": output_dir,
                }
                path = self.aosr_gen.generate(aosr_data)
                paths.append(path)

        # 3. КС-2/КС-3 если есть данные
        ks2_docs = [d for d in registry_data.get("documents", [])
                    if "кс-2" in d.get("name", "").lower() or "кс2" in d.get("note", "").lower()]
        if ks2_docs:
            ks2_data = {
                "ks2_number": self.numbering.next_number(project_code, "КС2"),
                "ks3_number": self.numbering.next_number(project_code, "КС3"),
                "project_name": registry_data.get("project_name", ""),
                "customer": registry_data.get("customer", "Заказчик"),
                "contractor": registry_data.get("contractor", "ООО «КСК №1»"),
                "date": datetime.now().strftime("%d.%m.%Y"),
                "lines": [],
                "output_dir": output_dir,
            }
            ks2_path = self.ks2_gen.generate(ks2_data)
            ks3_path = self.ks3_gen.generate(ks2_data)
            paths.extend([ks2_path, ks3_path])

        logger.info("Auto-generated %d files from registry for %s", len(paths), project_code)
        return paths


# =============================================================================
# Singleton
# =============================================================================

output_pipeline = OutputPipeline()
