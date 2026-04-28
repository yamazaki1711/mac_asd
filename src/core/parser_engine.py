"""
ASD v12.0 — Streaming Parser Engine.

Потоковый парсер для инвентаризации документов.
Поддерживает батчевую обработку с чекпоинтами (checkpoint/resume),
автоопределение типа файла, извлечение метаданных и VLM OCR.

Pipeline:
  file → detect_type → parse → chunk → metadata_extract → yield chunks

Поддерживаемые форматы:
  - PDF: текст (PyMuPDF) + Vision OCR fallback для сканов
  - DOCX: текст + таблицы
  - XLSX: данные таблиц (сметы, ВОР)
  - DWG: через ODA File Converter (если доступен)
  - Изображения: PNG/JPG/TIFF → Vision OCR

Checkpoint/Resume:
  - После каждых N страниц сохраняется прогресс
  - При сбое — возобновление с последнего чекпоинта

Usage:
  from src.core.parser_engine import StreamingParser

  parser = StreamingParser(checkpoint_dir="/tmp/asd_checkpoints")
  async for chunk in parser.parse_stream("документ.pdf", project_id=1):
      print(f"Page {chunk.page}: {chunk.content[:100]}...")

  # Batch mode:
  results = await parser.parse_batch(
      ["file1.pdf", "file2.docx"], 
      project_id=1,
      on_progress=lambda done, total: print(f"{done}/{total}")
  )
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import time
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, AsyncIterator, Callable, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


# =============================================================================
# Enums & Data Classes
# =============================================================================

class FileType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    DWG = "dwg"
    IMAGE = "image"
    UNKNOWN = "unknown"


class ParseMethod(str, Enum):
    PYMFIT_TEXT = "pymupdf_text"
    PYTHON_DOCX = "python_docx"
    OPENPYXL = "openpyxl"
    ODA_CONVERTER = "oda_converter"
    VISION_OCR = "vision_ocr"
    RAW_BINARY = "raw_binary"


@dataclass
class ParsedChunk:
    """Один чанк (обычно страница) распарсенного документа."""
    content: str
    page: int = 1
    total_pages: int = 1
    method: str = ParseMethod.PYMFIT_TEXT.value
    metadata: Dict[str, Any] = field(default_factory=dict)
    confidence: float = 1.0  # 0.0–1.0
    ocr_fallback: bool = False


@dataclass
class ParseResult:
    """Итог парсинга одного документа."""
    file_path: str
    file_type: FileType
    chunks: List[ParsedChunk]
    total_pages: int
    total_chars: int
    methods_used: List[str]
    duration_sec: float
    errors: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    @property
    def success(self) -> bool:
        return len(self.errors) == 0

    @property
    def full_text(self) -> str:
        """Склеить все чанки в один текст (с разделителями страниц)."""
        return "\n\n--- PAGE BREAK ---\n\n".join(
            c.content for c in self.chunks
        )


@dataclass
class Checkpoint:
    """Точка сохранения прогресса батчевого парсинга."""
    batch_id: str
    file_index: int
    page_number: int
    processed_files: List[str]
    total_files: int
    created_at: str


# =============================================================================
# File Type Detection
# =============================================================================

# Сигнатуры файлов (magic bytes)
MAGIC_SIGNATURES: Dict[bytes, FileType] = {
    b"%PDF": FileType.PDF,
    b"\x89PNG": FileType.IMAGE,
    b"\xff\xd8\xff": FileType.IMAGE,
    b"II*\x00": FileType.IMAGE,   # TIFF little-endian
    b"MM\x00*": FileType.IMAGE,    # TIFF big-endian
}

# Сигнатура DWG (первые 6 байт: "AC" + версия)
DWG_MAGIC_PREFIXES = [b"AC10", b"AC101", b"AC102", b"AC103"]


def detect_file_type(file_path: str) -> FileType:
    """
    Определить тип файла по расширению и magic bytes.

    Приоритет: magic bytes → расширение.
    """
    ext = Path(file_path).suffix.lower()

    # Расширение — первая линия
    ext_map = {
        ".pdf": FileType.PDF,
        ".docx": FileType.DOCX,
        ".xlsx": FileType.XLSX,
        ".xls": FileType.XLSX,
        ".dwg": FileType.DWG,
        ".png": FileType.IMAGE,
        ".jpg": FileType.IMAGE,
        ".jpeg": FileType.IMAGE,
        ".tiff": FileType.IMAGE,
        ".tif": FileType.IMAGE,
        ".bmp": FileType.IMAGE,
    }

    if ext in ext_map:
        return ext_map[ext]

    # Magic bytes — вторая линия
    try:
        with open(file_path, "rb") as f:
            header = f.read(8)
    except (IOError, OSError):
        return FileType.UNKNOWN

    for magic, ftype in MAGIC_SIGNATURES.items():
        if header.startswith(magic):
            return ftype

    # DWG проверка
    for prefix in DWG_MAGIC_PREFIXES:
        if header.startswith(prefix):
            return FileType.DWG

    return FileType.UNKNOWN


# =============================================================================
# Streaming Parser
# =============================================================================

class StreamingParser:
    """
    Потоковый парсер документов ASD v12.0.

    Поддерживает:
      - Поточный (streaming) и батчевый (batch) режимы
      - Автоопределение типа файла
      - Чекпоинты для возобновления после сбоя
      - Vision OCR fallback для сканированных PDF
      - Извлечение метаданных (автор, даты, кол-во страниц)
    """

    # Максимальный размер страницы текста для чанка (символов)
    MAX_CHUNK_CHARS = 10000

    # Размер батча для чекпоинтов (каждые N страниц)
    CHECKPOINT_INTERVAL_PAGES = 20

    def __init__(self, checkpoint_dir: str = "/tmp/asd_checkpoints"):
        self._checkpoint_dir = Path(checkpoint_dir)
        self._checkpoint_dir.mkdir(parents=True, exist_ok=True)
        self._active_checkpoints: Dict[str, Checkpoint] = {}

    # -------------------------------------------------------------------------
    # Streaming Mode
    # -------------------------------------------------------------------------

    async def parse_stream(
        self,
        file_path: str,
        project_id: int = 0,
        max_chars: int = 0,
        use_ocr: bool = True,
    ) -> AsyncIterator[ParsedChunk]:
        """
        Поточно парсить документ, отдавая чанки по мере готовности.

        Args:
            file_path: путь к файлу
            project_id: ID проекта (для метаданных)
            max_chars: лимит символов (0 = без лимита)
            use_ocr: использовать Vision OCR для сканов

        Yields:
            ParsedChunk по одному на страницу
        """
        file_type = detect_file_type(file_path)
        logger.info(
            "Streaming parse: %s [%s]", os.path.basename(file_path), file_type.value
        )

        start_time = time.time()
        total_chars = 0

        try:
            if file_type == FileType.PDF:
                async for chunk in self._parse_pdf_stream(file_path, use_ocr):
                    if max_chars and total_chars >= max_chars:
                        break
                    total_chars += len(chunk.content)
                    yield chunk
                    if max_chars and total_chars >= max_chars:
                        break

            elif file_type == FileType.DOCX:
                async for chunk in self._parse_docx_stream(file_path):
                    if max_chars and total_chars >= max_chars:
                        break
                    total_chars += len(chunk.content)
                    yield chunk

            elif file_type == FileType.XLSX:
                chunks = self._parse_xlsx(file_path)
                for chunk in chunks:
                    if max_chars and total_chars >= max_chars:
                        break
                    total_chars += len(chunk.content)
                    yield chunk

            elif file_type == FileType.IMAGE:
                chunk = await self._parse_image_ocr(file_path)
                yield chunk

            elif file_type == FileType.DWG:
                chunk = self._parse_dwg(file_path)
                yield chunk

            else:
                logger.warning("Unknown file type: %s", file_path)
                yield ParsedChunk(
                    content=f"[UNSUPPORTED FILE TYPE: {file_type.value}]",
                    method=ParseMethod.RAW_BINARY.value,
                    confidence=0.0,
                    metadata={"file_type": file_type.value},
                )

        except Exception as e:
            logger.error("Streaming parse error for %s: %s", file_path, e)
            yield ParsedChunk(
                content=f"[PARSE ERROR: {str(e)}]",
                method=ParseMethod.RAW_BINARY.value,
                confidence=0.0,
                metadata={"error": str(e)},
            )

        elapsed = time.time() - start_time
        logger.info(
            "Streaming parse done: %s — %d chars in %.1fs",
            os.path.basename(file_path), total_chars, elapsed,
        )

    # -------------------------------------------------------------------------
    # Batch Mode
    # -------------------------------------------------------------------------

    async def parse_batch(
        self,
        file_paths: List[str],
        project_id: int = 0,
        on_progress: Optional[Callable[[int, int], None]] = None,
        resume: bool = True,
    ) -> List[ParseResult]:
        """
        Батчевый парсинг с чекпоинтами.

        Args:
            file_paths: список путей к файлам
            project_id: ID проекта
            on_progress: колбэк (done, total)
            resume: возобновить с последнего чекпоинта

        Returns:
            Список ParseResult для каждого файла
        """
        total = len(file_paths)
        batch_id = self._make_batch_id(file_paths)

        # Проверяем чекпоинт
        start_idx = 0
        if resume:
            checkpoint = self._load_checkpoint(batch_id)
            if checkpoint:
                logger.info(
                    "Resuming batch %s from file %d/%d (page %d)",
                    batch_id, checkpoint.file_index, total, checkpoint.page_number,
                )
                start_idx = checkpoint.file_index

        results: List[ParseResult] = []

        for i, file_path in enumerate(file_paths):
            if i < start_idx:
                continue

            logger.info("Batch [%d/%d]: %s", i + 1, total, os.path.basename(file_path))

            try:
                result = await self.parse_file(file_path, project_id)
                results.append(result)

                if result.errors:
                    logger.warning(
                        "File %s had %d errors: %s",
                        file_path, len(result.errors), "; ".join(result.errors[:3]),
                    )
            except Exception as e:
                logger.error("Batch parse failed for %s: %s", file_path, e)
                results.append(ParseResult(
                    file_path=file_path,
                    file_type=detect_file_type(file_path),
                    chunks=[],
                    total_pages=0,
                    total_chars=0,
                    methods_used=[],
                    duration_sec=0,
                    errors=[str(e)],
                ))

            # Чекпоинт
            self._save_checkpoint(Checkpoint(
                batch_id=batch_id,
                file_index=i + 1,
                page_number=0,
                processed_files=file_paths[: i + 1],
                total_files=total,
                created_at=datetime.utcnow().isoformat(),
            ))

            if on_progress:
                on_progress(i + 1, total)

        # Очищаем чекпоинт после успешного завершения
        self._clear_checkpoint(batch_id)
        logger.info("Batch complete: %d files, %d errors", 
                     len(results), sum(1 for r in results if r.errors))

        return results

    async def parse_file(self, file_path: str, project_id: int = 0) -> ParseResult:
        """
        Полный парсинг одного файла (собирает все чанки).

        Returns:
            ParseResult с полным списком чанков
        """
        file_type = detect_file_type(file_path)
        start_time = time.time()
        chunks: List[ParsedChunk] = []
        methods: set = set()
        errors: List[str] = []

        try:
            async for chunk in self.parse_stream(file_path, project_id):
                chunks.append(chunk)
                methods.add(chunk.method)
                if chunk.metadata.get("error"):
                    errors.append(
                        f"Page {chunk.page}: {chunk.metadata['error']}"
                    )
        except Exception as e:
            errors.append(str(e))

        elapsed = time.time() - start_time
        total_pages = max((c.page for c in chunks), default=0)
        total_chars = sum(len(c.content) for c in chunks)

        # Извлечение метаданных
        metadata = self._extract_metadata(file_path, file_type, chunks)

        return ParseResult(
            file_path=file_path,
            file_type=file_type,
            chunks=chunks,
            total_pages=total_pages,
            total_chars=total_chars,
            methods_used=sorted(methods),
            duration_sec=round(elapsed, 2),
            errors=errors,
            metadata=metadata,
        )

    # -------------------------------------------------------------------------
    # Format-Specific Parsers
    # -------------------------------------------------------------------------

    async def _parse_pdf_stream(
        self, file_path: str, use_ocr: bool
    ) -> AsyncIterator[ParsedChunk]:
        """Поточный парсинг PDF с Vision OCR fallback."""
        import fitz

        doc = fitz.open(file_path)
        total_pages = doc.page_count

        try:
            for page_num in range(total_pages):
                page = doc[page_num]
                text = page.get_text().strip()

                if text:
                    # Stage 1: чистый текст
                    if len(text) > self.MAX_CHUNK_CHARS:
                        text = text[: self.MAX_CHUNK_CHARS] + "\n... [TRUNCATED]"
                    yield ParsedChunk(
                        content=text,
                        page=page_num + 1,
                        total_pages=total_pages,
                        method=ParseMethod.PYMFIT_TEXT.value,
                        metadata={"source": file_path},
                        confidence=0.95,
                    )
                elif use_ocr:
                    # Stage 2: Vision OCR
                    logger.info(
                        "PDF page %d/%d: no text, OCR fallback", 
                        page_num + 1, total_pages
                    )
                    ocr_text = await self._vision_ocr_page(page)
                    yield ParsedChunk(
                        content=ocr_text,
                        page=page_num + 1,
                        total_pages=total_pages,
                        method=ParseMethod.VISION_OCR.value,
                        metadata={"source": file_path},
                        confidence=0.7,
                        ocr_fallback=True,
                    )
                else:
                    yield ParsedChunk(
                        content=f"[EMPTY PAGE {page_num + 1} — OCR disabled]",
                        page=page_num + 1,
                        total_pages=total_pages,
                        method=ParseMethod.PYMFIT_TEXT.value,
                        confidence=0.0,
                    )
        finally:
            doc.close()

    async def _parse_docx_stream(
        self, file_path: str
    ) -> AsyncIterator[ParsedChunk]:
        """Поточный парсинг DOCX."""
        import docx as docx_lib

        doc = docx_lib.Document(file_path)

        # Собираем параграфы
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        # Добавляем таблицы
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(" | ".join(cells))
            text += "\n\n[TABLE]\n" + "\n".join(rows)

        if not text:
            text = "[EMPTY DOCUMENT]"

        # Разбиваем на страницы (приблизительно: ~3000 символов = 1 стр.)
        chars_per_page = 3000
        total_pages = max(1, len(text) // chars_per_page + 1)

        for i in range(total_pages):
            start = i * chars_per_page
            end = start + chars_per_page
            chunk_text = text[start:end]

            yield ParsedChunk(
                content=chunk_text,
                page=i + 1,
                total_pages=total_pages,
                method=ParseMethod.PYTHON_DOCX.value,
                metadata={"source": file_path},
                confidence=0.95,
            )

    def _parse_xlsx(self, file_path: str) -> List[ParsedChunk]:
        """Парсинг XLSX."""
        import openpyxl

        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
        except Exception as e:
            return [
                ParsedChunk(
                    content=f"[XLSX PARSE ERROR: {e}]",
                    method=ParseMethod.OPENPYXL.value,
                    confidence=0.0,
                    metadata={"error": str(e)},
                )
            ]

        chunks = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows_data = []

            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_str = " | ".join(str(c) if c is not None else "" for c in row)
                    rows_data.append(row_str)

            if rows_data:
                content = f"[SHEET: {sheet_name}]\n" + "\n".join(rows_data)
                chunks.append(ParsedChunk(
                    content=content,
                    page=len(chunks) + 1,
                    total_pages=len(wb.sheetnames),
                    method=ParseMethod.OPENPYXL.value,
                    metadata={"source": file_path, "sheet": sheet_name},
                    confidence=0.9,
                ))

        wb.close()
        return chunks

    async def _parse_image_ocr(self, file_path: str) -> ParsedChunk:
        """OCR изображения через Vision-модель."""
        logger.info("Image OCR: %s", os.path.basename(file_path))

        # На Mac Studio — Gemma 4 31B VLM, на dev — minicpm-v через Ollama
        try:
            ocr_text = await self._run_vision_ocr(file_path)
            return ParsedChunk(
                content=ocr_text,
                page=1,
                total_pages=1,
                method=ParseMethod.VISION_OCR.value,
                metadata={"source": file_path},
                confidence=0.7,
            )
        except Exception as e:
            logger.error("Image OCR failed: %s", e)
            return ParsedChunk(
                content=f"[OCR FAILED: {e}]",
                method=ParseMethod.RAW_BINARY.value,
                confidence=0.0,
                metadata={"error": str(e)},
            )

    def _parse_dwg(self, file_path: str) -> ParsedChunk:
        """Парсинг DWG через ODA File Converter (если доступен)."""
        # Проверяем наличие ODAFileConverter
        oda_path = os.environ.get("ODA_CONVERTER_PATH", "ODAFileConverter")
        dwg_info = f"[DWG FILE: {os.path.basename(file_path)}]"

        try:
            import subprocess
            result = subprocess.run(
                ["which", oda_path], capture_output=True, text=True, timeout=5
            )
            if result.returncode == 0:
                return ParsedChunk(
                    content=f"{dwg_info}\nODA File Converter available. Conversion to PDF required for parsing.",
                    method=ParseMethod.ODA_CONVERTER.value,
                    confidence=0.5,
                )
        except Exception:
            pass

        return ParsedChunk(
            content=f"{dwg_info}\n[ODA File Converter not available. Install for DWG support.]",
            method=ParseMethod.RAW_BINARY.value,
            confidence=0.0,
        )

    # -------------------------------------------------------------------------
    # Vision OCR
    # -------------------------------------------------------------------------

    async def _vision_ocr_page(self, page) -> str:
        """
        OCR одной страницы PDF через Vision-модель.

        На Mac Studio: Gemma 4 31B VLM (встроенное зрение).
        На dev: minicpm-v через Ollama.
        """
        try:
            # Рендерим страницу в изображение
            pix = page.get_pixmap(dpi=150)
            image_bytes = pix.tobytes("png")

            # Отправляем в LLM Engine (vision endpoint)
            from src.core.llm_engine import llm_engine

            result = await llm_engine.vision(
                "pto",
                image_bytes,
                "Извлеки ВЕСЬ текст с этой страницы документа. "
                "Сохрани форматирование таблиц и списков. "
                "Если это чертёж — опиши все размеры, метки и обозначения.",
            )
            return result or "[OCR: NO TEXT EXTRACTED]"

        except ImportError:
            logger.warning("llm_engine not available for OCR — returning placeholder")
            return "[OCR: LLM ENGINE NOT AVAILABLE]"
        except Exception as e:
            logger.error("Vision OCR failed: %s", e)
            return f"[OCR ERROR: {e}]"

    async def _run_vision_ocr(self, image_path: str) -> str:
        """OCR отдельного файла-изображения."""
        try:
            with open(image_path, "rb") as f:
                image_bytes = f.read()

            from src.core.llm_engine import llm_engine

            result = await llm_engine.vision(
                "pto",
                image_bytes,
                "Извлеки ВЕСЬ текст с этого изображения. Это может быть скан документа, "
                "чертёж или фотография. Извлеки все видимые данные: текст, цифры, подписи.",
            )
            return result or "[OCR: NO TEXT EXTRACTED]"
        except Exception as e:
            return f"[OCR ERROR: {e}]"

    # -------------------------------------------------------------------------
    # Metadata Extraction
    # -------------------------------------------------------------------------

    def _extract_metadata(
        self, file_path: str, file_type: FileType, chunks: List[ParsedChunk]
    ) -> Dict[str, Any]:
        """Извлечь метаданные документа."""
        path = Path(file_path)
        stat = path.stat() if path.exists() else None

        metadata = {
            "filename": path.name,
            "file_type": file_type.value,
            "extension": path.suffix.lower(),
            "size_bytes": stat.st_size if stat else 0,
            "size_mb": round(stat.st_size / 1024 / 1024, 2) if stat else 0,
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat() if stat else None,
            "total_chars": sum(len(c.content) for c in chunks),
            "has_ocr": any(c.ocr_fallback for c in chunks),
            "hash_sha256": self._file_hash(file_path),
        }

        # PDF метаданные
        if file_type == FileType.PDF:
            try:
                import fitz
                doc = fitz.open(file_path)
                pdf_meta = doc.metadata or {}
                metadata["pdf_author"] = pdf_meta.get("author", "")
                metadata["pdf_title"] = pdf_meta.get("title", "")
                metadata["pdf_created"] = pdf_meta.get("creationDate", "")
                metadata["pdf_pages"] = doc.page_count
                doc.close()
            except Exception:
                pass

        return metadata

    @staticmethod
    def _file_hash(file_path: str, algo: str = "sha256") -> str:
        """Хеш файла для дедупликации."""
        try:
            h = hashlib.new(algo)
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(8192), b""):
                    h.update(chunk)
            return h.hexdigest()[:16]
        except Exception:
            return ""

    # -------------------------------------------------------------------------
    # Checkpoint Management
    # -------------------------------------------------------------------------

    def _make_batch_id(self, file_paths: List[str]) -> str:
        """Сгенерировать ID батча."""
        key = "|".join(sorted(file_paths))
        return hashlib.md5(key.encode()).hexdigest()[:8]

    def _checkpoint_path(self, batch_id: str) -> Path:
        return self._checkpoint_dir / f"batch_{batch_id}.json"

    def _save_checkpoint(self, cp: Checkpoint) -> None:
        path = self._checkpoint_path(cp.batch_id)
        data = {
            "batch_id": cp.batch_id,
            "file_index": cp.file_index,
            "page_number": cp.page_number,
            "processed_files": cp.processed_files,
            "total_files": cp.total_files,
            "created_at": cp.created_at,
        }
        with open(path, "w") as f:
            json.dump(data, f)

    def _load_checkpoint(self, batch_id: str) -> Optional[Checkpoint]:
        path = self._checkpoint_path(batch_id)
        if not path.exists():
            return None
        try:
            with open(path) as f:
                data = json.load(f)
            return Checkpoint(
                batch_id=data["batch_id"],
                file_index=data["file_index"],
                page_number=data.get("page_number", 0),
                processed_files=data["processed_files"],
                total_files=data["total_files"],
                created_at=data["created_at"],
            )
        except Exception as e:
            logger.warning("Failed to load checkpoint: %s", e)
            return None

    def _clear_checkpoint(self, batch_id: str) -> None:
        path = self._checkpoint_path(batch_id)
        if path.exists():
            path.unlink()
            logger.debug("Checkpoint cleared: %s", batch_id)


# =============================================================================
# Module-level instances
# =============================================================================

streaming_parser = StreamingParser()

# Сохраняем старый интерфейс для обратной совместимости
parser_engine = streaming_parser
