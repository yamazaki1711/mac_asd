"""
Extract structured knowledge from 31 id-prosto.ru DOCX checklists.

Produces: data/knowledge/idprosto_worktype_docs.json
  - 31 work type entries
  - ~569 document rows
  - ~252 unique normative references
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from collections import defaultdict
from typing import List, Dict

try:
    import docx
except ImportError:
    print("ERROR: python-docx is required. Install: pip install python-docx")
    sys.exit(1)

LISTS_DIR = Path(r"C:\idprosto\downloads_idprosto_lists")
FORMS_DIR = Path(r"C:\idprosto\downloads_idprosto_forms")
OUTPUT_PATH = Path(r"C:\MAC_ASD\data\knowledge\idprosto_worktype_docs.json")

# Mapping from DOCX filenames to work type codes (same as id-prosto slugs)
SLUG_MAP = {
    "01_permit": "01_permit",
    "02_geodethic": "02_geodetic",
    "03_earth": "03_earth",
    "04_piles": "04_piles",
    "05_bored-piles": "05_bored-piles",
    "06_concrete": "06_concrete",
    "07_steel": "07_steel",
    "08_precast-concrete": "08_precast-concrete",
    "09_finishing-works": "09_finishing-works",
    "10_anticorrosive": "10_anticorrosive",
    "11_extwatersupply": "11_extwatersupply",
    "12_extsewerage": "12_extsewerage",
    "13_drilling": "13_drilling",
    "14_intwatersupply": "14_intwatersupply",
    "15_intsewerage": "15_intsewerage",
    "16_heating": "16_heating",
    "17_ventilation": "17_ventilation",
    "18_extinguishing": "18_extinguishing",
    "19_pipelines": "19_pipelines",
    "20_equipment": "20_equipment",
    "21_tanks": "21_tanks",
    "22_electric": "22_electric",
    "23_extelectric": "23_extelectric",
    "24_automatic": "24_automatic",
    "25_fire-alarm": "25_fire-alarm",
    "26_sks": "26_sks",
    "27_elevators": "27_elevators",
    "28_heat-pipelines": "28_heat-pipelines",
    "29_roads": "29_roads",
    "30_demolition": "30_demolition",
    "31_steam-boiler": "31_steam-boiler",
}


def classify_doc_category(name: str, form: str, norm: str) -> List[str]:
    """Classify a document row into categories: aosr, aook, journal, test_act, schema, certificate, other."""
    text = f"{name} {form} {norm}".lower()
    cats = []

    if any(w in text for w in ["аоср", "акт освидетельствования скрытых работ", "скрытых работ"]):
        cats.append("aosr")
    if any(w in text for w in ["аоок", "ответственных конструкций", "аоусито", "участков сетей"]):
        cats.append("aook")
    if any(w in text for w in ["журнал", "ведомость", "реестр"]):
        cats.append("journal")
    if any(w in text for w in ["протокол", "испытан", "определени", "измерен"]):
        cats.append("test_act")
    if any(w in text for w in ["исполнительн", "схем", "чертеж", "геодезическ"]):
        cats.append("schema")
    if any(w in text for w in ["паспорт", "сертификат", "декларац", "качеств", "входн",
                                 "удостоверен"]):
        cats.append("certificate")
    if any(w in text for w in ["разрешен", "допуск", "огражден", "проект производ",
                                 "приказ", "договор", "акт приём", "акт переда",
                                 "аогро", "агро", "ароокс"]):
        cats.append("other")

    return cats if cats else ["other"]


def extract_normative_refs(text: str) -> List[str]:
    """Extract individual normative references from a text cell."""
    refs = set()

    # Match patterns like: СП 70.13330.2025, ГОСТ Р 51872-2024, Приказ Минстроя №344/пр
    patterns = [
        r'(СП\s*[\d.]+[\w\d.-]*)',
        r'(ГОСТ\s*(?:Р\s*)?[\d.]+[\w\d.-]*)',
        r'(Приказ\s+Минстроя\s*№?\s*[\d/а-я]+(?:/\d+)?)',
        r'(ПП\s*РФ\s*№?\s*[\d]+)',
        r'(РД[\s-]*[\d.]+[\w\d.-]*)',
        r'(ВСН\s*[\d.]+[\w\d.-]*)',
        r'(СНиП\s*[\d.]+[\w\d.-]*)',
        r'(И\s*[\d.]+[\w\d.-]*)',
        r'(СП\s*\d{2}\.\d{7}\.\d{4})',  # СП 543.1325800.2024 format
    ]

    for pat in patterns:
        for match in re.findall(pat, text, re.IGNORECASE):
            refs.add(match.strip())

    return sorted(refs)


def extract_checklist(filepath: Path, code: str) -> dict:
    """Extract structured data from a single DOCX checklist."""
    doc = docx.Document(str(filepath))

    # Title
    title_text = ""
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip():
            title_text = p.text.strip().replace("\n", " — ")
            break
    if not title_text:
        title_text = doc.paragraphs[0].text.strip() if doc.paragraphs else code

    # Documents table
    documents = []
    norms_set = set()

    if doc.tables:
        table = doc.tables[0]
        for row in table.rows[1:]:  # skip header
            cells = [cell.text.strip().replace("\n", " ").replace("\r", " ") for cell in row.cells]
            if len(cells) < 3:
                continue
            num = cells[0]
            name = cells[1] if len(cells) > 1 else ""
            form = cells[2] if len(cells) > 2 else ""
            norm = cells[3] if len(cells) > 3 else ""

            if not name or name == "Наименование документа":
                continue

            categories = classify_doc_category(name, form, norm)
            doc_norms = extract_normative_refs(f"{name} {form} {norm}")
            norms_set.update(doc_norms)

            # Split raw normative text on multiple spaces into individual citations
            raw_norm_texts = []
            if norm:
                raw_norm_texts = [t.strip() for t in re.split(r'\s{2,}', norm) if t.strip()]

            documents.append({
                "num": num,
                "name": name,
                "form": form,
                "normative": norm,
                "categories": categories,
                "normative_refs": doc_norms,
                "raw_norm_texts": raw_norm_texts,
            })

    # Collect ALL raw normative texts (each unique text is a normative ref)
    for d in documents:
        for t in d.get("raw_norm_texts", []):
            norms_set.add(t)

    return {
        "code": code,
        "name": title_text,
        "total_docs": len(documents),
        "documents": documents,
        "normative_refs": sorted(norms_set),
    }


def extract_norms_from_text(text: str) -> List[str]:
    """Aggressive extraction of normative references from directory/file names."""
    refs = set()
    text_clean = text.replace("_", " ").replace("-", " ")

    # СП patterns
    for m in re.findall(r'СП\s*\d{2,3}[.\d]*\d{4}', text_clean):
        refs.add(m.strip())
    # ГОСТ patterns
    for m in re.findall(r'ГОСТ\s*(?:Р\s*)?[\d.]+[\w\d.-]*', text_clean):
        refs.add(m.strip())
    # СНиП
    for m in re.findall(r'СНиП\s*[\d.]+[\w\d.-]*', text_clean):
        refs.add(m.strip())
    # Приказ Минстроя
    for m in re.findall(r'Приказ[а]?\s+Минстроя\s*(?:России\s*)?(?:№|N|от\s+)?\s*[\d/а-я]+(?:/\d+)?', text_clean):
        refs.add(m.strip())
    # ПП РФ
    for m in re.findall(r'ПП\s*РФ\s*№?\s*[\d]+', text_clean):
        refs.add(m.strip())
    # РД
    for m in re.findall(r'РД[\s-]*[\d.]+[\w\d.-]*', text_clean):
        refs.add(m.strip())
    # ВСН
    for m in re.findall(r'ВСН\s*[\d.]+[\w\d.-]*', text_clean):
        refs.add(m.strip())
    # И
    for m in re.findall(r'\bИ\s+[\d.]+[\w\d.-]*', text_clean):
        refs.add(m.strip())
    # ТР, ТС
    for m in re.findall(r'\bТ[СР]\s*[\d.]+[\w\d.-]*', text_clean):
        refs.add(m.strip())
    # Приказ Ростехнадзора
    for m in re.findall(r'Приказ[а]?\s+Ростехнадзора\s*(?:№|от\s+)?\s*[\d/а-я]+', text_clean):
        refs.add(m.strip())
    # Федеральный закон
    for m in re.findall(r'(?:ФЗ|Федеральный\s+закон)\s*(?:№|от\s+)?\s*[\d/а-я-]+', text_clean):
        refs.add(m.strip())
    # СанПиН
    for m in re.findall(r'СанПиН\s*[\d.]+[\w\d.-]*', text_clean):
        refs.add(m.strip())

    return list(refs)


def catalogue_templates() -> List[Dict]:
    """Catalogue all DOCX/XLSX templates from the forms directory."""
    templates = []
    all_template_norms = set()

    for form_dir in sorted(FORMS_DIR.iterdir()):
        if not form_dir.is_dir():
            continue
        dir_name = form_dir.name
        # Extract short name from directory name
        parts = dir_name.split("_-_")
        if len(parts) >= 1:
            short_name = parts[0].replace("_", " ").replace("01 ", "").replace("02 ", "").replace("03 ", "")
            short_name = re.sub(r'^\d+_', '', short_name)
        else:
            short_name = dir_name

        # Extract normative refs from directory name
        dir_norms = extract_norms_from_text(dir_name)

        for f in sorted(form_dir.iterdir()):
            if not f.is_file():
                continue
            if not f.suffix.lower() in ('.docx', '.xlsx'):
                continue

            # Extract normative refs from filename
            file_norms = extract_norms_from_text(f.stem)

            templates.append({
                "file_name": f.stem,
                "extension": f.suffix.lower(),
                "full_path": str(f.relative_to(LISTS_DIR.parent.parent)),
                "regulation_package": short_name[:120],
                "regulation_dir": dir_name[:200],
                "normative_refs": list(set(dir_norms + file_norms)),
            })
            all_template_norms.update(dir_norms)
            all_template_norms.update(file_norms)

    return templates, list(all_template_norms)


def main():
    print("=== Extracting id-prosto.ru Knowledge Base ===")

    # 1. Extract all 31 checklists
    work_types = {}
    all_norms = set()
    total_docs = 0

    for fname in sorted(LISTS_DIR.glob("*.docx")):
        stem = fname.stem  # e.g. "06_concrete"
        code = SLUG_MAP.get(stem, stem)
        print(f"  Extracting: {fname.name} ->{code}")

        entry = extract_checklist(fname, code)
        work_types[code] = entry
        all_norms.update(entry["normative_refs"])
        total_docs += entry["total_docs"]

    print(f"  Work types: {len(work_types)}")
    print(f"  Total documents: {total_docs}")
    print(f"  Unique normative refs: {len(all_norms)}")

    # 2. Catalogue templates
    print("\n=== Cataloguing Templates ===")
    templates, template_norms = catalogue_templates()
    all_norms.update(template_norms)
    print(f"  Templates found: {len(templates)}")
    print(f"  Template normative refs: {len(template_norms)}")

    # Count by extension
    by_ext = defaultdict(int)
    for t in templates:
        by_ext[t["extension"]] += 1
    print(f"  By type: {dict(by_ext)}")

    # 2b. Catalogue sample PDFs
    SAMPLES_DIR = Path(r"C:\idprosto\id_samples")
    samples = []
    sample_norms = set()
    if SAMPLES_DIR.exists():
        for f in sorted(SAMPLES_DIR.iterdir()):
            if f.is_file() and f.suffix.lower() == '.pdf':
                fnorms = extract_norms_from_text(f.stem)
                samples.append({
                    "file_name": f.name,
                    "full_path": str(f.relative_to(LISTS_DIR.parent.parent)),
                    "normative_refs": fnorms,
                })
                sample_norms.update(fnorms)
        all_norms.update(sample_norms)
        print(f"\n=== Cataloguing Sample PDFs ===")
        print(f"  Samples found: {len(samples)}")
        print(f"  Sample normative refs: {len(sample_norms)}")

    print(f"\n  Total unique normative refs: {len(all_norms)}")

    # 3. Build final output
    output = {
        "meta": {
            "source": "id-prosto.ru",
            "lists_dir": str(LISTS_DIR),
            "forms_dir": str(FORMS_DIR),
            "extracted_at": __import__("datetime").datetime.now().isoformat(),
            "total_work_types": len(work_types),
            "total_documents": total_docs,
            "total_normative_refs": len(all_norms),
            "total_templates": len(templates),
        },
        "work_types": work_types,
        "all_normative_refs": sorted(all_norms),
        "templates": templates,
        "samples": samples,
    }

    # 4. Write output
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    print(f"\n=== Written: {OUTPUT_PATH} ===")
    print(f"  Size: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
