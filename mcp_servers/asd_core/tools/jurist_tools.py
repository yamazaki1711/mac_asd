"""
ASD v12.0 — Jurist MCP Tools.

9 tools for the Legal Agent with real LLM logic:
  1. asd_upload_document          — Parse PDF/DOCX → chunks
  2. asd_analyze_contract         — Full Map-Reduce + БЛС legal analysis
  3. asd_normative_search         — Hybrid search (Vector + Graph)
  4. asd_generate_protocol        — Protocol of disagreements (DOCX)
  5. asd_generate_claim           — Pre-trial claim (DOCX)
  6. asd_generate_lawsuit         — Arbitration lawsuit (DOCX)
  7. asd_add_trap                 — Manual trap addition to БЛС
  8. asd_list_telegram_channels   — List cataloged Telegram channels for БЛС
  9. asd_ingest_telegram          — Ingest Telegram export to БЛС (single or batch)
"""

import os
import logging
from typing import Dict, Any, Optional, List

from src.core.services.legal_service import legal_service
from src.schemas.legal import (
    LegalAnalysisRequest,
    ReviewType,
)

logger = logging.getLogger(__name__)


async def asd_upload_document(file_path: str) -> Dict[str, Any]:
    """
    Загрузка и парсинг документа (PDF, DOCX).

    Args:
        file_path: Путь к файлу документа

    Returns:
        {"status", "file_path", "chunks_extracted", "total_chars", "message"}
    """
    logger.info(f"asd_upload_document: {file_path}")

    try:
        result = await legal_service.upload_and_parse(file_path)
        return result.model_dump()
    except Exception as e:
        logger.error(f"Failed to upload document: {e}")
        return {"status": "error", "message": str(e)}


async def asd_analyze_contract(
    document_id: Optional[int] = None,
    document_text: Optional[str] = None,
    file_path: Optional[str] = None,
    review_type: str = "contract",
    chunk_size: int = 6000,
    chunk_overlap: int = 300,
) -> Dict[str, Any]:
    """
    Юридическая экспертиза контракта (БЛС + Map-Reduce + LightRAG).

    Полный цикл анализа:
      1. Парсинг документа (если file_path)
      2. Разбивка на чанки (если документ длинный)
      3. Поиск похожих ловушек в БЛС через RAG
      4. MAP: LLM анализирует каждый чанк
      5. REDUCE: агрегация результатов → вердикт

    Args:
        document_id: ID документа в БД (если уже загружен)
        document_text: Текст документа напрямую
        file_path: Путь к файлу (PDF, DOCX)
        review_type: Тип рассмотрения (contract, tender, compliance)
        chunk_size: Размер чанка для Map-Reduce (символов)
        chunk_overlap: Перекрытие чанков

    Returns:
        Полный LegalAnalysisResult как dict:
        {"findings", "verdict", "summary", "normative_refs", "contradictions", ...}
    """
    logger.info(
        f"asd_analyze_contract: doc_id={document_id}, "
        f"text={'yes' if document_text else 'no'}, "
        f"file={file_path}, review_type={review_type}"
    )

    try:
        request = LegalAnalysisRequest(
            document_id=document_id,
            document_text=document_text,
            file_path=file_path,
            review_type=ReviewType(review_type),
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

        result = await legal_service.analyze(request)
        return result.model_dump()

    except Exception as e:
        logger.error(f"Contract analysis failed: {e}")
        return {
            "status": "error",
            "message": str(e),
            "findings": [],
            "verdict": "approved",
            "summary": f"Ошибка при анализе: {str(e)[:200]}",
        }


async def asd_normative_search(query: str) -> Dict[str, Any]:
    """
    Поиск по нормативной базе (Graph + Vector).

    Args:
        query: Поисковый запрос (например, "неустойка ФЗ-44")

    Returns:
        {"status", "query", "results": {"vector_chunks", "graph_context"}}
    """
    logger.info(f"asd_normative_search: {query}")

    try:
        from src.core.rag_service import rag_service

        results = await rag_service.hybrid_search(query)
        return {"status": "success", "query": query, "results": results}
    except Exception as e:
        logger.error(f"Normative search failed: {e}")
        return {"status": "error", "query": query, "message": str(e)}


async def asd_generate_protocol(
    analysis_result: Dict[str, Any],
    contract_number: str = "___",
    contract_date: str = "___",
    customer_name: str = "Заказчик",
    contractor_name: str = "Подрядчик",
    output_path: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Протокол разногласий к договору (DOCX).

    v12.0.0: Генерирует полноценный DOCX с 3-колоночной таблицей:
    | Пункт, статья | Редакция Заказчика | Редакция Подрядчика |

    Каждая правка Подрядчика опирается на конкретную статью закона (ГК РФ,
    ГрК РФ, ФЗ-44/223). Это заставляет даже государственного Заказчика
    вернуться в правовое поле и умерить аппетиты.

    Args:
        analysis_result: Результат asd_analyze_contract (dict с findings)
        contract_number: Номер договора
        contract_date: Дата договора
        customer_name: Наименование Заказчика
        contractor_name: Наименование Подрядчика
        output_path: Путь для сохранения DOCX (если None — auto-generate)

    Returns:
        {"status", "file_path", "total_items", "message"}
    """
    logger.info(f"asd_generate_protocol: contract={contract_number}")

    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from src.config import settings

        # Extract protocol items from analysis
        findings = analysis_result.get("findings", [])
        protocol_items = analysis_result.get("protocol_items", [])

        # Build protocol items from findings if not already provided
        if not protocol_items:
            for i, f in enumerate(findings):
                if isinstance(f, dict):
                    protocol_items.append({
                        "clause_ref": f.get("clause_ref", "—"),
                        "customer_text": f.get("contract_quote", f.get("issue", "—")),
                        "contractor_text": f.get("contractor_edit", f.get("recommendation", "—")),
                        "legal_basis": f.get("legal_basis", "—"),
                        "severity": f.get("severity", "medium"),
                        "blc_match": f.get("blc_match"),
                    })

        if not protocol_items:
            return {
                "status": "success",
                "message": "Нет пунктов для протокола разногласий",
                "total_items": 0,
                "file_path": None,
            }

        # Create DOCX
        doc = Document()

        # Styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

        # Title
        title = doc.add_heading('ПРОТОКОЛ РАЗНОГЛАСИЙ', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(
            f'к Договору {contract_number} от {contract_date}'
        )
        run.font.size = Pt(12)
        run.bold = True

        # Parties
        doc.add_paragraph()
        parties = doc.add_paragraph()
        parties.add_run('Заказчик: ').bold = True
        parties.add_run(customer_name + '\n')
        parties.add_run('Подрядчик: ').bold = True
        parties.add_run(contractor_name)

        doc.add_paragraph()
        note = doc.add_paragraph()
        note.add_run(
            'Настоящий протокол разногласий составлен на основании '
            'действующего законодательства РФ (ГК РФ, ГрК РФ, ФЗ-44/223). '
            'Каждая предлагаемая редакция Подрядчика опирается на конкретную '
            'норму права, что обеспечивает правовую обоснованность позиции Подрядчика.'
        ).italic = True

        doc.add_paragraph()

        # 3-column table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        header_cells = table.rows[0].cells
        headers = ['№', 'Пункт, статья\nдоговора', 'Редакция\nЗаказчика', 'Редакция\nПодрядчика']
        widths = [Cm(1.2), Cm(3.5), Cm(6.5), Cm(6.5)]

        for i, (header, width) in enumerate(zip(headers, widths)):
            header_cells[i].text = header
            header_cells[i].width = width
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for idx, item in enumerate(protocol_items, 1):
            row = table.add_row()

            row.cells[0].text = str(idx)
            row.cells[1].text = item.get("clause_ref", "—")
            row.cells[2].text = item.get("customer_text", "—")
            
            # Contractor text with legal basis
            contractor_text = item.get("contractor_text", "—")
            legal_basis = item.get("legal_basis", "")
            if legal_basis and legal_basis != "—":
                contractor_text += f"\n\n(Основание: {legal_basis})"
            row.cells[3].text = contractor_text

            # Format
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)

            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Summary
        doc.add_paragraph()
        total = len(protocol_items)
        summary = doc.add_paragraph()
        summary.add_run(
            f'Итого: {total} {"пункт" if total == 1 else "пункта" if total < 5 else "пунктов"} разногласий.'
        ).bold = True

        # Severity summary
        severities = {}
        for item in protocol_items:
            sev = item.get("severity", "medium")
            severities[sev] = severities.get(sev, 0) + 1

        sev_text = []
        if severities.get("critical"):
            sev_text.append(f"{severities['critical']} критических")
        if severities.get("high"):
            sev_text.append(f"{severities['high']} высоких")
        if severities.get("medium"):
            sev_text.append(f"{severities['medium']} средних")

        if sev_text:
            sev_para = doc.add_paragraph()
            sev_para.add_run("Критичность: ").bold = True
            sev_para.add_run(", ".join(sev_text))

        # Signature lines
        doc.add_paragraph()
        doc.add_paragraph()

        sig_table = doc.add_table(rows=3, cols=2)
        sig_table.style = 'Table Grid'
        sig_table.cell(0, 0).text = "ЗАКАЗЧИК:"
        sig_table.cell(0, 1).text = "ПОДРЯДЧИК:"
        for cell in sig_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        sig_table.cell(1, 0).text = f"\n\n\n________________ / {customer_name} /"
        sig_table.cell(1, 1).text = f"\n\n\n________________ / {contractor_name} /"
        sig_table.cell(2, 0).text = "М.П."
        sig_table.cell(2, 1).text = "М.П."

        # Save
        if not output_path:
            artifacts = settings.artifacts_path
            artifacts.mkdir(parents=True, exist_ok=True)
            output_path = str(artifacts / f"Протокол_разногласий_{contract_number}.docx")

        doc.save(output_path)

        return {
            "status": "success",
            "file_path": output_path,
            "total_items": total,
            "message": f"Протокол разногласий создан: {total} пунктов ({', '.join(sev_text)})",
        }

    except Exception as e:
        logger.error(f"Protocol generation failed: {e}")
        return {"status": "error", "message": str(e)}


async def asd_generate_claim(
    contract_id: int,
    debt_amount: float,
    works_description: str,
    works_completed_date: Optional[str] = None,
    payment_deadline: Optional[str] = None,
    penalty: Optional[float] = None,
    output_path: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Претензия при неоплате СМР (DOCX).

    Генерирует досудебную претензию с расчётом неустойки по ст. 395 ГК РФ.
    Использует ProtocolPartyInfo из contracts для автозаполнения реквизитов.

    Args:
        contract_id: ID контракта
        debt_amount: Сумма задолженности (руб.)
        works_description: Описание выполненных работ
        works_completed_date: Дата завершения работ (ISO)
        payment_deadline: Срок оплаты по договору
        penalty: Неустойка (если не указана — расчёт по ставке ЦБ)
        output_path: Путь для сохранения DOCX

    Returns:
        {"status", "file_path", "debt_amount", "penalty_amount", "total_amount",
         "claim_deadline", "parties", "message"}
    """
    logger.info(f"asd_generate_claim: contract_id={contract_id}, debt={debt_amount}")

    try:
        from datetime import date, timedelta

        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from src.config import settings

        # Look up contract for party info (graceful fallback if DB unavailable)
        contractor_name = "ООО «КСК №1»"
        contractor_inn = ""
        contractor_address = ""
        customer_name = "Заказчик"
        customer_inn = ""
        customer_address = ""
        contract_number_str = str(contract_id)
        contract_date_str = ""

        try:
            from src.db.init_db import Session
            from src.db.models import Contract
            from sqlalchemy import select

            with Session() as session:
                contract = session.execute(
                    select(Contract).where(Contract.id == contract_id)
                ).scalar_one_or_none()
                if contract:
                    contract_number_str = contract.number or str(contract_id)
                    contract_date_str = contract.date.isoformat() if contract.date else ""
                    customer_name = contract.party_1 or customer_name
                    customer_inn = contract.party_1_inn or ""
                    customer_address = contract.party_1_address or ""
                    contractor_name = contract.party_2 or contractor_name
                    contractor_inn = contract.party_2_inn or ""
                    contractor_address = contract.party_2_address or ""
        except (ImportError, Exception) as e:
            logger.debug("Contract lookup skipped (DB models not available): %s", e)

        # Calculate penalty if not provided (ЦБ РФ key rate = 21% as of 2026)
        if penalty is None:
            key_rate = 0.21
            penalty = round(debt_amount * key_rate / 365 * 30, 2)

        total_amount = debt_amount + penalty
        claim_deadline = (date.today() + timedelta(days=30)).isoformat()

        # Create DOCX
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.add_run(f"Исх. № ___ от {date.today().strftime('%d.%m.%Y')}").font.size = Pt(10)

        doc.add_paragraph()
        recipient = doc.add_paragraph()
        recipient.add_run(f"Кому: {customer_name}\n").bold = True
        if customer_address:
            recipient.add_run(f"Адрес: {customer_address}\n")
        if customer_inn:
            recipient.add_run(f"ИНН: {customer_inn}")

        # Title
        doc.add_paragraph()
        title = doc.add_heading('ДОСУДЕБНАЯ ПРЕТЕНЗИЯ', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Body
        doc.add_paragraph()
        body = doc.add_paragraph()
        body.add_run(
            f'Между {contractor_name} (Подрядчик) и {customer_name} (Заказчик) '
            f'заключён договор подряда № {contract_number_str}'
        )
        if contract_date_str:
            body.add_run(f' от {contract_date_str}')
        body.add_run(
            f', предметом которого является выполнение строительно-монтажных работ.\n\n'
            f'Подрядчик надлежащим образом выполнил следующие работы: '
            f'{works_description}.\n\n'
        )
        if works_completed_date:
            body.add_run(f'Работы завершены {works_completed_date}. ')
        body.add_run(
            f'Заказчик в нарушение условий договора и ст. 746 ГК РФ '
            f'не произвёл оплату выполненных работ.\n\n'
        )

        # Amounts table
        doc.add_paragraph()
        amt_table = doc.add_table(rows=4, cols=2)
        amt_table.style = 'Table Grid'
        amounts = [
            ("Сумма основного долга", f"{debt_amount:,.2f} руб."),
            ("Неустойка (ст. 395 ГК РФ)", f"{penalty:,.2f} руб."),
            ("Итого к оплате", f"{total_amount:,.2f} руб."),
            ("Срок ответа на претензию", f"до {claim_deadline} (30 дней)"),
        ]
        for i, (label, value) in enumerate(amounts):
            amt_table.cell(i, 0).text = label
            amt_table.cell(i, 1).text = value
            for cell in amt_table.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
            if i == 2:
                for cell in amt_table.rows[i].cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True

        # Legal basis
        doc.add_paragraph()
        legal = doc.add_paragraph()
        legal.add_run('Правовое основание: ').bold = True
        legal.add_run(
            'ст. 309, 395, 708, 711, 746 ГК РФ, ст. 4 АПК РФ. '
        )

        # Demand
        doc.add_paragraph()
        demand = doc.add_paragraph()
        demand.add_run('На основании изложенного ТРЕБУЮ:').bold = True
        demands_list = [
            f'1. Оплатить задолженность в размере {debt_amount:,.2f} руб.',
            f'2. Уплатить неустойку в размере {penalty:,.2f} руб.',
            '3. В случае неудовлетворения претензии Подрядчик будет вынужден обратиться '
            'в Арбитражный суд с отнесением всех судебных расходов на Заказчика.',
        ]
        for d in demands_list:
            dp = doc.add_paragraph(d)

        # Signature
        doc.add_paragraph()
        doc.add_paragraph()
        sig = doc.add_paragraph()
        sig.add_run(f"{contractor_name}\n").bold = True
        sig.add_run("Генеральный директор _______________ /_______________/\n")
        sig.add_run("М.П.")

        # Save
        if not output_path:
            artifacts = settings.artifacts_path
            artifacts.mkdir(parents=True, exist_ok=True)
            ts = date.today().strftime('%Y%m%d')
            output_path = str(
                artifacts / f"Претензия_контракт_{contract_id}_{ts}.docx"
            )

        doc.save(output_path)

        return {
            "status": "success",
            "file_path": output_path,
            "debt_amount": debt_amount,
            "penalty_amount": penalty,
            "total_amount": total_amount,
            "claim_deadline": claim_deadline,
            "parties": {
                "customer": {"name": customer_name, "inn": customer_inn},
                "contractor": {"name": contractor_name, "inn": contractor_inn},
            },
            "message": f"Претензия сгенерирована. Срок ответа: до {claim_deadline}",
        }

    except Exception as e:
        logger.error(f"Claim generation failed: {e}")
        return {"status": "error", "message": str(e)}


async def asd_generate_lawsuit(
    contract_id: int,
    claim_id: int,
    court: Optional[str] = None,
    additional_amounts: Optional[List[Dict[str, Any]]] = None,
    output_path: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Исковое заявление в арбитражный суд (DOCX).

    Формируется на базе неудовлетворённой претензии с автоматическим
    извлечением юридически значимых данных из контракта и претензии.

    Args:
        contract_id: ID контракта
        claim_id: ID претензии (должна быть неудовлетворена)
        court: Подсудность (если не указана — берётся из договора)
        additional_amounts: Доп. суммы [{label, amount}, ...]
        output_path: Путь для сохранения DOCX

    Returns:
        {"status", "file_path", "court", "plaintiff", "defendant",
         "claim_amount", "attachments", "message"}
    """
    logger.info(f"asd_generate_lawsuit: contract_id={contract_id}, claim_id={claim_id}")

    try:
        from datetime import date

        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from src.config import settings

        # Look up contract and claim (graceful fallback if DB unavailable)
        contractor_name = "ООО «КСК №1»"
        contractor_inn = ""
        contractor_ogrn = ""
        contractor_address = ""
        customer_name = "Ответчик"
        customer_inn = ""
        customer_ogrn = ""
        customer_address = ""
        contract_number_str = str(contract_id)
        contract_date_str = ""
        claim_amount = 0.0
        debt_amount = 0.0
        penalty_amount = 0.0

        try:
            from src.db.init_db import Session
            from src.db.models import Contract, Claim
            from sqlalchemy import select

            with Session() as session:
                contract = session.execute(
                    select(Contract).where(Contract.id == contract_id)
                ).scalar_one_or_none()
                if contract:
                    contract_number_str = contract.number or str(contract_id)
                    contract_date_str = contract.date.isoformat() if contract.date else ""
                    customer_name = contract.party_1 or customer_name
                    customer_inn = contract.party_1_inn or ""
                    customer_ogrn = contract.party_1_ogrn or ""
                    customer_address = contract.party_1_address or ""
                    contractor_name = contract.party_2 or contractor_name
                    contractor_inn = contract.party_2_inn or ""
                    contractor_ogrn = contract.party_2_ogrn or ""
                    contractor_address = contract.party_2_address or ""
                    if not court:
                        court = contract.court or "Арбитражный суд г. Москвы"

                if claim_id:
                    claim = session.execute(
                        select(Claim).where(Claim.id == claim_id)
                    ).scalar_one_or_none()
                    if claim:
                        debt_amount = float(claim.debt_amount or 0)
                        penalty_amount = float(claim.penalty_amount or 0)
                        claim_amount = debt_amount + penalty_amount
        except (ImportError, Exception) as e:
            logger.debug("DB lookup skipped for lawsuit (models not available): %s", e)

        if not court:
            court = "Арбитражный суд г. Москвы"

        # Additional amounts
        legal_costs = 0.0
        extra_items = []
        if additional_amounts:
            for amt in additional_amounts:
                a = float(amt.get("amount", 0))
                legal_costs += a
                extra_items.append((amt.get("label", "—"), a))

        # State duty (ст. 333.21 НК РФ)
        if claim_amount <= 100_000:
            state_duty = max(claim_amount * 0.04, 2000)
        elif claim_amount <= 200_000:
            state_duty = 4000 + (claim_amount - 100_000) * 0.03
        elif claim_amount <= 1_000_000:
            state_duty = 7000 + (claim_amount - 200_000) * 0.02
        elif claim_amount <= 2_000_000:
            state_duty = 23000 + (claim_amount - 1_000_000) * 0.01
        else:
            state_duty = 33000 + (claim_amount - 2_000_000) * 0.005
            state_duty = min(state_duty, 200_000)
        state_duty = round(state_duty, 2)

        total_claim = claim_amount + legal_costs + state_duty

        # Create DOCX
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Court header
        court_para = doc.add_paragraph()
        court_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        court_para.add_run(f"В {court}")

        # Parties
        doc.add_paragraph()
        parties_header = doc.add_paragraph()
        parties_header.add_run("ИСТЕЦ: ").bold = True
        parties_header.add_run(f"{contractor_name}")
        doc.add_paragraph(f"ИНН: {contractor_inn}")
        doc.add_paragraph(f"ОГРН: {contractor_ogrn}")
        doc.add_paragraph(f"Адрес: {contractor_address}")

        doc.add_paragraph()
        resp_header = doc.add_paragraph()
        resp_header.add_run("ОТВЕТЧИК: ").bold = True
        resp_header.add_run(f"{customer_name}")
        doc.add_paragraph(f"ИНН: {customer_inn}")
        doc.add_paragraph(f"ОГРН: {customer_ogrn}")
        doc.add_paragraph(f"Адрес: {customer_address}")

        # Title
        doc.add_paragraph()
        title = doc.add_heading('ИСКОВОЕ ЗАЯВЛЕНИЕ', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(
            f'о взыскании задолженности по договору подряда '
            f'№ {contract_number_str} от {contract_date_str}'
        ).font.size = Pt(12)

        # Claim amounts
        doc.add_paragraph()
        doc.add_heading('Цена иска:', level=2)
        rows_data = [
            ("Основной долг", f"{debt_amount:,.2f} руб."),
            ("Неустойка (ст. 395 ГК РФ)", f"{penalty_amount:,.2f} руб."),
        ]
        for label, amount in extra_items:
            rows_data.append((label, f"{amount:,.2f} руб."))
        rows_data.append(("Госпошлина (ст. 333.21 НК РФ)", f"{state_duty:,.2f} руб."))
        rows_data.append(("ИТОГО", f"{total_claim:,.2f} руб."))

        amt_table = doc.add_table(rows=len(rows_data), cols=2)
        amt_table.style = 'Table Grid'
        for i, (label, value) in enumerate(rows_data):
            amt_table.cell(i, 0).text = label
            amt_table.cell(i, 1).text = value
            for cell in amt_table.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
            if i == len(rows_data) - 1:
                for cell in amt_table.rows[i].cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True

        # Factual background
        doc.add_paragraph()
        doc.add_heading('Обстоятельства дела:', level=2)
        facts = doc.add_paragraph()
        facts.add_run(
            f'{contractor_name} (Подрядчик) и {customer_name} (Заказчик) заключили '
            f'договор подряда № {contract_number_str} от {contract_date_str}.\n\n'
            f'Подрядчик надлежащим образом выполнил строительно-монтажные работы, '
            f'что подтверждается актами о приёмке выполненных работ (КС-2), '
            f'справками о стоимости (КС-3), актами освидетельствования скрытых работ '
            f'(АОСР) и записями в Общем журнале работ.\n\n'
            f'Заказчик в нарушение ст. 309, 711, 746 ГК РФ не произвёл оплату '
            f'выполненных работ. Досудебная претензия (исх. от {date.today().strftime("%d.%m.%Y")}) '
            f'оставлена без удовлетворения.\n\n'
            f'Согласно ст. 395 ГК РФ на сумму долга подлежат начислению проценты '
            f'за пользование чужими денежными средствами.'
        )

        # Legal basis
        doc.add_paragraph()
        doc.add_heading('Правовое основание:', level=2)
        doc.add_paragraph('ст. 309, 310, 395, 702, 711, 746 ГК РФ')
        doc.add_paragraph('ст. 4, 27, 35, 125, 126 АПК РФ')
        doc.add_paragraph('ст. 333.21 НК РФ')

        # Demands
        doc.add_paragraph()
        doc.add_heading('Прошу суд:', level=2)
        demands = [
            f'1. Взыскать с {customer_name} в пользу {contractor_name} '
            f'задолженность в размере {debt_amount:,.2f} руб.',
            f'2. Взыскать неустойку по ст. 395 ГК РФ в размере {penalty_amount:,.2f} руб.',
        ]
        for label, amount in extra_items:
            demands.append(
                f'{len(demands) + 1}. Взыскать {label.lower()} в размере {amount:,.2f} руб.'
            )
        demands.append(
            f'{len(demands) + 1}. Взыскать расходы по уплате государственной пошлины '
            f'в размере {state_duty:,.2f} руб.'
        )
        for d in demands:
            doc.add_paragraph(d)

        # Attachments
        doc.add_paragraph()
        doc.add_heading('Приложения:', level=2)
        attachments = [
            f'1. Копия договора № {contract_number_str}',
            '2. Копии актов выполненных работ (КС-2)',
            '3. Копии справок о стоимости (КС-3)',
            '4. Копия досудебной претензии',
            '5. Доказательства направления претензии ответчику',
            '6. Расчёт исковых требований',
            '7. Копия свидетельства о регистрации истца',
            f'8. Квитанция об уплате госпошлины ({state_duty:,.2f} руб.)',
        ]
        for a in attachments:
            doc.add_paragraph(a)

        # Signature
        doc.add_paragraph()
        doc.add_paragraph()
        sig = doc.add_paragraph()
        sig.add_run(f"Генеральный директор\n{contractor_name}\n").bold = True
        sig.add_run("_______________ /_______________/\n")
        sig.add_run(f"Дата: {date.today().strftime('%d.%m.%Y')}")

        # Save
        if not output_path:
            artifacts = settings.artifacts_path
            artifacts.mkdir(parents=True, exist_ok=True)
            ts = date.today().strftime('%Y%m%d')
            output_path = str(
                artifacts / f"Иск_контракт_{contract_id}_{ts}.docx"
            )

        doc.save(output_path)

        return {
            "status": "success",
            "file_path": output_path,
            "court": court,
            "plaintiff": contractor_name,
            "defendant": customer_name,
            "claim_amount": total_claim,
            "attachments": [a.split(". ", 1)[1] if ". " in a else a for a in attachments],
            "message": f"Исковое заявление сгенерировано. Проверьте перед подачей.",
        }

    except Exception as e:
        logger.error(f"Lawsuit generation failed: {e}")
        return {"status": "error", "message": str(e)}


async def asd_add_trap(
    title: str,
    description: str,
    source: str,
    mitigation: str,
    domain: str = "legal",
    channel: Optional[str] = None,
    category: Optional[str] = None,
    court_cases: Optional[List[str]] = None,
) -> Dict[str, Any]:
    """
    Ручное добавление доменной ловушки в БЛС.

    v12.0.0: Обобщено с LegalTrap до DomainTrap — поддержка всех доменов.
    v12.0.0: Добавлен параметр domain (legal, pto, smeta, logistics, procurement).

    Args:
        title: Краткое название ловушки
        description: Полное описание
        source: Источник (например, "Telegram @advokatgrikevich", "Internal")
        mitigation: Рекомендация по защите
        domain: Домен агента (legal, pto, smeta, logistics, procurement)
        channel: Username Telegram-канала (например, "advokatgrikevich")
        category: Категория источника
        court_cases: Список судебных дел (например, ["А40-123/2023"])

    Returns:
        {"status", "message", "trap_id", "domain"}
    """
    logger.info(f"asd_add_trap: {title} (domain={domain})")

    try:
        from src.db.init_db import SessionLocal
        from src.db.models import DomainTrap
        from src.core.llm_engine import llm_engine

        db = SessionLocal()
        try:
            # Get embedding via llm_engine
            embedding = await llm_engine.embed(description)

            # Determine weight from category
            weight_map = {
                "legal_practice": 100, "pto_practice": 100, "smeta_practice": 100,
                "legal_association": 80, "pto_news": 60, "smeta_news": 60,
                "legal_education": 60, "logistics_practice": 80, "procurement_practice": 80,
                "legal_news": 40, "logistics_news": 40, "procurement_news": 40,
                "unknown": 50,
            }
            weight = weight_map.get(category, 50) if category else 50

            trap = DomainTrap(
                domain=domain,
                title=title,
                description=description,
                source=source,
                channel=channel or source,
                category=category or "unknown",
                weight=weight,
                court_cases=court_cases or [],
                mitigation=mitigation,
                embedding=embedding,
            )
            db.add(trap)
            db.commit()
            db.refresh(trap)

            return {
                "status": "success",
                "message": f"Trap '{title}' added to domain={domain} (id={trap.id}, category={category or 'unknown'}, weight={weight}).",
                "trap_id": trap.id,
                "domain": domain,
            }
        except Exception as e:
            db.rollback()
            raise
        finally:
            db.close()

    except Exception as e:
        logger.error(f"Failed to add trap: {e}")
        return {"status": "error", "message": str(e)}


async def asd_list_telegram_channels(
    priority: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Вывод каталога Telegram-каналов для БЛС.

    v12.0.0: Показывает все каналы из config/telegram_channels.yaml
    с метаданными (категория, приоритет, фокус-области).

    Args:
        priority: Фильтр по приоритету (critical, high, medium, low).
                  Если None — показать все каналы.

    Returns:
        {"status", "channels": [...], "total_count"}
    """
    logger.info(f"asd_list_telegram_channels: priority={priority}")

    try:
        from src.scripts.ingest_blc_telegram import ChannelCatalog

        catalog = ChannelCatalog()

        # Collect unique channels
        seen = set()
        channels = []
        for key, ch in catalog.channels.items():
            username = ch.get("username", "")
            if username in seen:
                continue
            seen.add(username)

            if priority and ch.get("priority") != priority:
                continue

            channels.append({
                "username": username,
                "display_name": ch.get("display_name", ""),
                "url": ch.get("url", f"https://t.me/{username}"),
                "category": ch.get("category", "unknown"),
                "priority": ch.get("priority", "low"),
                "focus_areas": ch.get("focus_areas", []),
                "description": ch.get("description", ""),
            })

        # Sort by priority
        priority_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
        channels.sort(key=lambda x: priority_order.get(x["priority"], 4))

        return {
            "status": "success",
            "channels": channels,
            "total_count": len(channels),
        }

    except Exception as e:
        logger.error(f"Failed to list channels: {e}")
        return {"status": "error", "message": str(e), "channels": []}


async def asd_ingest_telegram(
    file_path: Optional[str] = None,
    batch_dir: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Импорт Telegram JSON-экспорта в БЛС (одиночный или пакетный).

    v12.0.0: Поддерживает каталог каналов для автокатегоризации
    и приоритизации ловушек.

    Args:
        file_path: Путь к result.json (одиночный импорт)
        batch_dir: Путь к папке с JSON-экспортами (пакетный импорт)

    Returns:
        {"status", "results": {...}, "total_traps_found"}
    """
    logger.info(f"asd_ingest_telegram: file={file_path}, batch_dir={batch_dir}")

    try:
        from src.scripts.ingest_blc_telegram import (
            ChannelCatalog,
            parse_telegram_export,
            batch_parse_exports,
        )

        catalog = ChannelCatalog()

        if batch_dir:
            results = await batch_parse_exports(batch_dir, catalog)
            total_traps = sum(s.get("trapped", 0) for s in results.values())
            return {
                "status": "success",
                "mode": "batch",
                "results": results,
                "total_traps_found": total_traps,
            }
        elif file_path:
            stats = await parse_telegram_export(file_path, catalog)
            return {
                "status": "success",
                "mode": "single",
                "results": {os.path.basename(file_path): stats},
                "total_traps_found": stats.get("trapped", 0),
            }
        else:
            return {
                "status": "error",
                "message": "Specify either file_path or batch_dir",
            }

    except Exception as e:
        logger.error(f"Failed to ingest telegram: {e}")
        return {"status": "error", "message": str(e)}
